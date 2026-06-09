"""Soften the +15.108d framing in the Abstract."""
from docx import Document
from copy import deepcopy
from lxml import etree

DOC_PATH = '/home/user/workspace/RiceBaCI-GEE/manuscript/Manuscript.docx'
doc = Document(DOC_PATH)
p = doc.paragraphs[14]

# Each run holds part of the text. Strategy: rebuild paragraph as a single run with replaced text.
old_text = p.text

# Targeted replacement:
old_chunk = "+15.108 d (WCR p = 0.4065) after correction — directionally confirming the pre-registered prediction τ_raw > τ_corrected > 0 while the 95% CI still includes zero."
new_chunk = ("+15.108 d (WCR p = 0.4065) after correction — directionally consistent with the pre-registered prediction τ_raw > τ_corrected > 0 but indicative rather than confirmatory given that the 95 % CI brackets zero and that diagnostic re-inspection of the phenometric panel reveals double-logistic curve-fit instability in a non-trivial subset of district-year cells (Section 5.4).")

if old_chunk not in old_text:
    print("ERROR: target chunk not found verbatim. Trying soft match.")
    # show near-context for debug
    idx = old_text.find("directionally")
    print(old_text[max(0, idx-80): idx+200])
else:
    new_text = old_text.replace(old_chunk, new_chunk)
    # Wipe runs and rebuild as a single run
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    p_xml = p._element
    # Preserve the first run's rPr (formatting) if present
    first_run = p_xml.find('w:r', ns)
    rpr = first_run.find('w:rPr', ns) if first_run is not None else None
    rpr_copy = deepcopy(rpr) if rpr is not None else None
    for r in p_xml.findall('w:r', ns):
        p_xml.remove(r)
    new_run = etree.SubElement(p_xml, W + 'r')
    if rpr_copy is not None:
        new_run.append(rpr_copy)
    new_t = etree.SubElement(new_run, W + 't')
    new_t.text = new_text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    doc.save(DOC_PATH)
    print("OK — abstract softened.")

# Verify
doc2 = Document(DOC_PATH)
p2 = doc2.paragraphs[14]
idx = p2.text.find("+15.108")
print("VERIFY:", p2.text[idx:idx+350])
