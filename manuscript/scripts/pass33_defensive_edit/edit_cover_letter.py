"""Insert a sensitivity-analysis disclosure paragraph in the cover letter."""
from docx import Document
from copy import deepcopy
from lxml import etree

DOC_PATH = '/home/user/workspace/RiceBaCI-GEE/manuscript/Cover_Letter.docx'
doc = Document(DOC_PATH)

new_text = (
    "Transparent reporting of post-hoc sensitivities. During final pre-submission diagnostic "
    "review we identified that the Beck et al. (2006) double-logistic curve-fit in Module 04 "
    "exhibits a quantisation pattern in a non-trivial subset of district-year cells, a known "
    "limitation of this functional form on tropical multi-cropped pixels (Atkinson et al. 2012; "
    "Cao et al. 2015). We document this finding in full in Section 5.4 of the manuscript and in "
    "the new Supplementary Note S10, where we also report five additional post-hoc sensitivity "
    "analyses probing COVID-cyclone collinearity, the Sentinel-1B 2021 mission failure, "
    "monsoon-onset heterogeneity, a continuous-exposure specification, and a salinity-carryover "
    "lag specification. The headline qualitative results (the indistinguishability of saline-surge "
    "from agronomic-flood SAR signatures, the classifier OA = 0.844 SAR-only / 0.990 full-feature, "
    "the EOS null) are unaffected by these sensitivities. The absolute magnitude of the "
    "τ̂_corrected_SOS coefficient is qualified as indicative rather than confirmatory, and a "
    "TIMESAT-based re-implementation of the phenometric extractor is registered for follow-up "
    "work. We believe this transparent disclosure strengthens rather than weakens the manuscript "
    "and provide it here for the Editors' consideration."
)

# Find paragraph 8 (Headline empirical findings) — insert AFTER it
template = doc.paragraphs[8]
template_xml = template._element
new_p_xml = deepcopy(template_xml)

ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
for r in new_p_xml.findall('w:r', ns):
    new_p_xml.remove(r)

new_run = etree.SubElement(new_p_xml, W + 'r')
new_t = etree.SubElement(new_run, W + 't')
new_t.text = new_text
new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

template_xml.addnext(new_p_xml)
doc.save(DOC_PATH)

# Verify
doc2 = Document(DOC_PATH)
for i, p in enumerate(doc2.paragraphs[7:11]):
    print(f"  para[{7+i}] style={p.style.name!r}")
    print(f"     {p.text[:160]!r}")
