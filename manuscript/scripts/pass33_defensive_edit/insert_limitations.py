"""
Insert the new 'Sixth' Limitations paragraph after paragraph 198 (5.4 body)
and before paragraph 199 (5.5 Future Work heading).
"""
from docx import Document
from copy import deepcopy
from lxml import etree

DOC_PATH = '/home/user/workspace/RiceBaCI-GEE/manuscript/Manuscript.docx'

with open('/home/user/workspace/RiceBaCI-GEE/manuscript/_edits/limitations_new_paragraph.md', 'r') as f:
    new_text = f.read().strip()

doc = Document(DOC_PATH)

# Template = the existing 5.4 limitations body paragraph (paragraph 198)
template_p = doc.paragraphs[198]
template_xml = template_p._element

# Create a copy of the template element, clear all runs, set new text
new_p_xml = deepcopy(template_xml)

# Remove all <w:r> run elements from the new paragraph (keep pPr style)
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for r in new_p_xml.findall('w:r', ns):
    new_p_xml.remove(r)

# Build a single new run with the text (preserve line breaks as <w:br/>)
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
new_run = etree.SubElement(new_p_xml, W + 'r')

# Join text into single paragraph (replace newlines with single space, since this
# is prose meant to flow). The user's text used line-wrap for editing only.
flat_text = ' '.join(line.strip() for line in new_text.split('\n') if line.strip())
new_t = etree.SubElement(new_run, W + 't')
new_t.text = flat_text
new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

# Insert immediately after the template paragraph (before 5.5 heading)
template_xml.addnext(new_p_xml)

doc.save(DOC_PATH)
print("OK — inserted new Limitations paragraph after 5.4 body.")

# Verify
doc2 = Document(DOC_PATH)
for i, p in enumerate(doc2.paragraphs[197:202]):
    snippet = p.text[:120].replace('\n', ' ')
    print(f"  para[{197+i}] style={p.style.name!r} text={snippet!r}")
