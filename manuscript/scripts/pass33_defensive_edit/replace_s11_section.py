"""
Remove the broken (pandoc-built) S11 section from the supplement and append the
clean v2 (python-docx native) S11 section.

Strategy:
1. Restore supplement from the pre-pass33 backup (which has no S11 at all).
2. Re-insert Table S10 (rf_feature_importance) — using python-docx merge.
3. Append clean S11 (v2) from the rebuilt source.
"""
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

BACKUP_PATH = '/home/user/workspace/RiceBaCI-GEE/manuscript/Supplement_Combined.docx.bak_pre_pass33_defensive'
TARGETS = [
    '/home/user/workspace/RiceBaCI-GEE/manuscript/supplement/Supplement_Combined.docx',
    '/home/user/workspace/RiceBaCI-GEE/manuscript/Supplement_Combined.docx',
]
TABLE_S10_PATH = '/home/user/workspace/RiceBaCI-GEE/manuscript/supplement/Table_S10_rf_feature_importance.docx'
S11_V2_PATH = '/home/user/workspace/RiceBaCI-GEE/manuscript/supplement/S11_sensitivity_analyses_v2.docx'

import shutil

def append_docx(target_path, addition_path, add_page_break=True):
    main = Document(target_path)
    add = Document(addition_path)
    if add_page_break:
        main.add_page_break()
    for element in add.element.body:
        if element.tag == qn('w:sectPr'):
            continue
        main.element.body.append(deepcopy(element))
    main.save(target_path)

for target in TARGETS:
    # Restore from backup
    shutil.copy(BACKUP_PATH, target)
    print(f"Restored {target.split('/')[-1]} from backup")
    # Append Table S10 (the orphan that the manuscript references)
    append_docx(target, TABLE_S10_PATH, add_page_break=True)
    # Append S11 (clean rebuild)
    append_docx(target, S11_V2_PATH, add_page_break=True)
    # Verify
    d = Document(target)
    t10_at = -1; s11_at = -1
    for i, p in enumerate(d.paragraphs):
        if 'Table S10' in p.text and 'feature importance' in p.text.lower() and t10_at == -1:
            t10_at = i
        if 'Supplementary Note S11' in p.text and s11_at == -1:
            s11_at = i
    n_tables = len(d.tables)
    print(f"  verify: T10@{t10_at}  S11@{s11_at}  tables={n_tables}")
