"""Trim abstract back toward 300 words while preserving the new defensive language."""
from docx import Document
from copy import deepcopy
from lxml import etree

DOC_PATH = '/home/user/workspace/RiceBaCI-GEE/manuscript/Manuscript.docx'
doc = Document(DOC_PATH)
p = doc.paragraphs[14]
old = p.text

# Trim 1: condense the "cross-validated by five robustness instruments" enumeration
t1_old = "cross-validated by five robustness instruments: wild-cluster restricted (WCR) bootstrap, leave-one-out jackknife, in-space and in-time placebo tests, post-hoc minimum-detectable-effect analysis, and out-of-sample transferability to Cyclone Bulbul."
t1_new = "cross-validated by a five-instrument robustness suite (wild-cluster bootstrap, leave-one-out jackknife, in-space and in-time placebo tests, minimum-detectable-effect analysis, and out-of-sample transferability to Cyclone Bulbul)."

# Trim 2: condense "The C-band SAR backscatter decrease during agronomic transplanting flooding..."
t2_old = "The C-band SAR backscatter decrease during agronomic transplanting flooding — the primary phenological anchor of published rice mapping algorithms — is nearly indistinguishable from the signal produced by storm-surge inundation four to six weeks earlier, silently corrupting start-of-season (SOS), peak-of-season (POS), and end-of-season (EOS) dates."
t2_new = "The C-band SAR backscatter decrease during agronomic transplanting flooding — the primary phenological anchor of published rice algorithms — is nearly indistinguishable from the signal produced by storm-surge inundation four to six weeks earlier, silently corrupting SOS, POS, and EOS dates."

# Trim 3: shorten "We developed a random-forest classifier fusing..."
t3_old = "We developed a random-forest classifier fusing Sentinel-1 backscatter, Sentinel-2 spectral indices, Joint Research Centre (JRC) Global Surface Water permanence, and ECMWF ERA5 maximum wind to discriminate the two flood types across five coastal Odisha districts and eight Kharif seasons (2017–2024)."
t3_new = "We developed a random-forest classifier fusing Sentinel-1 backscatter, Sentinel-2 spectral indices, JRC Global Surface Water permanence, and ERA5 maximum wind to discriminate the two flood types across five coastal Odisha districts and eight Kharif seasons (2017–2024)."

# Trim 4: shorten closing sentence
t4_old = "We provide the first empirical characterisation of the cyclone-flood confound in SAR rice phenology and an open, Google Earth Engine-deployable correction framework for cyclone-exposed Asian deltas."
t4_new = "We provide the first empirical characterisation of the cyclone-flood confound in SAR rice phenology and an open, Google Earth Engine-deployable correction framework for cyclone-exposed deltas."

new_text = old
for o, n in [(t1_old, t1_new), (t2_old, t2_new), (t3_old, t3_new), (t4_old, t4_new)]:
    if o not in new_text:
        print(f"WARNING: not found: {o[:60]!r}")
    else:
        new_text = new_text.replace(o, n)

# Rebuild paragraph
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
p_xml = p._element
first_run = p_xml.find('w:r', ns)
rpr = first_run.find('w:rPr', ns) if first_run is not None else None
rpr_copy = deepcopy(rpr) if rpr is not None else None
for r in p_xml.findall('w:r', ns):
    p_xml.remove(r)
new_run = etree.SubElement(p_xml, W + 'r')
if rpr_copy is not None:
    new_run.append(rpr_copy)
new_t = etree.SubElement(new_run, W + 't')
new_t.text = new_text
new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
doc.save(DOC_PATH)

# Verify
doc2 = Document(DOC_PATH)
abs_text = doc2.paragraphs[14].text
print("Abstract word count after trim:", len(abs_text.split()))
print()
print(abs_text)
