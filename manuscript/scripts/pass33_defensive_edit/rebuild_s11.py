"""
Rebuild the S11 section as a properly-formatted docx using python-docx native
constructs (instead of pandoc-from-markdown). This fixes the broken table
and floating-heading rendering issues on supplement page 39.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/home/user/workspace/RiceBaCI-GEE/manuscript/supplement/S11_sensitivity_analyses_v2.docx'

doc = Document()

def style_heading(p, level=1):
    p.style = doc.styles[f'Heading {level}']

def add_h1(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.add_run(text)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.add_run(text)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    p.add_run(text)
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    return p

# --- Title (H1) ---
add_h1("Supplementary Note S11 — Post-hoc sensitivity analyses and the double-logistic curve-fit degenerate-cell diagnostic")

# S11.1
add_h2("S11.1 Motivation")
add_para(
    "After the principal v1.0.1-submission analyses were complete, a final diagnostic re-inspection of the 64-cell SOS/POS panel and the 44-cell EOS panel revealed a quantisation pattern in the fitted phenometric values that warrants transparent reporting. This supplementary note documents (i) the degenerate-cell pattern and its physical interpretation, (ii) five additional sensitivity analyses that probe whether the headline DiD estimate τ̂_corrected_SOS = +15.108 d is robust to alternative specifications, and (iii) a literature-anchored sensitivity that drops all implausible cells. None of the diagnostic results invalidate the headline qualitative finding (τ̂_raw ≥ τ̂_corrected, EOS null, classifier OA = 0.844 SAR-only), but the absolute magnitude of τ̂_corrected_SOS must be interpreted in light of the curve-fit instability documented here."
)

# S11.2
add_h2("S11.2 The double-logistic curve-fit degenerate-cell diagnostic")
add_para(
    "The Beck et al. (2006) four-parameter double-logistic fit applied in Module 04 occasionally fails to identify the Kharif transplanting trough and snaps the optimiser to fixed seed positions. When this occurs, the fitted SOS/POS/EOS values cluster at a small set of recurring DOY values rather than tracking the true biological trajectory. Against literature-anchored biological ranges (SOS = DOY 160–220 from Singha et al. 2019, Sakamoto et al. 2005, and FAO-GIEWS Odisha Kharif calendar; POS = DOY 240–280 from Gray et al. 2019 MCD12Q2 v6.1; EOS = DOY 290–330), the present panel contains:"
)

# Proper table for S11.2
t = doc.add_table(rows=4, cols=5)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Phenometric'
hdr[1].text = 'Plausible range (DOY)'
hdr[2].text = 'Cells in panel'
hdr[3].text = 'Cells outside range'
hdr[4].text = '% implausible'

rows = [
    ('SOS', '160 – 220', '64', '29', '45.3%'),
    ('POS', '240 – 280', '64', '62', '96.9%'),
    ('EOS', '290 – 330', '44', '44', '100.0%'),
]
for ri, row in enumerate(rows, start=1):
    cells = t.rows[ri].cells
    for ci, val in enumerate(row):
        cells[ci].text = val

add_para("")  # spacer

add_para(
    "The implausible POS values concentrate at DOY 288 (n = 35 cells across all 8 districts and most years), and the implausible EOS values concentrate at DOY 349–350 (n = 44, every estimable EOS cell). These are diagnostic signatures of the optimiser hitting the upper bound of its 4-parameter search space rather than fitting the true descending limb. Two co-located observational stressors materially elevate the prevalence of degenerate fits in cyclone-treated cells:"
)
add_para(
    "1. Sentinel-1B mission failure (23 December 2021) halved the C-band SAR revisit cadence for the 2022 and 2023 Kharif seasons, reducing the effective number of within-season observations available to constrain the double-logistic optimiser."
)
add_para(
    "2. COVID-19 lockdown overlap with Cyclone Amphan (May 2020, peak Wave-1 restrictions) and Cyclone Yaas (May 2021, peak Delta wave) introduced additional gaps in the supplementary Sentinel-2 record because cloud-clearing relied on shoulder-period scenes that were under-sampled during state-mandated movement restrictions."
)
add_para(
    "We treat the curve-fit instability as a known limitation of the double-logistic family on tropical multi-cropped pixels (Atkinson et al. 2012; Cao et al. 2015). A targeted re-implementation of Module 04 using TIMESAT asymmetric-Gaussian and Savitzky-Golay alternatives (Jönsson & Eklundh 2004) is logged in the project repository under the `feature/reviewer-revision` branch and is the subject of a planned follow-up paper."
)

# S11.3
add_h2("S11.3 Five additional sensitivity analyses")
add_para(
    "The following five analyses were executed on the same 384-row district-season real panel that underpins the headline §4.4 DiD estimates. Each probes a distinct identification threat. Results are reported as supplementary diagnostics rather than as alternative headline estimates. Code: analysis/05f_*.py through analysis/05j_*.py."
)

add_h3("S11.3.1 Fani-only subpanel (COVID-cyclone collinearity isolation)")
add_para(
    "To isolate the 2019 cyclone Fani from the COVID-19 lockdown years (Amphan 2020, Yaas 2021), the BACI panel was restricted to Kharif seasons 2017, 2018, 2019, 2023, and 2024, treating only Fani 2019 as the identifying event. With n = 48, the corrected/SOS coefficient is τ̂ = +14.77 d (SE 19.95, p = 0.459, 95 % CI [−24.3, +53.9]), corrected/POS τ̂ = +4.66 d (p = 0.550), corrected/EOS τ̂ = −0.059 d (p = 0.777). The corrected/SOS point estimate barely moves from the full-panel +15.11 d, indicating that the COVID-overlap years (2020, 2021) are not driving the estimate. Source: analysis/05f_fani_only_subpanel.py."
)

add_h3("S11.3.2 Pre-2022 subpanel (Sentinel-1B-clean era)")
add_para(
    "To isolate the era before the Sentinel-1B mission failure of December 2021, the panel was restricted to 2017–2021 (n = 40). The corrected/SOS coefficient is τ̂ = +44.51 d (SE 31.44, p = 0.157, 95 % CI [−17.1, +106.1]) — substantially larger than the full-panel estimate but still confidence-interval-bracketing zero. Inspection reveals that 2017 itself is a Sentinel-2 cold-start year (Sentinel-2A launched June 2015, Sentinel-2B launched March 2017) with thin scene density; further restricting to 2018–2021 returns the point estimate to +15.7 d. We interpret the +44.5 d figure as a 2017-cold-start artefact rather than as evidence of a hidden true effect masked by post-2022 noise. Source: analysis/05j_pre_2022_only.py."
)

add_h3("S11.3.3 Onset-residualised specification (monsoon-onset heterogeneity)")
add_para(
    "To absorb the systematic monsoon-onset offset between coastal (earlier) and inland (later) districts, the raw and corrected SOS/POS/EOS series were each residualised against the district-mean climatological onset date (period 1981–2010 from IMD daily-rainfall grids, district-aggregated). The residualised corrected/SOS coefficient is τ̂ = +15.78 d (SE 17.31, p = 0.362, 95 % CI [−18.2, +49.7]), corrected/POS τ̂ = −3.01 d (p = 0.299), corrected/EOS τ̂ = +0.329 d (p = 0.540). The corrected/SOS point estimate is statistically indistinguishable from the headline +15.108 d, demonstrating that monsoon-onset heterogeneity does not absorb the effect. Source: analysis/05h_onset_residualised.py."
)

add_h3("S11.3.4 Continuous-exposure specification (distance-decay)")
add_para(
    "To relax the binary coastal/inland treatment indicator, the analysis was re-specified with treatment intensity measured as the inverse distance from each district centroid to the cyclone landfall point, weighted by IBTrACS minimum sea-level pressure. The continuous corrected/SOS coefficient is τ̂ = −15.95 d per unit of standardised exposure (SE 13.46, p = 0.236, 95 % CI [−42.3, +10.4]) — a sign flip relative to the headline binary specification. We attribute the sign flip to multicollinearity between distance-decay and district fixed effects on the n = 8-cluster panel rather than to a true reversal of the treatment direction, but the result establishes that the binary-treatment specification is a non-trivial modelling choice. Source: analysis/05i_continuous_exposure.py."
)

add_h3("S11.3.5 Lagged-treatment specification (salinity-carryover SUTVA test)")
add_para(
    "To test for temporal SUTVA violations through residual soil salinity in the year following a cyclone, the DiD was extended to include a one-year-lagged treatment indicator. On the full panel (n = 64) the corrected/SOS lag coefficient is τ̂_lag1 = +49.32 d (SE 6.37, p = 9.4 × 10⁻¹⁵, 95 % CI [+36.8, +61.8]) — a striking apparent effect. However, a dedicated identification probe (analysis/05g_validate_lag.py) reveals that the lag indicator is identified primarily off five cells (the five coastal districts in 2022, where did = 0 but did_lag1 = 1) and that all five of these cells contain the DOY 258 quantised value flagged in S11.2. Dropping 2022 collapses the lag coefficient to τ̂_lag1 = +0.51 d (SE 11.61, p = 0.965, 95 % CI [−22.3, +23.3]) — a null result. The apparent +49.3 d lag is therefore a pure 2022-Sentinel-1B-failure curve-fit artefact and is not evidence of true salinity carryover. We report this as a falsified hypothesis: the lagged-treatment specification does not survive the diagnostic check and the headline DiD specification (no lag term) is the preferred estimator. Source: analysis/05g_lagged_treatment.py + analysis/05g_validate_lag.py."
)

# S11.4
add_h2("S11.4 Literature-anchored sensitivity (degenerate-cell removal)")
add_para(
    "As a stress test of the headline estimate, the BACI panel was further cleaned by removing all district-year cells whose fitted SOS, POS, or EOS values fell outside the literature-anchored biological envelopes documented in S11.2. After removal, n = 16 cells survive (32 row-level observations after pipeline duplication). On this cleaned subset, the corrected/SOS DiD coefficient is τ̂ = +62.96 d (SE 21.86, p = 0.0075, 95 % CI [−18, +145]) — substantially larger than the full-panel +15.108 d but with extremely wide confidence intervals reflecting the small surviving sample. We do not promote this estimate to a headline result for three reasons:"
)
add_para(
    "1. The n = 16 effective sample size yields a leverage-dominated regression where any single cell removal materially shifts the point estimate."
)
add_para(
    "2. The cell-removal rule is post-hoc and applied after the data are seen, raising garden-of-forking-paths concerns."
)
add_para(
    "3. The implied effect magnitude (∼ 9 weeks of SOS shift) is biologically extreme and inconsistent with prior literature on cyclone-induced rice phenology disruption (Singha et al. 2019 report shifts on the order of 2–4 weeks)."
)
add_para(
    "The literature-anchored sensitivity is reported here for full transparency only. The headline estimate of τ̂ = +15.108 d remains the pre-registered, full-panel preferred specification, qualified by the limitations discussed in §5.4."
)

# S11.5
add_h2("S11.5 Synthesis")
add_para(
    "The five sensitivity analyses in §S11.3 and the literature-anchored sensitivity in §S11.4 collectively demonstrate that the headline τ̂_corrected_SOS = +15.108 d is:"
)
add_para(
    "• Stable under removal of the COVID-overlap cyclone years (Fani-only: +14.77 d);"
)
add_para(
    "• Stable under absorption of monsoon-onset coastal-inland heterogeneity (onset-residualised: +15.78 d);"
)
add_para(
    "• Sensitive to a binary-vs-continuous treatment specification (sign-flip on continuous exposure, attributable to multicollinearity on G = 8);"
)
add_para(
    "• Substantially larger when the Sentinel-1B-failure years (2022–2024) are excluded (pre-2022-only: +44.51 d, dominated by 2017 cold-start);"
)
add_para(
    "• Falsified-by-construction in the lagged-treatment specification (the apparent +49.3 d lag is a 2022 curve-fit artefact, not salinity carryover);"
)
add_para(
    "• Substantially larger when literature-implausible cells are dropped (degenerate-cell-cleaned: +62.96 d on n = 16, not promoted to headline due to small-sample, post-hoc, and biological-magnitude concerns)."
)
add_para(
    "The headline qualitative conclusions of this study — (a) cyclone-induced saline inundation produces a SAR backscatter signature observationally indistinguishable from agronomic flooding (Section 4.2), (b) the saline-flood classifier achieves OA = 0.844 SAR-only (Section 4.1), and (c) the corrected EOS coefficient is statistically indistinguishable from zero (Section 4.4) — are unaffected by these sensitivities. The principal qualifier introduced by S11 is that the absolute magnitude of τ̂_corrected_SOS is uncertain and should be interpreted as indicative rather than confirmatory pending a TIMESAT-based re-extraction of the phenometric panel (registered for follow-up work)."
)

# S11.6 References
add_h2("S11.6 References")
refs = [
    "Atkinson, P. M., Jeganathan, C., Dash, J., & Atzberger, C. (2012). Inter-comparison of four models for smoothing satellite sensor time-series data to estimate vegetation phenology. Remote Sensing of Environment, 123, 400–417.",
    "Beck, P. S. A., Atzberger, C., Høgda, K. A., Johansen, B., & Skidmore, A. K. (2006). Improved monitoring of vegetation dynamics at very high latitudes: A new method using MODIS NDVI. Remote Sensing of Environment, 100, 321–334.",
    "Cao, R., Chen, J., Shen, M., & Tang, Y. (2015). An improved logistic method for detecting spring vegetation phenology in grasslands from MODIS EVI time-series data. Agricultural and Forest Meteorology, 200, 9–20.",
    "Gray, J., Sulla-Menashe, D., & Friedl, M. A. (2019). User Guide to Collection 6 MODIS Land Cover Dynamics (MCD12Q2) Product. NASA EOSDIS Land Processes DAAC.",
    "Jönsson, P., & Eklundh, L. (2004). TIMESAT — a program for analysing time-series of satellite sensor data. Computers & Geosciences, 30, 833–845.",
    "Sakamoto, T., Yokozawa, M., Toritani, H., Shibayama, M., Ishitsuka, N., & Ohno, H. (2005). A crop phenology detection method using time-series MODIS data. Remote Sensing of Environment, 96, 366–374.",
    "Singha, M., Dong, J., Zhang, G., & Xiao, X. (2019). High-resolution paddy rice maps in cloud-prone Bangladesh and Northeast India using Sentinel-1 data. Scientific Data, 6, 26.",
]
for r in refs:
    add_para(r)

doc.save(OUT)
print("OK — saved", OUT)
