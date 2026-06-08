"""Embed figures into Manuscript.docx and Supplement_Combined.docx.

Strategy: for each (Figure N, image path, caption) entry, find the paragraph
that contains the first reference (e.g. "Figure 1" / "Figure S1") and insert
a new paragraph with the image + caption immediately after it.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from lxml import etree

ROOT = Path("/home/user/workspace/RiceBaCI-GEE")
FIG = ROOT / "figures"
ASSETS = ROOT / "assets"

# ---- MANUSCRIPT ----
MANU_FIGS = [
    # (label, file, width_cm, anchor regex (case-sensitive, word boundary), caption)
    ("Figure 1",
     FIG / "figure1_study_area.png",
     15.0,
     r"\(Figure 1\)",
     "Figure 1. Study area — five coastal districts of Odisha (Balasore, Bhadrak, Kendrapara, Jagatsinghpur, Puri) with three control districts inland (Cuttack, Dhenkanal, Angul). Cyclone tracks for Fani (2019), Amphan (2020), and Yaas (2021) shown with 50-km IBTrACS buffers. ESA WorldCover v200 cropland mask overlay."),
    ("Figure 1B",
     FIG / "fig1b_identification_dag.png",
     14.0,
     r"Figure 1B",
     "Figure 1B. Pearl-style identification DAG. Cyclone landfall (T) opens two pathways into the SAR backscatter trough: a legitimate transplanting-flood pathway and a confounding saline-surge pathway. Module 02 (the saline-flood classifier) intercepts the confounding pathway; the remaining backscatter trough identifies the agronomic SOS."),
    ("Figure 2",
     ASSETS / "Fig02_workflow.png",
     16.0,
     r"\(Figure 2\)",
     "Figure 2. Analytical workflow. Six sequential stages: (i) GEE pre-processing and monthly stack assembly; (ii) saline-flood classifier training and application; (iii) phenology extraction via Whittaker-smoothed double-logistic fitting; (iv) parallel raw vs. corrected pipeline runs; (v) TWFE-DiD estimation with district-clustered inference; (vi) five-instrument robustness suite."),
    ("Figure 3",
     FIG / "real_v21" / "fig3_event_study.png",
     16.0,
     r"event-study leads in Figure 3",
     "Figure 3. Event-study coefficients (Eq. 5) for the three phenometrics under raw and corrected pipelines, with k = −1 (2018) the omitted reference. Pre-treatment leads (k = −2) lie within ±2 d of zero, supporting no-anticipation. CR1-clustered 95% CIs shown."),
    ("Figure 4",
     FIG / "real_v21" / "fig4_district_sos_panel.png",
     16.0,
     r"spatial distribution of classified cyclone-flood pixels \(Figure 4\)",
     "Figure 4. Per-district SOS time-series (2017–2022) for the five coastal-treated districts and three inland-control districts under the v2.1-corrected pipeline. Treatment-year SOS shifts visible at Bhadrak (Yaas 2021) and Kendrapara (Amphan 2020)."),
    ("Figure 6",
     FIG / "fig6_placebo_distribution.png",
     15.0,
     r"in-space donor-swap permutation test \(Table S7, Figure 6\)",
     "Figure 6. In-space placebo distribution from donor-swap permutation (56 reassignments of treatment status to inland-control districts). Real τ-hat marked by red dashed line. Corrected/EOS cell hits the design-floor p_perm = 0.018 at G = 8."),
    # Figure 9 (pixel-level uncertainty rasters) removed: source figure carried
    # an 'ILLUSTRATIVE - REPLACE WITH REAL DATA' watermark and synthetic gradients.
    # Real per-district uncertainty summary is now reported in Supplementary Table
    # S12; pixel-level GeoTIFFs are released via the Mendeley deposit.
]

# ---- SUPPLEMENT ----
SUPP_FIGS = [
    ("Figure S1",
     FIG / "figS1_cyclone_climatology.png",
     16.0,
     # Anchor: the dedicated figure-section header at the back of the
     # supplement (uses em-dash). This avoids matching front-matter
     # bullets or in-text 'Figure S1' references.
     r"^Figure S1\s+[\u2014\-]\s+Pre-Kharif cyclone climatology$",
     "Figure S1. Pre-Kharif cyclone climatology. (A) Track-genesis schematic; (B) intensity-vs-day-of-year (DOY) scatter on Saffir-Simpson bands. Fani (2019), Amphan (2020), Yaas (2021) shown against the n = 19 IBTrACS 1981–2018 reference distribution."),
    ("Figure S2",
     FIG / "figS2_backscatter_signatures.png",
     14.0,
     # Anchor: dedicated figure-section header at back of supplement.
     r"^Figure S2\s+[\u2014\-]\s+Canonical Sentinel-1 backscatter signatures$",
     "Figure S2. Canonical Sentinel-1 backscatter signatures (VH, VV, CR vs. DOY) for the three mechanisms: transplanting flood, saline storm-surge, freshwater rainfall ponding. Onset markers and ΔVH calibrations from Hoshikawa et al. (2023), Wali et al. (2020), Filipponi (2019), Konkathi et al. (2024)."),
]


def insert_figure_after(doc, anchor_para, img_path, width_cm, caption_text, label):
    """Insert a figure paragraph and caption paragraph after anchor_para."""
    # Create new paragraph for image
    new_p = anchor_para._p.makeelement(
        anchor_para._p.tag, anchor_para._p.attrib
    )
    # Use python-docx machinery: create runs manually via inserted paragraphs
    # Easier: use add_paragraph then move XML
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))

    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption_text)
    cap_run.font.name = 'Arial'
    cap_run.font.size = Pt(9)
    cap_run.italic = True
    # Pin caption to following figure so they never split across a page break
    cap_para.paragraph_format.keep_with_next = True
    cap_para.paragraph_format.keep_together = True
    fig_para.paragraph_format.keep_with_next = False
    fig_para.paragraph_format.keep_together = True

    # Move both new paragraphs from end of doc to immediately after anchor.
    # Order on page: caption first (above), then figure (below) — prevents
    # page-break orphaning the caption when image is tall.
    anchor_para._p.addnext(fig_para._p)
    anchor_para._p.addnext(cap_para._p)
    return True


def find_anchor_paragraph(doc, regex):
    pat = re.compile(regex)
    for p in doc.paragraphs:
        if pat.search(p.text):
            return p
    return None


def embed(docx_path, figure_specs):
    doc = Document(str(docx_path))
    inserted = []
    missed = []
    for label, img, width, anchor_re, caption in figure_specs:
        if not Path(img).exists():
            missed.append(f"{label}: image missing at {img}")
            continue
        ap = find_anchor_paragraph(doc, anchor_re)
        if ap is None:
            missed.append(f"{label}: no anchor matching /{anchor_re}/")
            continue
        insert_figure_after(doc, ap, img, width, caption, label)
        inserted.append(label)
    doc.save(str(docx_path))
    return inserted, missed


if __name__ == "__main__":
    manu = ROOT / "manuscript" / "Manuscript.docx"
    supp = ROOT / "manuscript" / "Supplement_Combined.docx"

    print("=== Manuscript ===")
    ins, miss = embed(manu, MANU_FIGS)
    print("Inserted:", ins)
    print("Missed:", miss)

    # NOTE: Supplement figures (S1, S2) are embedded directly by
    # build_supplement_bundle.py via run.add_picture() in the dedicated
    # figure section at the back. Running embed() on the supplement here
    # would duplicate figures and insert captions at the wrong anchor
    # (e.g. the front-matter bullets that mention 'Figure S1').
    # Therefore the supplement is intentionally skipped.
    print("\n=== Supplement ===")
    print("Skipped: supplement figures already embedded by build_supplement_bundle.py")
