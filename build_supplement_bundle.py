"""Build combined Supplement.docx — TOC, page numbers, three Notes, nine Tables, two Figures.

Bundle order:
  Cover / Contents / Provenance declaration
  Note S1 — Bulbul transferability  (methods_module05b_bulbul.md)
  Note S2 — Cyclone climatology     (methods_module11_climatology.md)
  Note S3 — Backscatter physics     (methods_module12_backscatter.md)
  Tables S1–S9 (existing standalone docx)
  Figure 1B (identification DAG, panel B)
  Figure S1 (cyclone climatology)
  Figure S2 (backscatter signatures)

Approach: pandoc each Note .md → docx fragments; then python-docx merges
fragments + pre-rendered table docx + figure images into one master docx,
applying consistent A4 / Arial / heading styles, page numbers in footer,
TOC bookmarks, and section breaks.
"""
from __future__ import annotations
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path("/home/user/workspace/RiceBaCI-GEE")
SUPPL = ROOT / "manuscript/supplement"
FIGS = ROOT / "figures"
OUT = SUPPL / "Supplement_v0.3.0.docx"
TMP = Path(tempfile.mkdtemp(prefix="suppl_"))

# ------------------------------------------------------------------ helpers
def _add_page_number_footer(section):
    """Insert 'Supplement — page N of M' centred in the footer."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.text = ""
    run = p.add_run("Supplement — page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    # PAGE field
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r2 = p.add_run(); r2.font.name = "Arial"; r2.font.size = Pt(9)
    r2._r.append(fld_begin); r2._r.append(instr); r2._r.append(fld_end)
    r3 = p.add_run(" of "); r3.font.name = "Arial"; r3.font.size = Pt(9)
    fld_begin2 = OxmlElement("w:fldChar"); fld_begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText"); instr2.text = "NUMPAGES"
    fld_end2 = OxmlElement("w:fldChar"); fld_end2.set(qn("w:fldCharType"), "end")
    r4 = p.add_run(); r4.font.name = "Arial"; r4.font.size = Pt(9)
    r4._r.append(fld_begin2); r4._r.append(instr2); r4._r.append(fld_end2)


def _set_section_a4(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7); section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    _add_page_number_footer(section)


def _md_to_docx(md_path: Path, out_docx: Path) -> None:
    """Pandoc convert with tex_math_dollars preprocessing (matches build_s3_docx pipeline)."""
    src = md_path.read_text(encoding="utf-8")
    src = re.sub(r"\\\[(.+?)\\\]", lambda m: f"$${m.group(1)}$$", src, flags=re.DOTALL)
    src = re.sub(r"\\\((.+?)\\\)", lambda m: f"${m.group(1)}$", src, flags=re.DOTALL)
    tmp_md = TMP / (md_path.stem + ".md")
    tmp_md.write_text(src, encoding="utf-8")
    subprocess.run(
        ["pandoc", str(tmp_md),
         "-f", "markdown+tex_math_dollars+pipe_tables",
         "-t", "docx", "-o", str(out_docx)],
        check=True,
    )


_MATH_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def _cell_text(tc):
    """Read all text including OMath nodes from a w:tc cell.

    OMath inline math is emitted by pandoc for $...$ source; python-docx's
    cell.text skips it, so we walk the XML in document order and gather both
    w:t runs and m:t (OMath) text leaves. Inline math like '$\Delta_{obs}$ (d)'
    becomes 'Delta_obs (d)' with subscripts in textual form.
    """
    parts = []
    for el in tc.iter():
        tag = el.tag
        if tag == qn("w:t") and el.text:
            parts.append(el.text)
        elif tag == _MATH_NS + "t" and el.text:
            parts.append(el.text)
    return "".join(parts)


def _extract_table_data(src_tbl):
    """Extract [[cell_text,...],...] from a pandoc-emitted w:tbl element."""
    rows = []
    for tr in src_tbl.findall(qn("w:tr")):
        row = []
        for tc in tr.findall(qn("w:tc")):
            row.append(_cell_text(tc))
        rows.append(row)
    return rows


def _build_native_table(dst_doc, rows, header=True, font_pt=9):
    """Build a properly-styled python-docx table from extracted rows."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    nrows = len(rows)
    tbl = dst_doc.add_table(rows=nrows, cols=ncols)
    tbl.autofit = False
    tbl.style = "Table Grid"  # built-in style guaranteed to exist
    # Set table width = full text area (16 cm), columns equal.
    total_dxa = 9100
    col_w = int(total_dxa / ncols)
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
        tblW.set(qn("w:type"), "dxa"); tblW.set(qn("w:w"), str(total_dxa))
        layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        # Tight cell margins (top/bot 40 twips ≈ 0.07cm, left/right 80 twips).
        tblCellMar = OxmlElement("w:tblCellMar")
        for side, w in (("top", 40), ("left", 80), ("bottom", 40), ("right", 80)):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(w)); el.set(qn("w:type"), "dxa")
            tblCellMar.append(el)
        tblPr.append(tblCellMar)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for c in list(grid): grid.remove(c)
        for _ in range(ncols):
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(col_w))
            grid.append(gc)
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = tbl.cell(ri, ci)
            cell.width = Cm(16.0 / ncols)
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.name = "Arial"; r.font.size = Pt(font_pt)
            if header and ri == 0:
                r.font.bold = True


def _copy_body(src_doc: Document, dst_doc: Document) -> None:
    """Append all body elements (paragraphs, tables) from src to dst.

    Pandoc tables are rebuilt as native python-docx tables with the built-in
    'Table Grid' style and explicit equal-column widths — LibreOffice mis-
    renders pandoc tables (tblW='auto', no tblLayout) as vertically-stacked
    cells that orphan body paragraphs.
    """
    src_body = src_doc.element.body
    dst_body = dst_doc.element.body
    sect_pr = dst_body.find(qn("w:sectPr"))
    for child in list(src_body):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:tbl"):
            # Rebuild ALL tables natively. Pandoc tables (style='Table') and
            # pre-rendered LightGrid-Accent1 tables both mis-render in LO due to
            # tblW='auto' / no tblLayout + narrow auto-derived column widths
            # that wrap long numeric values mid-digit (e.g. '+15.2 89').
            rows = _extract_table_data(child)
            if rows:
                ncols = max(len(r) for r in rows)
                # Wide numeric tables (≥8 cols) use 8pt; otherwise 9pt.
                font_size = 8 if ncols >= 8 else 9
                _build_native_table(dst_doc, rows, header=True, font_pt=font_size)
                new_tbl = dst_body.findall(qn("w:tbl"))[-1]
                dst_body.remove(new_tbl)
                if sect_pr is not None:
                    sect_pr.addprevious(new_tbl)
                else:
                    dst_body.append(new_tbl)
                continue
        if sect_pr is not None:
            sect_pr.addprevious(child)
        else:
            dst_body.append(child)


def _add_section_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _heading(doc, text, level=1, color=(0x14, 0x14, 0x14), bookmark=None):
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if bookmark:
        # insert bookmark element
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(abs(hash(bookmark)) % 100000))
        bm_start.set(qn("w:name"), bookmark)
        p._p.append(bm_start)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    if bookmark:
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(abs(hash(bookmark)) % 100000))
        p._p.append(bm_end)
    return p


def _restyle_normal_runs(doc):
    """Force every run to Arial 10pt where no explicit size is set."""
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.font.name:
                run.font.name = "Arial"
            if run.font.size is None:
                # only set if it's body — heading runs have explicit sizes
                if not para.style.name.startswith("Heading"):
                    run.font.size = Pt(10)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.font.name:
                            run.font.name = "Arial"
                        if run.font.size is None:
                            run.font.size = Pt(9)


# ------------------------------------------------------------------ build
print("=== Building combined supplement bundle ===")
print(f"   workdir: {TMP}")

# Master document
doc = Document()

# Section 1 — cover (portrait, A4)
sect0 = doc.sections[0]
_set_section_a4(sect0, landscape=False)

# Set Normal style baseline
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.15

# --- COVER PAGE ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(80)
r = title.add_run("Supplementary Material")
r.font.name = "Arial"; r.font.size = Pt(26); r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(8)
r = sub.add_run(
    "Cyclonic damage to coastal Odisha rice systems:\n"
    "a Sentinel-driven Before–After / Control–Impact (BACI) "
    "design with pre-registered identification"
)
r.font.name = "Arial"; r.font.size = Pt(14); r.font.italic = True
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# Author block
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.paragraph_format.space_before = Pt(40)
r = auth.add_run("Supranab Panda")
r.font.name = "Arial"; r.font.size = Pt(12); r.font.bold = True
auth.add_run("\n").font.name = "Arial"
r = auth.add_run("ORCID 0009-0009-6496-6545")
r.font.name = "Arial"; r.font.size = Pt(10)
auth.add_run("\n").font.name = "Arial"
r = auth.add_run("Corresponding author: pandasupranab@gmail.com")
r.font.name = "Arial"; r.font.size = Pt(10)

# Repo / OSF block
ident = doc.add_paragraph()
ident.alignment = WD_ALIGN_PARAGRAPH.CENTER
ident.paragraph_format.space_before = Pt(40)
for line, sz in [
    ("Pre-registration: OSF c4mp8 (DOI 10.17605/OSF.IO/C4MP8)", 10),
    ("Code: github.com/pandasupranab/RiceBaCI-GEE (release v1.0.1-submission)", 10),
    ("Software archive: Zenodo 10.5281/zenodo.20587316 (concept 10.5281/zenodo.20024578)", 10),
    ("Data: Mendeley 10.17632/z3zxk4xy3c.1", 10),
    ("Compiled: 2026-06-08", 10),
]:
    r = ident.add_run(line + "\n")
    r.font.name = "Arial"; r.font.size = Pt(sz)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

_add_section_break(doc)

# --- DATA PROVENANCE DECLARATION ---
_heading(doc, "Data provenance declaration", level=1,
         color=(0xA1, 0x2C, 0x7B), bookmark="provenance")
prov = doc.add_paragraph()
r = prov.add_run(
    "This supplement accompanies release v1.0.1-submission of the "
    "github.com/pandasupranab/RiceBaCI-GEE harness (archived at Zenodo "
    "10.5281/zenodo.20587316). All numerical results in Tables S1–S9 and "
    "Figures S1–S2 are derived from real public Earth-observation inputs, "
    "as follows:"
)
r.font.name = "Arial"; r.font.size = Pt(10)

bullet_items = [
    ("Real Sentinel-2 phenology BACI panel (n = 64 district-year SOS/POS rows, "
     "n = 44 estimable EOS rows after right-censoring 20 post-cyclone cells, "
     "across 8 coastal/inland Odisha districts, 2017–2024) — derived from "
     "Sentinel-2 L2A NDVI/LSWI time-series via Whittaker smoothing and "
     "double-logistic curve fitting in Module 03; drives Tables S1, S2, S4, "
     "S5, S6, S7 and the Module 05/05a/05b/05d/05e/06/07/09 results. "
     "Mendeley deposit 10.17632/z3zxk4xy3c.1."),
    ("Real cyclone metadata (IMD/IBTrACS) — Fani 2019, Amphan 2020, Yaas 2021, "
     "Bulbul 2019 landfall parameters and the 1981–2018 reference distribution "
     "in Module 11. Embedded from public best-track records; "
     "Table S8 / Figure S1."),
    ("Real Sentinel-1 dual-polarisation backscatter signatures (Module 12) — "
     "per-district ΔVH/ΔVV/ΔCR profiles computed on Sentinel-1 IW GRD "
     "time-series, with literature anchors from Hoshikawa 2023, Wali 2020, "
     "Filipponi 2019, Konkathi 2024, Lee & Pottier 2009; Table S9 / Figure S2."),
    ("Module 02 random-forest classifier (Table S10) — v0.3.0-tagged real-data "
     "retraining on n = 480 automated reference labels (80 Copernicus EMS "
     "EMSR357 Fani delineation + 160 UN-SPIDER 2019 Sentinel-1 change-detection "
     "+ 240 Sentinel-1∩WorldCover∩JRC mask); n_train = 384, n_test = 96; "
     "OA = 0.990 full-feature, OA = 0.844 SAR-only, 5-fold CV OA = 0.831. "
     "The v2.1 per-cell correction summary (Table S13a) is computed directly "
     "from the real GEE-export panel."),
]
for it in bullet_items:
    bp = doc.add_paragraph(style="List Bullet")
    r = bp.add_run(it); r.font.name = "Arial"; r.font.size = Pt(10)

closing = doc.add_paragraph()
r = closing.add_run(
    "All structural claims (study design, identification strategy, "
    "falsifiability conditions, mechanism physics, climatological framing, "
    "code organisation) carry through unchanged from the pre-registered "
    "protocol (OSF c4mp8). All artefacts are version-pinned: GitHub "
    "v1.0.1-submission, Zenodo 10.5281/zenodo.20587316, Mendeley "
    "10.17632/z3zxk4xy3c.1. No modelling decision in this document is "
    "post-hoc to the data."
)
r.font.name = "Arial"; r.font.size = Pt(10); r.font.italic = True

_add_section_break(doc)

# --- TABLE OF CONTENTS ---
_heading(doc, "Contents", level=1, bookmark="toc")
toc_entries = [
    ("Note S1", "Bulbul (2019) transferability probe", "S1"),
    ("Note S2", "Pre-Kharif Bay-of-Bengal cyclone climatology", "S2"),
    ("Note S3", "Sentinel-1 dual-polarisation backscatter signatures", "S3"),
    ("Table S1", "Static TWFE-DiD point estimates (SOS, POS, EOS)", "T1"),
    ("Table S2", "Pre-trend tests on the BACI panel", "T2"),
    ("Table S3", "Bulbul transferability — district-level residuals", "T3"),
    ("Table S4", "Wild-cluster bootstrap 95% CIs (G = 8)", "T4"),
    ("Table S5", "Jackknife leave-one-out (district + year)", "T5"),
    ("Table S6", "MDE and power curves at G = 8", "T6"),
    ("Table S7", "Placebo / falsification permutation tests", "T7"),
    ("Table S8", "Pre-Kharif cyclone climatology — per-storm percentiles", "T8"),
    ("Table S9", "Canonical Sentinel-1 dual-pol backscatter features", "T9"),
    ("Figure S1", "Pre-Kharif cyclone climatology — track + intensity-vs-DOY", "F1"),
    ("Figure S2", "Canonical backscatter signatures — VH/VV/CR vs DOY", "F2"),
]
for label, title_text, anchor in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(4.0))
    r = p.add_run(label); r.font.name = "Arial"; r.font.size = Pt(10); r.font.bold = True
    r = p.add_run("\t" + title_text)
    r.font.name = "Arial"; r.font.size = Pt(10)

_add_section_break(doc)

# --- NOTES S1, S2, S3 ---
note_specs = [
    ("S1", "Note S1 — Bulbul (2019) transferability probe",
     SUPPL / "methods_module05b_bulbul.md"),
    ("S2", "Note S2 — Pre-Kharif Bay-of-Bengal cyclone climatology",
     SUPPL / "methods_module11_climatology.md"),
    ("S3", "Note S3 — Sentinel-1 dual-polarisation backscatter signatures",
     SUPPL / "methods_module12_backscatter.md"),
]
for sid, heading_text, md_path in note_specs:
    _heading(doc, heading_text, level=1,
             color=(0x1F, 0x3A, 0x5F), bookmark=f"note_{sid.lower()}")
    fragment = TMP / f"note_{sid}.docx"
    _md_to_docx(md_path, fragment)
    src = Document(str(fragment))
    _copy_body(src, doc)
    _add_section_break(doc)

# --- TABLES S1–S9 ---
table_specs = [
    ("S1", "Table S1 — Static TWFE-DiD point estimates",
     "Table_S1_did_static.docx"),
    ("S2", "Table S2 — Pre-trend tests",
     "Table_S2_pretrends.docx"),
    ("S3", "Table S3 — Bulbul transferability residuals",
     "Table_S3_bulbul_transferability.docx"),
    ("S4", "Table S4 — Wild-cluster bootstrap (G = 8)",
     "Table_S4_wild_bootstrap.docx"),
    ("S5", "Table S5 — Jackknife leave-one-out",
     "Table_S5_jackknife.docx"),
    ("S6", "Table S6 — MDE / power",
     "Table_S6_mde.docx"),
    ("S7", "Table S7 — Placebo / falsification",
     "Table_S7_placebo.docx"),
    ("S8", "Table S8 — Pre-Kharif cyclone climatology",
     "Table_S8_cyclone_climatology.docx"),
    ("S9", "Table S9 — Canonical S1 dual-pol features",
     "Table_S9_backscatter_features.docx"),
    ("S10", "Table S10 — RF feature importance and falsifiability checks",
     "Table_S10_rf_feature_importance.docx"),
    ("S13a", "Table S13a — Per-(district, year, metric) v2.1 correction summary",
     "Table_S13a_v21_correction_summary.docx"),
]
for tid, heading_text, fname in table_specs:
    _heading(doc, heading_text, level=1,
             color=(0x1F, 0x3A, 0x5F), bookmark=f"table_{tid.lower()}")
    src = Document(str(SUPPL / fname))
    _copy_body(src, doc)
    _add_section_break(doc)

# --- FIGURES S1, S2 ---
fig_specs = [
    ("S1", "Figure S1 — Pre-Kharif cyclone climatology",
     FIGS / "figS1_cyclone_climatology.png",
     "Track-genesis schematic (left) and intensity-vs-DOY scatter on "
     "Saffir-Simpson bands (right). Fani (2019), Amphan (2020), Yaas (2021) "
     "shown against the n = 19 IBTrACS 1981–2018 reference distribution "
     "of pre-Kharif Bay-of-Bengal systems."),
    ("S2", "Figure S2 — Canonical Sentinel-1 backscatter signatures",
     FIGS / "figS2_backscatter_signatures.png",
     "Three-panel stack (VH cross-pol, VV co-pol, CR cross-ratio vs day-of-year) "
     "for transplanting flood, saline storm-surge, and freshwater rainfall "
     "ponding mechanisms. Onset markers and ΔVH callouts annotate the "
     "discriminating feature space exploited by the Module 02 random-forest "
     "classifier."),
]
for fid, heading_text, png_path, caption in fig_specs:
    _heading(doc, heading_text, level=1,
             color=(0x1F, 0x3A, 0x5F), bookmark=f"figure_{fid.lower()}")
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic.add_run()
    run.add_picture(str(png_path), width=Inches(6.5))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = cap.add_run(f"Figure {fid}. ")
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.bold = True
    r = cap.add_run(caption)
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.italic = True
    if fid != fig_specs[-1][0]:
        _add_section_break(doc)

# --- Final pass: enforce Arial / sizing on stragglers ---
_restyle_normal_runs(doc)

doc.save(str(OUT))
print(f"\n[OK] wrote {OUT}")
print(f"     size = {OUT.stat().st_size} bytes")
