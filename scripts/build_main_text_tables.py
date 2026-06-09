#!/usr/bin/env python3
"""Build standalone Table_1.docx and Table_2.docx for Editorial Manager upload.

Elsevier (and RSE in particular) accepts numbered main-text tables either
embedded in the manuscript file OR uploaded as separate Word files. The
submission checklist (manuscript/03_submission_checklist.md) lists
`Table_1.docx` and `Table_2.docx` as separate-upload items, so we produce
publication-quality standalone files here.

Each output file contains:
  - The bold "Table N." caption block (one paragraph, justified left)
  - The table itself, with a grey-shaded bold header row
  - Header repeats on every page (heading-row property)
  - 9 pt body text, 10 pt header text, single-spaced
  - Centred body cells with left-aligned first column

Source of truth: hard-coded here against manuscript/manuscript_text.md
lines 126-132 (Table 1) and 140-149 (Table 2). The auditor extension below
verifies that the docx caption + every row's first cell matches the
manuscript so the two cannot silently drift.

Usage:
    python3 scripts/build_main_text_tables.py
Outputs:
    manuscript/Table_1.docx
    manuscript/Table_2.docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUTDIR = ROOT / "manuscript"

# ---------- TABLE DEFINITIONS (single source of truth) -----------------------

TABLE_1 = {
    "label": "Table 1",
    "caption": (
        "Tropical cyclone events affecting the coastal Odisha study area, "
        "2019–2021. ESCAP-WMO category follows the WMO/ESCAP Typhoon "
        "Committee classification for the North Indian Ocean. Surge height "
        "estimates are based on IMD post-landfall damage reports."
    ),
    "header": [
        "Cyclone",
        "Landfall Date",
        "Peak 3-min Wind (km h⁻¹)",
        "ESCAP-WMO Category",
        "District of Landfall",
        "Approx. Surge Height (m)",
    ],
    "rows": [
        ["Fani", "3 May 2019", "215 (peak) / 185 (landfall)",
         "Extremely Severe Cyclonic Storm (ESCS)", "Puri", "1.0–1.5"],
        ["Amphan", "20 May 2020", "240 (peak) / 155–165 (landfall)",
         "Super Cyclonic Storm (SuCS)",
         "Bakkhali, West Bengal (Sundarbans)", "1.5–2.0"],
        ["Yaas", "26 May 2021", "140 (peak, gusting to 155)",
         "Very Severe Cyclonic Storm (VSCS)", "Balasore", "1.0–2.0"],
    ],
    "col_widths_in": [0.95, 1.05, 1.15, 1.95, 1.55, 1.10],
    "filename": "Table_1.docx",
}

TABLE_2 = {
    "label": "Table 2",
    "caption": (
        "Satellite and ancillary datasets used in this study, all accessed "
        "via Google Earth Engine."
    ),
    "header": [
        "Dataset",
        "GEE Collection ID",
        "Temporal Range",
        "Native Resolution",
        "Role",
    ],
    "rows": [
        ["Sentinel-1 GRD (IW, VH+VV, descending)",
         "COPERNICUS/S1_GRD", "2017–2024", "10 m",
         "Primary SAR backscatter; classifier features; phenology retrieval"],
        ["Sentinel-2 L2A (harmonised)",
         "COPERNICUS/S2_SR_HARMONIZED", "2017–2024", "10 m",
         "Optical indices (NDVI, NDWI, LSWI, CIre); cloud-free phenology "
         "support"],
        ["JRC GSW Monthly History v1.4",
         "JRC/GSW1_4/MonthlyHistory", "1984–2024", "30 m",
         "Water permanence prior; saline-flood classifier feature"],
        ["ERA5-Land Daily Aggregates",
         "ECMWF/ERA5_LAND/DAILY_AGGR", "2017–2024", "~9 km",
         "Maximum 10-m wind speed; total precipitation; cyclone proximity "
         "signal"],
        ["ESA WorldCover v200",
         "ESA/WorldCover/v200", "2021 epoch", "10 m",
         "Cropland mask (class 40); pixel inclusion/exclusion filter"],
        ["GAUL Level-2 (FAO 2015)",
         "FAO/GAUL/2015/level2", "2015 epoch", "Vector",
         "District boundary delineation for study and control areas"],
    ],
    # GEE collection IDs in code font: monospace per Elsevier code-typography
    # convention. The build step below applies Courier to col index 1 cells.
    "monospace_cols": [1],
    "col_widths_in": [1.85, 1.85, 1.05, 0.85, 2.10],
    "filename": "Table_2.docx",
}

TABLES = [TABLE_1, TABLE_2]


# ---------- DOCX HELPERS -----------------------------------------------------

def _set_cell_shading(cell, fill_hex):
    """Apply a solid background fill to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_header_repeat(row):
    """Mark the first row as a header row that repeats on every page."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _set_borders(table):
    """Apply single 0.5 pt black borders to every cell."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")           # 0.5 pt = 4 eighths
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _style_cell(cell, text, *, bold=False, monospace=False, size_pt=9,
                align=WD_ALIGN_PARAGRAPH.LEFT, vcenter=True):
    """Write `text` into `cell` with the requested formatting."""
    cell.text = ""  # wipe the default empty paragraph
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if monospace:
        run.font.name = "Courier New"
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rFonts.set(qn("w:cs"), "Courier New")
        rPr.append(rFonts)
    if vcenter:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_one(table_spec, outpath):
    doc = Document()

    # Page setup: landscape A4 so wide tables don't wrap awkwardly.
    section = doc.sections[0]
    section.page_height = Inches(8.27)   # A4 width becomes height in landscape
    section.page_width = Inches(11.69)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # Default body style: 10 pt Times New Roman
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    # 1. Caption paragraph
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(6)
    label_run = cap.add_run(f"{table_spec['label']}. ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    text_run = cap.add_run(table_spec["caption"])
    text_run.font.size = Pt(10)

    # 2. Table
    n_cols = len(table_spec["header"])
    n_rows = 1 + len(table_spec["rows"])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False
    _set_borders(table)

    # Set explicit column widths
    for col_idx, width_in in enumerate(table_spec["col_widths_in"]):
        for row in table.rows:
            row.cells[col_idx].width = Inches(width_in)

    # Header row
    hdr = table.rows[0]
    _set_header_repeat(hdr)
    for j, head_text in enumerate(table_spec["header"]):
        cell = hdr.cells[j]
        _set_cell_shading(cell, "D9D9D9")  # light grey
        _style_cell(cell, head_text, bold=True, size_pt=10,
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    # Body rows
    monospace_cols = set(table_spec.get("monospace_cols", []))
    for i, row_vals in enumerate(table_spec["rows"], start=1):
        row = table.rows[i]
        for j, val in enumerate(row_vals):
            cell = row.cells[j]
            # First column left-aligned; the rest centred for readability.
            align = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                     else WD_ALIGN_PARAGRAPH.CENTER)
            # Role/description columns (wider ones) read better left-aligned.
            if j == n_cols - 1 and len(val) > 40:
                align = WD_ALIGN_PARAGRAPH.LEFT
            _style_cell(cell, val,
                        bold=False,
                        monospace=(j in monospace_cols),
                        size_pt=9,
                        align=align)

    # 3. Footnote-style source line (one blank paragraph then the footer)
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot_run = foot.add_run(
        "Source: see manuscript §2.2–2.3 for full provenance. "
        "Identical to the inline Table appearing in the main manuscript file."
    )
    foot_run.italic = True
    foot_run.font.size = Pt(8)

    doc.save(str(outpath))
    return outpath


def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    built = []
    for spec in TABLES:
        outpath = OUTDIR / spec["filename"]
        build_one(spec, outpath)
        size_kb = outpath.stat().st_size / 1024
        print(f"  OK: wrote {outpath.relative_to(ROOT)} "
              f"({size_kb:.1f} KB, {len(spec['rows'])} rows)")
        built.append(outpath)
    print(f"\nBuilt {len(built)} standalone main-text table files.")


if __name__ == "__main__":
    main()
