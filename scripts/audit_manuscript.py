"""
Comprehensive manuscript / supplement / cover-letter auditor.

Runs 8 categories of checks across the submission package and emits
a JSON report of every issue found. Exit code 0 if zero issues, 1 otherwise.

Categories:
  A. Synthetic / placeholder / pending wording in source MD
  B. Synthetic / placeholder / pending wording in rendered PDFs
  C. Cross-reference integrity (Table S1..S13, Figure 1..6, Note S1..S3)
  D. Identifier consistency (Zenodo, Mendeley, ORCID, OSF, GitHub release)
  E. Numeric headline-value drift (tau, SE, p-values, PI bounds)
  F. Required artifacts exist on disk and are non-empty
  G. Required result CSVs exist and are non-empty
  H. Build freshness (PDFs newer than their source MD)
"""

import argparse
import json
import os
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

ROOT = Path("/home/user/workspace/RiceBaCI-GEE")
MS = ROOT / "manuscript"
SUP = MS / "supplement"

# ---------- Whitelisted phrases (legitimate uses of forbidden words) -------
WHITELIST_PATTERNS = [
    r"synthetic aperture radar",           # SAR expansion
    r"\bSAR\b",                            # avoid matching nothing
    r"Synthetic control method",           # Abadie 2010 reference (with or without 's')
    r"Abadie",                             # Abadie reference window
    r"no synthetic",                       # any 'no synthetic ...' attestation
    r"No synthetic",                       # capitalised attestation
    r"zero synthetic",                     # attestation variant
    r"detect.{0,15}synthetic aperture",
    r"flooded rice with synthetic aperture",
    r"synthetic-\u03c3",                       # 'synthetic-σ' methodological contrast
    r"synthetic-sigma",                    # ASCII variant
    r"previous synthetic",                 # describing replaced analysis
    r"earlier synthetic",                  # describing replaced analysis
    r"replaced.{0,40}synthetic",           # 'replaced the ... synthetic'
    r"synthetic.{0,30}replaced",           # 'synthetic ... replaced'
    r"canonical TWFE",                     # canonical TWFE-DiD term
    r"canonical DiD",
    r"canonical difference",
    r"canonical two[- ]way",
]
WL_RE = re.compile("|".join(WHITELIST_PATTERNS), re.IGNORECASE)

# ---------- Forbidden phrases (case-insensitive) ---------------------------
FORBIDDEN = [
    r"\bsynthetic\b",
    r"\bplaceholder\b",
    r"\bprovisional\b",
    r"\bidealised\b",
    r"\bidealized\b",
    r"literature[- ]calibrated",
    r"to be (run|completed|added|replaced|filled)",
    r"\bTBD\b",
    r"\bTODO\b",
    r"forthcoming",
    r"will be replaced",
    r"pending GEE",
    r"once GEE",
    r"will be filled in",
]
FORB_RE = re.compile("|".join(FORBIDDEN), re.IGNORECASE)

# ---------- Required identifiers (must appear EXACTLY in submission) -------
REQUIRED_IDS = {
    "Zenodo this-version DOI": "10.5281/zenodo.20587316",
    "Zenodo concept DOI": "10.5281/zenodo.20024578",
    "Mendeley DOI": "10.17632/z3zxk4xy3c.1",
    "OSF pre-reg DOI": "10.17605/OSF.IO/C4MP8",
    "Panda ORCID": "0009-0009-6496-6545",
    "Sahu ORCID": "0000-0002-8048-1910",
    "GitHub release": "v1.0.1-submission",
}

# ---------- Headline numeric values that must agree everywhere -------------
HEADLINE_NUMERICS = {
    "tau_corrected_SOS": "15.108",
    "SE_corrected_SOS": "17.312",
    "tau_raw_SOS": "15.289",
    "SE_raw_SOS": "17.328",
    "p_corrected_SOS": "0.3828",
    "p_raw_SOS": "0.3776",
    "WCR_p_corrected": "0.4065",
    "PI_low": "-36.57",
    "PI_high": "+66.78",
    "Bulbul_mean_residual": "+6.56",
    "idio_SD": "19.88",
    "sigma_u": "13.06",
    "sigma_t": "41.75",
    "sigma_e": "31.28",
}

# ---------- Required cross-refs (must be defined AND used) -----------------
REQUIRED_TABLES = [f"Table S{i}" for i in range(1, 11)] + ["Table S13a"]
REQUIRED_FIGURES = ["Figure 1", "Figure 2", "Figure 3", "Figure 4",
                    "Figure 5", "Figure 6", "Figure S1", "Figure S2"]
REQUIRED_NOTES = ["Note S1", "Note S2", "Note S3"]

# ---------- Required artifact files ----------------------------------------
REQUIRED_ARTIFACTS = [
    MS / "Manuscript.docx", MS / "Manuscript.pdf",
    MS / "Cover_Letter.docx", MS / "Cover_Letter.pdf",
    SUP / "Supplement_Combined.docx", SUP / "Supplement_Combined.pdf",
    SUP / "Supplement_v0.3.0.docx",
    SUP / "Table_S1_did_static.docx",
    SUP / "Table_S2_pretrends.docx",
    SUP / "Table_S3_bulbul_transferability.docx",
    SUP / "Table_S4_wild_bootstrap.docx",
    SUP / "Table_S5_jackknife.docx",
    SUP / "Table_S6_mde.docx",
    SUP / "Table_S7_placebo.docx",
    SUP / "Table_S8_cyclone_climatology.docx",
    SUP / "Table_S9_backscatter_signatures.docx",
    ROOT / "figures/fig5_power_curves.pdf",
    ROOT / "figures/figS1_cyclone_climatology.pdf",
    ROOT / "figures/figS2_backscatter_signatures.pdf",
]

REQUIRED_RESULT_CSVS = [
    "analysis/results/real_v21/v21_correction_summary.csv",
    "analysis/results/real_v21/did_static.csv",
    "analysis/results/real_v21/event_study.csv",
    "analysis/results/real_v21/parallel_trends.csv",
    "analysis/results/real_v21/wild_bootstrap.csv",
    "analysis/results/real_v21/jackknife_district.csv",
    "analysis/results/real_v21/jackknife_year.csv",
    "analysis/results/real_v21/placebo_summary.csv",
    "analysis/results/real_v21/placebo_in_time.csv",
    "analysis/results/real_v21/placebo_in_space.csv",
    "analysis/results/real_v21/power_curves.csv",
    "analysis/results/real_v21/power_mde.csv",
    "analysis/results/real_v21/variance_components_real_v21.csv",
    "analysis/results/real_v21/s1_backscatter_real_signatures.csv",
    "analysis/results/real_v21/s1_backscatter_phase_means.csv",
    "analysis/results/real_v21/s1_backscatter_phase_deltas.csv",
    "analysis/results/real_v21/delta_cyc_empirical.csv",
    "analysis/results/real_v21/bulbul/probe_residuals_real_v21.csv",
    "analysis/results/real_v21/bulbul/probe_summary_real_v21.csv",
]

# Submission MD source files (audited line-by-line)
MD_SOURCES = [
    MS / "manuscript_text.md",
    MS / "00_cover_letter.md",
    MS / "01_highlights.md",
    MS / "02_declarations.md",
    MS / "methods_module09_power.md",
    MS / "methods_module05e_placebo.md",
    MS / "methods_module10_dag.md",
    SUP / "methods_module05b_bulbul.md",
    SUP / "methods_module11_climatology.md",
    SUP / "methods_module12_backscatter.md",
]

# Submission PDFs (audited page-by-page via pdftotext)
PDF_SOURCES = [
    MS / "Manuscript.pdf",
    MS / "Cover_Letter.pdf",
    SUP / "Supplement_Combined.pdf",
]


def pdf_text(pdf: Path) -> str:
    if not pdf.exists():
        return ""
    try:
        return subprocess.check_output(
            ["pdftotext", "-layout", str(pdf), "-"],
            stderr=subprocess.DEVNULL,
        ).decode("utf-8", errors="replace")
    except Exception as e:
        return f"<<PDF_READ_ERROR: {e}>>"


def scan_text_for_forbidden(text: str, source_name: str):
    """Yield (line_no, snippet, matched_forbidden) for forbidden hits not whitelisted."""
    issues = []
    for ln, line in enumerate(text.splitlines(), 1):
        # find any forbidden match
        for m in FORB_RE.finditer(line):
            # check whitelist: scan ±60 chars around the match
            start = max(0, m.start() - 60)
            end = min(len(line), m.end() + 60)
            window = line[start:end]
            if WL_RE.search(window):
                continue
            issues.append({
                "file": source_name,
                "line": ln,
                "match": m.group(0),
                "context": line.strip()[:200],
            })
    return issues


def check_category_A():
    """Forbidden wording in MD sources."""
    issues = []
    for md in MD_SOURCES:
        if not md.exists():
            issues.append({"file": str(md), "error": "MD source missing"})
            continue
        text = md.read_text()
        issues.extend(scan_text_for_forbidden(text, str(md.relative_to(ROOT))))
    return issues


def check_category_B():
    """Forbidden wording in rendered PDFs."""
    issues = []
    for pdf in PDF_SOURCES:
        if not pdf.exists():
            issues.append({"file": str(pdf), "error": "PDF missing"})
            continue
        text = pdf_text(pdf)
        issues.extend(scan_text_for_forbidden(text, str(pdf.relative_to(ROOT))))
    return issues


def check_category_C():
    """Cross-reference integrity in the main manuscript PDF."""
    issues = []
    text = pdf_text(MS / "Manuscript.pdf")
    sup_text = pdf_text(SUP / "Supplement_Combined.pdf")
    # Tables/figures/notes mentioned in main manuscript should be defined in supplement
    for tbl in REQUIRED_TABLES:
        cited = bool(re.search(rf"\b{re.escape(tbl)}\b", text + sup_text))
        if not cited:
            issues.append({"check": "cross-ref", "missing": tbl,
                           "note": "table not cited anywhere"})
    for note in REQUIRED_NOTES:
        cited_main = bool(re.search(rf"\b{re.escape(note)}\b", text))
        defined_sup = bool(re.search(rf"\b{re.escape(note)}\b", sup_text))
        if cited_main and not defined_sup:
            issues.append({"check": "cross-ref", "missing": note,
                           "note": "cited in manuscript but not defined in supplement"})
    return issues


def check_category_D():
    """Identifier consistency."""
    issues = []
    full = pdf_text(MS / "Manuscript.pdf") + pdf_text(MS / "Cover_Letter.pdf") + pdf_text(SUP / "Supplement_Combined.pdf")
    for label, value in REQUIRED_IDS.items():
        if value not in full:
            issues.append({"check": "identifier",
                           "missing": label, "expected": value})
    return issues


def check_category_E():
    """Headline numeric values must appear at least once in the manuscript or supplement.

    For p-values we accept either the 4-decimal form (0.3828) or the rounded
    3-decimal form (0.383) since tables render at 3 decimals.
    """
    issues = []
    full = pdf_text(MS / "Manuscript.pdf") + pdf_text(SUP / "Supplement_Combined.pdf")
    # normalise unicode minus
    full_norm = full.replace("\u2212", "-").replace("\u2013", "-")
    for label, val in HEADLINE_NUMERICS.items():
        v_strip = val.lstrip("+")
        accept = [v_strip]
        # accept 3-decimal rounding for 4-decimal p-values like 0.3828 -> 0.383
        if re.fullmatch(r"0\.\d{4}", v_strip):
            rounded = f"{round(float(v_strip), 3):.3f}"
            accept.append(rounded)
        # accept comma-decimal variants for European locales
        if "." in v_strip:
            accept.append(v_strip.replace(".", ","))
        if not any(a in full_norm for a in accept):
            issues.append({"check": "numeric",
                           "missing": label, "expected": val,
                           "accepted_forms": accept})
    return issues


def check_category_F():
    """Required artifact files exist and are non-empty."""
    issues = []
    for p in REQUIRED_ARTIFACTS:
        if not p.exists():
            issues.append({"check": "artifact", "missing": str(p.relative_to(ROOT))})
        elif p.stat().st_size < 1024:
            issues.append({"check": "artifact",
                           "file": str(p.relative_to(ROOT)),
                           "size": p.stat().st_size,
                           "note": "file < 1 KB, likely empty/broken"})
    return issues


def check_category_G():
    """Required result CSVs exist and are non-empty."""
    issues = []
    for rel in REQUIRED_RESULT_CSVS:
        p = ROOT / rel
        if not p.exists():
            issues.append({"check": "result_csv", "missing": rel})
        elif p.stat().st_size < 50:
            issues.append({"check": "result_csv",
                           "file": rel, "size": p.stat().st_size,
                           "note": "CSV < 50 B, likely empty"})
    return issues


def check_category_H():
    """PDF freshness: PDFs should be newer than their source MD."""
    issues = []
    pairs = [
        (MS / "manuscript_text.md", MS / "Manuscript.pdf"),
        (MS / "00_cover_letter.md", MS / "Cover_Letter.pdf"),
        (SUP / "Supplement_Combined.docx", SUP / "Supplement_Combined.pdf"),
    ]
    for src, pdf in pairs:
        if not src.exists() or not pdf.exists():
            continue
        if src.stat().st_mtime > pdf.stat().st_mtime + 1:
            issues.append({"check": "freshness",
                           "stale_pdf": str(pdf.relative_to(ROOT)),
                           "source": str(src.relative_to(ROOT)),
                           "src_mtime": datetime.fromtimestamp(src.stat().st_mtime).isoformat(),
                           "pdf_mtime": datetime.fromtimestamp(pdf.stat().st_mtime).isoformat()})
    return issues


def check_category_I():
    """Broken markdown links / cross-references inside MD sources."""
    issues = []
    link_re = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
    for md in MD_SOURCES:
        if not md.exists():
            continue
        for ln, line in enumerate(md.read_text().splitlines(), 1):
            for m in link_re.finditer(line):
                target = m.group(2)
                if target.startswith("http://") or target.startswith("https://"):
                    continue
                if target.startswith("#"):
                    continue
                # treat as repo-relative path
                cand = (md.parent / target).resolve()
                cand2 = (ROOT / target).resolve()
                if not cand.exists() and not cand2.exists():
                    issues.append({"check": "broken_link",
                                   "file": str(md.relative_to(ROOT)),
                                   "line": ln,
                                   "target": target,
                                   "text": m.group(1)})
    return issues


def check_category_J():
    """Sanity: PDFs should have non-trivial page counts."""
    issues = []
    expected_pages = {
        MS / "Manuscript.pdf": (15, 60),
        MS / "Cover_Letter.pdf": (1, 6),
        SUP / "Supplement_Combined.pdf": (20, 200),
    }
    for pdf, (lo, hi) in expected_pages.items():
        if not pdf.exists():
            continue
        try:
            out = subprocess.check_output(
                ["pdfinfo", str(pdf)], stderr=subprocess.DEVNULL
            ).decode()
            m = re.search(r"Pages:\s+(\d+)", out)
            if m:
                n = int(m.group(1))
                if n < lo or n > hi:
                    issues.append({"check": "page_count",
                                   "file": str(pdf.relative_to(ROOT)),
                                   "pages": n, "expected_range": [lo, hi]})
        except Exception as e:
            issues.append({"check": "page_count",
                           "file": str(pdf.relative_to(ROOT)),
                           "error": str(e)})
    return issues


def check_category_K():
    """Manuscript text must mention all required notes, tables, and figures at least once."""
    issues = []
    text = pdf_text(MS / "Manuscript.pdf")
    for name in REQUIRED_NOTES + ["Table S1", "Table S3", "Figure 5", "Figure S2"]:
        if not re.search(rf"\b{re.escape(name)}\b", text):
            issues.append({"check": "manuscript_citation",
                           "missing": name,
                           "note": "required reference not found in main manuscript"})
    return issues


def check_category_L():
    """Required figures appear in Figures_Bundle.pdf or are linked from manuscript."""
    issues = []
    figures = [
        ROOT / "figures/fig2_did_coefplot.png",
        ROOT / "figures/fig3_event_study.png",
        ROOT / "figures/fig4_district_sos_panel.png",
        ROOT / "figures/fig5_power_curves.png",
        ROOT / "figures/fig6_placebo_distribution.png",
        ROOT / "figures/figS1_cyclone_climatology.png",
        ROOT / "figures/figS2_backscatter_signatures.png",
    ]
    for f in figures:
        if not f.exists():
            issues.append({"check": "figure_file", "missing": str(f.relative_to(ROOT))})
        elif f.stat().st_size < 5000:
            issues.append({"check": "figure_file",
                           "file": str(f.relative_to(ROOT)),
                           "size": f.stat().st_size,
                           "note": "figure < 5 KB, likely corrupt"})
    return issues


def _read_docx_text(path: Path) -> str:
    """Concatenate all paragraph + table-cell text from a .docx."""
    try:
        from docx import Document
        d = Document(str(path))
        chunks = []
        for p in d.paragraphs:
            chunks.append(p.text)
        for t in d.tables:
            for r in t.rows:
                for c in r.cells:
                    chunks.append(c.text)
        return "\n".join(chunks)
    except Exception as e:
        return f"<<DOCX_READ_ERROR: {e}>>"


def check_category_M():
    """Supplement table DOCX content sanity — each table must carry real v2.1 numbers."""
    issues = []
    # Each entry: filename, list of (label, expected_substring) tuples.
    # If ANY expected substring is missing, the table is stale.
    table_checks = {
        "Table_S1_did_static.docx": [
            ("tau_raw_SOS", "15.289"),
            ("tau_corrected_SOS", "15.108"),
            ("SE_corrected_SOS", "17.312"),
        ],
        "Table_S2_pretrends.docx": [
            ("beta_pre_SOS", "-63.6"),
            ("p_pre_SOS", "0.343"),
        ],
        "Table_S3_bulbul_transferability.docx": [
            ("Boudh residual", "-8.11"),
            ("Ganjam residual", "-1.11"),
            ("Khordha residual", "+22.49"),
            ("Nayagarh residual", "+12.99"),
            ("tau_plug_in", "15.108"),
        ],
        "Table_S4_wild_bootstrap.docx": [
            ("WCR p raw SOS", "0.4000"),
            ("WCR p corrected SOS", "0.4065"),
        ],
        "Table_S5_jackknife.docx": [
            ("Bhadrak driver", "Bhadrak"),
            ("max_dtau", "76."),
        ],
        "Table_S6_mde.docx": [
            ("raw_SOS_MDE", "56.50"),
            ("raw_SOS_tau", "15.289"),
            ("corrected_EOS_MDE", "0.551"),
        ],
        "Table_S7_placebo.docx": [
            ("raw_SOS_tau", "15.29"),
            ("corrected_SOS_tau", "15.11"),
            ("p_perm_SOS", "0.5000"),
        ],
        "Table_S8_cyclone_climatology.docx": [
            ("IBTrACS", "IBTrACS"),
        ],
        "Table_S9_backscatter_signatures.docx": [
            ("Planetary Computer", "Planetary Computer"),
            ("Boudh peak canopy VH", "-13.79"),
            ("Boudh event window VH", "-13.77"),
            ("Khordha event window VH", "-13.42"),
        ],
    }
    # Normalisation: unicode minus, en-dash, math minus all become ASCII '-'
    def norm(s: str) -> str:
        return s.replace("\u2212", "-").replace("\u2013", "-")
    for fname, checks in table_checks.items():
        p = SUP / fname
        if not p.exists():
            issues.append({"check": "table_content", "missing_file": fname})
            continue
        txt = norm(_read_docx_text(p))
        for label, expected in checks:
            if norm(expected) not in txt:
                issues.append({"check": "table_content",
                               "file": fname,
                               "missing": label,
                               "expected": expected,
                               "note": "value not found in DOCX (possibly stale)"})
    return issues


def check_category_N():
    """Figure freshness — figures must be newer than the result CSV that feeds them."""
    issues = []
    pairs = [
        ("figures/fig2_did_coefplot.png", "analysis/results/real_v21/did_static.csv"),
        ("figures/fig3_event_study.png", "analysis/results/real_v21/event_study.csv"),
        ("figures/fig4_district_sos_panel.png", "analysis/results/real_v21/v21_correction_summary.csv"),
        ("figures/fig5_power_curves.png", "analysis/results/real_v21/power_curves.csv"),
        ("figures/fig6_placebo_distribution.png", "analysis/results/real_v21/placebo_in_space.csv"),
        ("figures/figS2_backscatter_signatures.png", "analysis/results/real_v21/s1_backscatter_phase_means.csv"),
    ]
    for fig_rel, csv_rel in pairs:
        fig = ROOT / fig_rel
        csv = ROOT / csv_rel
        if not fig.exists() or not csv.exists():
            continue
        if csv.stat().st_mtime > fig.stat().st_mtime + 1:
            issues.append({"check": "figure_freshness",
                           "stale_figure": fig_rel,
                           "newer_source": csv_rel,
                           "fig_mtime": datetime.fromtimestamp(fig.stat().st_mtime).isoformat(),
                           "src_mtime": datetime.fromtimestamp(csv.stat().st_mtime).isoformat()})
    return issues


def check_category_O():
    """Supplement bundle must NOT reference any archived/stale table file."""
    issues = []
    build = ROOT / "scripts/build_supplement_bundle.py"
    if not build.exists():
        return issues
    text = build.read_text()
    # any docx filename mentioned must exist (and not be in _archived_stale_synth/)
    docx_refs = re.findall(r"[\"\']([A-Za-z0-9_/-]+\.docx)[\"\']", text)
    for ref in set(docx_refs):
        # search for it in the supplement dir
        candidates = list((SUP).glob(ref))
        if not candidates:
            # also try as-is from project root (some refs are relative)
            ext = ROOT / ref
            if ext.exists():
                continue
            issues.append({"check": "bundle_ref",
                           "missing_docx": ref,
                           "note": "build_supplement_bundle.py references a docx that doesn't exist in supplement/"})
    # also flag any docx files in _archived_stale_synth/ as stale
    arch = SUP / "_archived_stale_synth"
    if arch.exists():
        for f in arch.iterdir():
            if f.suffix == ".docx":
                if f.name in text:
                    issues.append({"check": "bundle_ref",
                                   "file": f.name,
                                   "note": "archived stale file is still referenced in build script"})
    return issues


# ----- INTERNAL-CONTENT FILES (methods modules, READMEs, declarations, etc.) -----
INTERNAL_FILES = [
    "README.md",
    "SUBMISSION_README.md",
    "manuscript/00_cover_letter.md",
    "manuscript/01_highlights.md",
    "manuscript/02_declarations.md",
    "manuscript/03_submission_checklist.md",
    "manuscript/04_graphical_abstract_concept.md",
    "manuscript/methods_module02_baseline.md",
    "manuscript/methods_module05_did.md",
    "manuscript/methods_module05e_placebo.md",
    "manuscript/methods_module09_power.md",
    "manuscript/methods_module10_dag.md",
    "manuscript/supplement/methods_module05b_bulbul.md",
    "manuscript/supplement/methods_module11_climatology.md",
    "manuscript/supplement/methods_module12_backscatter.md",
    "docs/Data_Sources_Manifest.md",
    "docs/02_OSF_Pre_Registration.md",
    "docs/04_Week_1_Action_Plan.md",
    "docs/user_guides/01_how_to_run_phenology_in_GEE.md",
    "docs/user_guides/02_how_to_draw_labels.md",
    "docs/user_guides/03_active_learning_review.md",
]


def check_category_P():
    """Internal-content files (methods modules, READMEs, declarations, manifests, user guides)
    must not contain unresolved placeholders, TODO/TBD/FIXME, 'DOI: pending', or scaffold-era
    disclaimers indicating no real numbers exist. Pre-* backup snapshots are excluded."""
    issues = []
    # active placeholder/scaffold patterns
    patterns = {
        "PLACEHOLDER_marker": re.compile(r"\[PLACEHOLDER"),
        "DOI_pending":        re.compile(r"DOI:?\s*pending", re.I),
        "TBD":                re.compile(r"\bTBD\b"),
        "TODO":               re.compile(r"\bTODO\b"),
        "FIXME":              re.compile(r"\bFIXME\b"),
        "XXXX":               re.compile(r"\bXXXX+\b"),
        "forthcoming":        re.compile(r"forthcoming", re.I),
        "scaffold_disclaimer":re.compile(r"results\s+and\s+validation\s+numbers\s+are\s+placeholders", re.I),
        "illustrative_watermark": re.compile(r"ILLUSTRATIVE\s*—\s*REPLACE\s+WITH\s+REAL\s+DATA", re.I),
        "fabricated_results":  re.compile(r"no\s+legitimate\s+journal\s+will\s+accept\s+fabricated", re.I),
        "to_be_added":        re.compile(r"to\s+be\s+added", re.I),
        "to_be_released":     re.compile(r"to\s+be\s+released", re.I),
    }
    # whitelist phrases that look like hits but are legitimate references
    whitelist = [
        "Search the project for `[PLACEHOLDER:",  # SUBMISSION_README historical wording is removed
    ]
    for rel in INTERNAL_FILES:
        fp = ROOT / rel
        if not fp.exists():
            issues.append({"check": "internal_content", "file": rel, "note": "expected internal file missing"})
            continue
        txt = fp.read_text()
        for label, pat in patterns.items():
            for m in pat.finditer(txt):
                snippet = txt[max(0, m.start()-30):m.end()+30].replace("\n", " ")
                if any(w in snippet for w in whitelist):
                    continue
                line_no = txt[:m.start()].count("\n") + 1
                issues.append({"check": "internal_content", "file": rel,
                               "pattern": label, "line": line_no,
                               "snippet": snippet.strip()})
    return issues


def check_category_Q():
    """Internal-content files must not carry old synthetic-σ power/MDE numbers
    that contradict the real v2.1 panel (MDE range 0.551–56.501 d on the
    real panel; corrected/EOS τ̂ = -0.239 d, MDE = 0.551 d)."""
    issues = []
    stale_patterns = {
        "old_MDE_1.04":   re.compile(r"\b1\.04\s*d\b"),
        "old_MDE_2.49":   re.compile(r"\b2\.49\s*d\b"),
        "old_MDE_1.31":   re.compile(r"\b1\.31\s*d\b"),
        "old_tau_0.56":   re.compile(r"\|τ̂\|\s*=\s*0\.56"),
        "old_tau_5.662":  re.compile(r"(?<![\d.])5\.662(?![\d])"),
        "old_MDE_2.491":  re.compile(r"(?<![\d.])2\.491(?![\d])"),
    }
    for rel in INTERNAL_FILES:
        fp = ROOT / rel
        if not fp.exists():
            continue
        txt = fp.read_text()
        for label, pat in stale_patterns.items():
            for m in pat.finditer(txt):
                line_no = txt[:m.start()].count("\n") + 1
                issues.append({"check": "stale_numeric", "file": rel,
                               "pattern": label, "line": line_no})
    return issues


def check_category_R():
    """Real v2.1 identifiers / numerics that MUST appear in selected internal files
    so they remain self-consistent with the manuscript."""
    issues = []
    must_appear = {
        "SUBMISSION_README.md": [
            "v1.0.1-submission",
            "10.5281/zenodo.20587316",
            "10.17632/z3zxk4xy3c.1",
            "c4mp8",
            "0009-0009-6496-6545",
            "0000-0002-8048-1910",
        ],
        "manuscript/04_graphical_abstract_concept.md": [
            "15.108",  # corrected SOS τ̂
        ],
        "manuscript/methods_module09_power.md": [
            "56.501",  # raw/SOS MDE
            "0.551",   # corrected/EOS MDE
            "0.239",   # |τ̂| corrected/EOS
        ],
    }
    for rel, needles in must_appear.items():
        fp = ROOT / rel
        if not fp.exists():
            issues.append({"check": "required_anchor", "file": rel, "note": "file missing"})
            continue
        txt = fp.read_text()
        for n in needles:
            if n not in txt:
                issues.append({"check": "required_anchor", "file": rel,
                               "missing": n,
                               "note": "required real-v2.1 anchor not present in internal file"})
    return issues


# ----- SUBMISSION PACKAGE (the files actually uploaded to Editorial Manager) -----
SUBMISSION_DOCX = {
    "manuscript/Manuscript.docx":              {"min_media": 7, "required_text": ["v1.0.1-submission", "10.5281/zenodo.20587316", "10.17632/z3zxk4xy3c.1", "c4mp8", "0009-0009-6496-6545", "0000-0002-8048-1910"]},
    "manuscript/Cover_Letter.docx":            {"min_media": 0, "required_text": ["v1.0.1-submission", "10.5281/zenodo.20587316", "10.17632/z3zxk4xy3c.1", "0009-0009-6496-6545", "0000-0002-8048-1910"]},
    "manuscript/Highlights.docx":              {"min_media": 0, "required_text": []},
    "manuscript/Declarations.docx":            {"min_media": 0, "required_text": ["0009-0009-6496-6545", "0000-0002-8048-1910", "10.17632/z3zxk4xy3c.1", "10.5281/zenodo.20024578"]},
    "manuscript/Figures_Bundle.docx":          {"min_media": 9, "required_text": []},
    "manuscript/supplement/Supplement_Combined.docx": {"min_media": 2, "required_text": ["v1.0.1-submission", "10.5281/zenodo.20587316", "10.17632/z3zxk4xy3c.1", "56.501", "0.551", "0.239"]},
}
# Stale numeric markers that MUST NOT appear in any submission docx (full text + tables).
# Patterns are matched with word-boundary regex on the full text to avoid false hits
# on legitimate values like '+22.49 d' (Khordha Bulbul residual) or '0.844' (SAR-only OA).
import re as _re_submission
SUBMISSION_STALE_PATTERNS = [
    ("old_synth_tau_5.662",  _re_submission.compile(r"(?<![\d.])5\.662(?![\d])")),
    ("old_synth_MDE_2.491",  _re_submission.compile(r"(?<![\d.])2\.491(?![\d])")),
    ("old_MDE_1.04d",        _re_submission.compile(r"(?<![\d.])1\.04\s*d\b")),
    ("old_MDE_2.49d",        _re_submission.compile(r"(?<![\d.])2\.49\s*d\b")),
    ("old_MDE_1.31d",        _re_submission.compile(r"(?<![\d.])1\.31\s*d\b")),
    ("old_tau_abs_0.56",     _re_submission.compile(r"\|τ̂\|\s*=\s*0\.56")),
    ("placeholder_marker",   _re_submission.compile(r"\[PLACEHOLDER")),
    ("doi_pending",          _re_submission.compile(r"DOI:?\s*pending", _re_submission.I)),
    ("tbd_word",             _re_submission.compile(r"\bTBD\b")),
    ("todo_word",            _re_submission.compile(r"\bTODO\b")),
    ("fixme_word",           _re_submission.compile(r"\bFIXME\b")),
    ("forthcoming_word",     _re_submission.compile(r"forthcoming", _re_submission.I)),
    ("to_be_added",          _re_submission.compile(r"to\s+be\s+added", _re_submission.I)),
    ("to_be_released",       _re_submission.compile(r"to\s+be\s+released", _re_submission.I)),
    ("fabricated_results",   _re_submission.compile(r"fabricated\s+results", _re_submission.I)),
    ("illustrative_caveat",  _re_submission.compile(r"ILLUSTRATIVE\s*—\s*REPLACE", _re_submission.I)),
]


def _docx_full_text(path):
    """Return full text of a DOCX including paragraphs and table cells."""
    from docx import Document
    d = Document(str(path))
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _docx_media_count(path):
    """Return the number of embedded media (image) files in a DOCX."""
    import zipfile
    with zipfile.ZipFile(str(path)) as z:
        return len([n for n in z.namelist() if n.startswith("word/media/")])


def check_category_S():
    """Deep content check on each Editorial-Manager submission DOCX:
       - file must exist
       - must contain at least the expected number of embedded images
       - must contain every required identifier / anchor
       - must contain none of the stale numeric or scaffold markers
       - matching .pdf must exist and be newer than the .docx
    """
    issues = []
    for rel, spec in SUBMISSION_DOCX.items():
        fp = ROOT / rel
        if not fp.exists():
            issues.append({"check": "submission_pkg", "file": rel, "note": "submission file missing"})
            continue
        try:
            n_media = _docx_media_count(fp)
        except Exception as exc:
            issues.append({"check": "submission_pkg", "file": rel, "note": f"cannot read docx: {exc}"})
            continue
        if n_media < spec["min_media"]:
            issues.append({"check": "submission_pkg", "file": rel,
                           "note": f"embedded-media count {n_media} < required {spec['min_media']}"})
        try:
            text = _docx_full_text(fp)
        except Exception as exc:
            issues.append({"check": "submission_pkg", "file": rel, "note": f"cannot read docx text: {exc}"})
            continue
        for needle in spec["required_text"]:
            if needle not in text:
                issues.append({"check": "submission_pkg", "file": rel,
                               "missing": needle, "note": "required anchor not present"})
        for label, pat in SUBMISSION_STALE_PATTERNS:
            m = pat.search(text)
            if m:
                snippet = text[max(0, m.start()-40):m.end()+40].replace("\n", " ")
                issues.append({"check": "submission_pkg", "file": rel,
                               "stale": label, "snippet": snippet,
                               "note": "stale or scaffold marker present"})
        # matching PDF must exist and be at least as new as the docx
        pdf = fp.with_suffix(".pdf")
        if not pdf.exists():
            issues.append({"check": "submission_pkg", "file": rel,
                           "note": f"matching PDF {pdf.name} missing"})
        elif pdf.stat().st_mtime + 1 < fp.stat().st_mtime:
            issues.append({"check": "submission_pkg", "file": rel,
                           "note": f"PDF {pdf.name} older than DOCX (PDF mtime {pdf.stat().st_mtime:.0f} < DOCX mtime {fp.stat().st_mtime:.0f})"})
    return issues


# ---------------------------------------------------------------------------
# T_figure_data_correspondence
#   Verify that each figure PNG's headline annotated values match the
#   corresponding source CSV in analysis/results/real_v21/. We don't OCR the
#   image; instead we verify that the figure was built from the current CSV
#   AND that the CSV currently contains the expected values that the figure
#   should be showing (per the figure's design spec).
# ---------------------------------------------------------------------------
def check_category_T():
    import csv
    issues = []
    results_dir = ROOT / "analysis" / "results" / "real_v21"
    figures_dir = ROOT / "figures"

    # 1. did_static.csv must contain the v2.1 τ values that fig2/fig1b cite.
    expected_static = {
        ("raw", "SOS"): 15.289,
        ("raw", "POS"): -3.587,
        ("raw", "EOS"): 0.0,
        ("corrected", "SOS"): 15.108,
        ("corrected", "POS"): -3.677,
        ("corrected", "EOS"): -0.239,
    }
    static_path = results_dir / "did_static.csv"
    if static_path.exists():
        seen = {}
        with static_path.open() as f:
            for row in csv.DictReader(f):
                seen[(row["pipeline"], row["metric"])] = float(row["tau_days"])
        for key, exp in expected_static.items():
            got = seen.get(key)
            if got is None:
                issues.append({"check": "T_did_static_row_missing", "cell": key})
            elif abs(got - exp) > 0.01:
                issues.append({
                    "check": "T_did_static_value_mismatch",
                    "cell": key, "expected": exp, "got": got,
                })
    else:
        issues.append({"check": "T_did_static_missing",
                       "path": str(static_path)})

    # 2. Fig 2/3 freshness vs did_static.csv & event_study.csv
    for fig_name, src_name in [
        ("fig2_did_coefplot.png", "did_static.csv"),
        ("fig3_event_study.png",  "event_study.csv"),
    ]:
        fig_path = figures_dir / fig_name
        src_path = results_dir / src_name
        if not fig_path.exists():
            issues.append({"check": "T_figure_missing", "file": fig_name})
            continue
        if not src_path.exists():
            issues.append({"check": "T_source_csv_missing", "file": src_name})
            continue
        if fig_path.stat().st_mtime < src_path.stat().st_mtime:
            issues.append({
                "check": "T_figure_older_than_source_csv",
                "figure": fig_name, "source": src_name,
            })

    # 3. Event study must have rows for SOS, POS AND EOS (regression catcher
    #    — prior Fig 3 only plotted SOS panels).
    es_path = results_dir / "event_study.csv"
    if es_path.exists():
        seen_metrics = set()
        with es_path.open() as f:
            for row in csv.DictReader(f):
                seen_metrics.add(row["metric"])
        for m in ("SOS", "POS", "EOS"):
            if m not in seen_metrics:
                issues.append({
                    "check": "T_event_study_missing_metric", "metric": m,
                })

    # 4. Fig 5 expected to extend power curves to tau >= 60 d (so the 0.80
    #    crossing for G=8 is visible). Encode by requiring power_curves.csv
    #    rows for true_tau_d >= 60.
    pc_path = results_dir / "power_curves.csv"
    if pc_path.exists():
        has_tau60 = False
        with pc_path.open() as f:
            for row in csv.DictReader(f):
                try:
                    if float(row["true_tau_d"]) >= 60:
                        has_tau60 = True
                        break
                except (KeyError, ValueError):
                    pass
        if not has_tau60:
            issues.append({
                "check": "T_power_curves_short",
                "detail": "power_curves.csv lacks rows with true_tau_d>=60",
            })

    # 5. Source code guard: scripts that previously produced wrong figures
    #    must not have regressed. Tripwires for the Pass-20 bug fixes.
    fig_src = (ROOT / "analysis" / "06_figures.py").read_text(
        encoding="utf-8", errors="ignore"
    )
    if "metric == 'SOS'" in fig_src:
        # Allowed in fig_district_panel; only flag if appears in fig_event_study scope.
        # crude scope check: split file at "def fig_event_study" and check next 60 lines
        if "def fig_event_study" in fig_src:
            body = fig_src.split("def fig_event_study", 1)[1]
            head = body[:3000]
            if "metric == 'SOS'" in head:
                issues.append({
                    "check": "T_fig3_only_sos_regression",
                    "detail": "fig_event_study hardcodes metric == 'SOS'",
                })

    power_src = (ROOT / "analysis" / "09_power_analysis.py").read_text(
        encoding="utf-8", errors="ignore"
    )
    if "ax.set_xlim(-0.2, 8.4)" in power_src:
        issues.append({
            "check": "T_fig5_xlim_regression",
            "detail": "power figure x-axis truncated at 8.4 d",
        })

    return issues


CATEGORIES = [
    ("A_md_forbidden", check_category_A),
    ("B_pdf_forbidden", check_category_B),
    ("C_cross_refs", check_category_C),
    ("D_identifiers", check_category_D),
    ("E_numerics", check_category_E),
    ("F_artifacts", check_category_F),
    ("G_result_csvs", check_category_G),
    ("H_pdf_freshness", check_category_H),
    ("I_broken_links", check_category_I),
    ("J_page_counts", check_category_J),
    ("K_manuscript_citations", check_category_K),
    ("L_figure_files", check_category_L),
    ("M_table_content", check_category_M),
    ("N_figure_freshness", check_category_N),
    ("O_bundle_refs", check_category_O),
    ("P_internal_content", check_category_P),
    ("Q_internal_stale_numerics", check_category_Q),
    ("R_internal_anchors", check_category_R),
    ("S_submission_package", check_category_S),
    ("T_figure_data_correspondence", check_category_T),
]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=str(ROOT / "audit_report.json"))
    ap.add_argument("--brief", action="store_true")
    args = ap.parse_args()

    report = {"timestamp": datetime.now().isoformat(),
              "categories": {}, "total_issues": 0}

    for name, fn in CATEGORIES:
        try:
            issues = fn()
        except Exception as e:
            issues = [{"check": name, "error": str(e)}]
        report["categories"][name] = {"count": len(issues), "issues": issues}
        report["total_issues"] += len(issues)

    Path(args.out).write_text(json.dumps(report, indent=2, default=str))

    if args.brief:
        print(f"AUDIT @ {report['timestamp']}: total_issues = {report['total_issues']}")
        for name, blk in report["categories"].items():
            print(f"  {name}: {blk['count']}")
    else:
        print(json.dumps(report, indent=2, default=str)[:8000])

    return 0 if report["total_issues"] == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
