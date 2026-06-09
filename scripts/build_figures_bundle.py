#!/usr/bin/env python3
"""Build manuscript/Figures_Bundle.docx and the matching PDF.

Each figure file in `figures/` is embedded as a centred 6.0-inch image
preceded by an Editorial-Manager-style caption header. This produces the
single 'Figures' upload that Editorial Manager expects, with every figure
in the correct order (Figure 1, Figure 1B, Figure 2-6, then Figure S1, S2).

Run after any change to the figures/ directory:
    python3 scripts/build_figures_bundle.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
FIGDIR = ROOT / "figures"
OUT = ROOT / "manuscript" / "Figures_Bundle.docx"

FIGURES = [
    ("Figure 1",  "figure1_study_area.png",
     "Study area: coastal and inland Odisha BACI districts. Treated coastal districts (Balasore, Bhadrak, Kendrapara, Jagatsinghpur, Puri) and inland control districts (Dhenkanal, Angul, Cuttack), with IBTrACS cyclone tracks for Fani (May 2019), Bulbul (Nov 2019), Amphan (May 2020), and Yaas (May 2021)."),
    ("Figure 1B", "fig1b_identification_dag.png",
     "Causal identification DAG for the saline-flood / agronomic-flood decoupling. The cyclone landfall is a single exogenous shock that opens parallel legitimate (transplanting flooding) and confounding (saline storm-surge) pathways into the same SAR backscatter trough; the Module 02 saline-flood classifier intercepts the confounding pathway."),
    ("Figure 2",  "fig2_did_coefplot.png",
     "TWFE-DiD coefficients (raw vs. classifier-corrected panel). Point estimates with district-clustered (CR1) standard-error bars and 95% wild-cluster restricted bootstrap intervals for SOS, POS, and EOS, raw and corrected pipelines on the real v2.1 panel."),
    ("Figure 3",  "fig3_event_study.png",
     "Event-study plot, raw vs. corrected pipelines (Kharif 2017-2024). Coefficients at relative-year leads (k = -2, -1) test for pre-trends; coefficients at lags (k = 0, 1, 2) trace dynamic treatment effects; k = -1 (2018) is the omitted reference."),
    ("Figure 4",  "fig4_district_sos_panel.png",
     "District-level SOS trajectories, raw vs. corrected (2017-2024). Solid lines show raw-pipeline SOS dates; dashed lines show classifier-corrected SOS dates; vertical bands mark cyclone landfall years. Treated coastal districts are coloured; inland controls are grey."),
    ("Figure 5",  "fig5_power_curves.png",
     "Empirical power curves from the Monte-Carlo simulation (Note S4). Rejection rate of H0 against true effect size tau over 999 replications per grid point at G = {4, 6, 8, 12} with empirical variance components sigma_u = 13.06 d, sigma_t = 41.75 d, sigma_epsilon = 31.28 d estimated from the real v2.1 corrected-SOS panel. At G = 8 (this study) power >= 0.80 is reached only for tau >= 60 d; the type-I rate under H0 is 0.07, close to nominal 0.05."),
    ("Figure 6",  "fig6_placebo_distribution.png",
     "In-space placebo distributions (donor-swap, 56 permutations). Histograms of placebo tau_hat values from all swaps of treated/control labels across district pairs, with the observed tau_hat marked as a vertical line. The corrected/EOS cell yields p_perm = 0.018 (non-significant after Bonferroni correction across the six-cell family)."),
    ("Figure S1", "figS1_cyclone_climatology.png",
     "Cyclone climatology for the Bay of Bengal, 1980-2024. Annual landfall counts in the 50-km IBTrACS buffer around the coastal Odisha districts, stratified by Saffir-Simpson category and pre-monsoon vs. post-monsoon timing."),
    ("Figure S2", "figS2_backscatter_signatures.png",
     "Sentinel-1 RTC dual-polarisation VH/VV/CR backscatter signatures across four phenological phases (pre-canopy, early canopy, peak canopy, event window) for the four Bulbul out-of-sample probe districts (Boudh, Ganjam, Khordha, Nayagarh), retrieved from Microsoft Planetary Computer."),
]


def build():
    doc = Document()
    # Title page
    h = doc.add_heading("Figures Bundle", level=0)
    h.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph()
    para.alignment = 1
    para.add_run(
        "Decoupling Cyclone-Induced Saline Inundation from Agronomic Flooding "
        "in Sentinel-1/2 Rice Phenology Retrieval: A Multi-Year Bay-of-Bengal "
        "Coastal Framework (2017-2024)"
    ).italic = True
    para = doc.add_paragraph()
    para.alignment = 1
    para.add_run("Supranab Panda, Sarat Chandra Sahu").bold = True
    doc.add_paragraph(
        "All figures are embedded at publication size (6.0-inch column width) "
        "from PNGs in figures/ regenerated from analysis/results/real_v21/. "
        "Each figure carries a short caption immediately below the image."
    ).alignment = 1
    doc.add_page_break()

    missing = []
    for label, fname, caption in FIGURES:
        fp = FIGDIR / fname
        if not fp.exists():
            missing.append((label, fname))
            continue
        # Heading
        h = doc.add_heading(f"{label}. {fname}", level=2)
        # Image
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run()
        run.add_picture(str(fp), width=Inches(6.0))
        # Caption
        cap = doc.add_paragraph()
        cap.add_run(f"{label}. ").bold = True
        cap.add_run(caption)
        doc.add_page_break()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"[OK] wrote {OUT.relative_to(ROOT)}  ({len(FIGURES) - len(missing)} figures embedded)")
    if missing:
        print(f"WARNING missing files: {missing}")
        raise SystemExit(2)


if __name__ == "__main__":
    build()
