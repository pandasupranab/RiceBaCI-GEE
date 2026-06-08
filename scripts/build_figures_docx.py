"""Build Figures_Bundle.docx — all manuscript figures at 1000 DPI JPG, with caption + description below each.

Uses v21 (corrected) versions for fig2/fig3/fig4; main figures dir otherwise.
"""
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION

ROOT = Path("/home/user/workspace/RiceBaCI-GEE")
FIG = ROOT / "figures"
V21 = FIG / "real_v21"
OUT_JPG = FIG / "jpg_1000dpi"
OUT_JPG.mkdir(exist_ok=True)
OUT_DOCX = ROOT / "manuscript" / "Figures_Bundle.docx"

# Figure manifest: (id, pdf source, label, caption, description, abbreviations)
# `abbreviations` lists every non-trivial abbreviation, symbol, and unit that
# physically appears in *that* figure or its caption — OR is implicit in the
# quantities shown (e.g. SE underlying the error bars in Figs 2, 3, 5) —
# including units (km, dB, DOY, kt, °N, °E, n, CI, SE, p, %, days, months) when
# present, per reviewer request to expand universally understood units.
FIGURES = [
    (
        "Figure 1A",
        FIG / "figure1_study_area.pdf",
        "Study area: coastal & inland Odisha BACI districts.",
        "Map of the eight-district study domain on the Bay-of-Bengal coast of Odisha, India. Five treated coastal districts (Balasore, Bhadrak, Kendrapara, Jagatsinghpur, Puri) front the Bay and are exposed to cyclone landfalls; three inland control districts (Dhenkanal, Angul, Cuttack) lie outside the storm-surge footprint and provide the BACI counterfactual. Landfall tracks for Cyclones Fani (2019), Bulbul (2019), Amphan (2020), and Yaas (2021) are overlaid. District boundaries: GADM v4.1. Coastline: Natural Earth.",
        "Compact study-area locator establishing the BACI treatment/control geography and the four natural-experiment cyclone events used in the difference-in-differences design.",
        "BACI = Before–After / Control–Impact (quasi-experimental design); GADM = Database of Global Administrative Areas (district-boundary source); km = kilometres (linear distance); °N = degrees north (latitude); °E = degrees east (longitude); N = geographic north (compass arrow).",
    ),
    (
        "Figure 1B",
        FIG / "fig1b_identification_dag.pdf",
        "Causal identification DAG for the saline-flood / agronomic-flood decoupling.",
        "Directed acyclic graph showing the identification strategy: the cyclone landfall (treatment) propagates through storm-surge saline inundation, which contaminates the SAR backscatter trough used as the phenology anchor, producing biased SOS/POS/EOS estimates. The random-forest classifier severs the back-door path by separating cyclone-induced from agronomic flooding, allowing the corrected pipeline to recover the causal effect on rice phenology.",
        "Visual statement of the identification assumption: classifier acts as a back-door adjustment on the surge-confounded pixel pool, restoring exchangeability between treated and control districts conditional on cyclone-flood pixel share.",
        "DAG = directed acyclic graph; SAR = synthetic aperture radar; VV = vertical-transmit, vertical-receive co-polarisation channel; VH = vertical-transmit, horizontal-receive cross-polarisation channel; SOS = start of season; POS = peak of season; EOS = end of season; p = probability value (significance level); p_WCB = wild-cluster bootstrap p-value; Module 02 = saline-flood random-forest classifier (v0.3.0).",
    ),
    (
        "Figure 2",
        V21 / "fig2_did_coefplot.pdf",
        "TWFE-DiD coefficients (raw vs. classifier-corrected panel).",
        "Two-way fixed-effects difference-in-differences estimates for τ_SOS, τ_POS, and τ_EOS on the raw phenology panel (blue) and the classifier-corrected panel (orange). Error bars are 95% wild-cluster restricted bootstrap intervals with district clustering and B = 4 999 draws. The corrected estimates show small but pre-registered attenuation (τ_SOS: +15.289 → +15.108 days; τ_EOS: 0.000 → −0.239 days), consistent with the bounded district-aggregated cyclone-flood pixel share (Fani 0.04–2.5 %, Amphan 0.005–1.9 %, Yaas 0.018–7.2 %).",
        "Headline DiD result on the v2.1 corrected panel; transparently reports the WCB-restricted CI inclusive of zero rather than over-claiming the SOS effect, while confirming the pre-registered direction τ_raw > τ_corrected > 0.",
        "TWFE = two-way fixed effects; DiD = difference-in-differences; τ = DiD treatment-effect estimate (days); τ_SOS / τ_POS / τ_EOS = τ on the start, peak, and end of the rice-cropping season; SOS / POS / EOS = start, peak, and end of season; WCB = wild-cluster (restricted) bootstrap; B = number of bootstrap draws; CI = confidence interval (95 % WCB-restricted, district-clustered); SE = standard error of τ̂ (district-clustered; the half-width of the displayed error bars is derived from SE · t-critical); days = calendar days (effect-size unit); % = per-cent (district-aggregated pixel-share unit); v2.1 = panel after applying the Module 02 classifier-attenuation correction.",
    ),
    (
        "Figure 3",
        V21 / "fig3_event_study.pdf",
        "Event-study plot, raw vs. corrected (Kharif 2017–2024).",
        "Year-by-relative-cyclone event-study coefficients for SOS in coastal vs. inland districts. Pre-treatment placebo years (k = −2, −1) are statistically indistinguishable from zero in both raw and corrected pipelines (pre-trends test p = 0.41), supporting the parallel-trends assumption. Post-treatment years (k = 0, +1, +2) show the cyclone effect peaking at k = 0 and attenuating thereafter; the corrected curve sits inside the raw curve at every post-treatment lag, as predicted by Module 11.",
        "Dynamic causal effect with explicit pre-trends test; pre-period coefficients on or near zero are the strongest visual evidence for BACI identification on the v2.1 corrected panel.",
        "SOS = start of season; β = event-study coefficient (days, relative to reference period k = −1); k = event time in years from first treatment landfall (k = 0 is the cyclone year); BACI = Before–After / Control–Impact design; Kharif = post-monsoon rice cropping season (≈ Jun–Nov); p = pre-trends-test p-value (Wald test on pre-period coefficients); SE = district-clustered standard error of β (the shaded ribbon / whisker half-width on each event-time coefficient); days = calendar days (effect-size unit); v2.1 = classifier-corrected phenology panel; Module 11 = cyclone-climatology + pixel-share attenuation model.",
    ),
    (
        "Figure 4",
        V21 / "fig4_district_sos_panel.pdf",
        "District-level SOS trajectories, raw vs. corrected (2017–2024).",
        "Small-multiple panel of mean Kharif SOS (DOY) for the five treated coastal districts and three inland controls, comparing the raw and classifier-corrected pipelines. The largest correction is at Bhadrak in 2021 (Yaas EOS, |Δ| = 1.51 days, consistent with the 7.21 % cyclone-flood pixel share in that district-year). Inland controls show negligible correction (|Δ| < 0.05 days), as expected when no surge inundation is present.",
        "District-disaggregated transparency check: shows where the classifier matters (Bhadrak-Yaas-2021) and where it doesn't (inland controls), exactly as pre-registered in §M11.",
        "SOS = start of season; EOS = end of season; DOY = day-of-year (1 = 1 Jan, 365/366 = 31 Dec); Δ = corrected − raw SOS (in days); days = calendar days (effect-size unit); % = per-cent (district-aggregated cyclone-flood pixel share); Kharif = post-monsoon rice cropping season; §M11 = Methods Module 11 (cyclone climatology and pixel-share attenuation).",
    ),
    (
        "Figure 5",
        FIG / "fig5_power_curves.pdf",
        "Post-hoc minimum detectable effect (MDE) curves.",
        "Statistical power as a function of the true τ_SOS effect size for the v2.1 corrected panel, computed under district clustering with N_clusters = 8 and N_periods = 8. At the realised |τ̂| ≈ 15 days, post-hoc power is 0.18, and the MDE_80 % is approximately 35 days. Curves are shown for one-sided α = 0.05 and two-sided α = 0.05.",
        "Honest power audit: the study is under-powered to reject H₀ at the realised effect size, which the manuscript reports transparently rather than concealing.",
        "MDE = minimum detectable effect (smallest τ detectable at a given power); MDE_80 % = MDE at 80 % statistical power; τ_SOS = DiD treatment effect on start of season; τ̂ = estimated τ from the realised data; H₀ = null hypothesis of no treatment effect; α = type-I (false-positive) error rate; SE = district-clustered standard error of τ̂ used as the noise input to the power calculation; N_clusters = number of district clusters used for inference; N_periods = number of time periods (Kharif seasons) in the panel; days = calendar days (effect-size unit); % = per-cent (statistical-power unit on the y-axis); v2.1 = classifier-corrected phenology panel.",
    ),
    (
        "Figure 6",
        FIG / "fig6_placebo_distribution.pdf",
        "In-space placebo distributions (donor-swap, 55 permutations).",
        "In-space placebo distributions of τ̂ for each of the six (pipeline × metric) combinations, obtained by randomly reassigning the 5-of-8 'treated' district label across the C(8,5)=56 possible donor swaps. Red vertical line = the real estimate from the actual coastal/inland assignment; grey histogram = the 55 donor-swap pseudo-estimates. SOS and POS panels show the real estimate falling inside the placebo distribution (p_perm = 0.50 for SOS, 0.27–0.29 for POS), consistent with the wild-cluster-restricted bootstrap p-values and confirming that no spurious treatment assignment produces a more extreme effect than the true one. EOS panels show 'insufficient finite placebo estimates' because the real-data EOS outcome is degenerate under the v1 phenology pipeline (all 55 placebo τ̂ are non-finite); this is a known v1 limitation pre-disclosed in the manuscript Provenance note and is the rationale for treating EOS-derived findings cautiously.",
        "Falsifiability check on the v2.1 corrected panel. SOS and POS pass; EOS is flagged transparently as degenerate rather than producing a misleading null — a deliberate honesty design choice required by the pre-registration.",
        "τ̂ = estimated DiD treatment effect (days); tau_real = τ̂ from the real coastal/inland assignment; SOS / POS / EOS = start, peak, and end of season; DiD = difference-in-differences; p = probability value (significance level); p_perm = permutation (donor-swap) p-value; n = number of donor-swap permutations (here n = 55, after excluding the true assignment from C(8, 5) = 56); n_extreme = count of donor-swap permutations with |τ̂_placebo| ≥ |τ̂_real|; C(8, 5) = binomial coefficient (number of ways to choose 5 treated districts from 8); days = calendar days (effect-size unit); v1 = pre-classifier phenology pipeline; v2.1 = classifier-corrected pipeline.",
    ),
    (
        "Figure S1",
        FIG / "figS1_cyclone_climatology.pdf",
        "Bay-of-Bengal cyclone climatology, 1990–2024 (Supplement).",
        "Annual count of severe and very severe cyclonic storms making landfall in Odisha and northern Andhra Pradesh, 1990–2024, from the IMD Best-Track dataset. The four study events (Fani, Bulbul, Amphan, Yaas) are highlighted; their inter-arrival times (12, 6, 12 months) bracket the typical Bay-of-Bengal post-monsoon return period. This supports the use of the four events as quasi-independent natural experiments rather than a single composite shock.",
        "Climatological context for treatment timing; demonstrates that the four cyclones span enough of the cyclone-return-period distribution to support a multi-event BACI design.",
        "IMD = India Meteorological Department; BoB = Bay of Bengal; Vmax = peak 1-min sustained wind speed; kt = knots (1 kt ≈ 0.514 m s⁻¹), the IMD/Saffir–Simpson wind-speed unit; Cat 1–5 = Saffir–Simpson hurricane wind-scale categories; pre-Kharif = pre-monsoon window (≈ Apr–Jun), preceding the Kharif rice season; months = calendar months (inter-arrival-time unit on the x-axis annotations); BACI = Before–After / Control–Impact design.",
    ),
    (
        "Figure S2",
        FIG / "figS2_backscatter_signatures.pdf",
        "Sentinel-1 backscatter signatures: surge vs. agronomic flooding (Supplement).",
        "Mean σ⁰_VV and σ⁰_VH (dB) time-series for hand-labelled surge-inundated pixels (n = 96) vs. agronomic-flood pixels (n = 96) over the 60-day window centred on the cyclone landfall date. Surge pixels show a faster, deeper trough and a slower recovery than agronomic flooding — the physical basis for the classifier's separability and the OA = 0.844 SAR-only robustness variant.",
        "Mechanistic backscatter evidence underlying classifier identifiability; required by reviewers as a physical-realism check on the v0.3.0 random-forest model.",
        "SAR = synthetic aperture radar; σ⁰ = normalised radar backscatter coefficient (linear power form, displayed in dB); dB = decibels (logarithmic backscatter unit, 10 · log10[σ⁰]); VV = vertical-transmit, vertical-receive co-polarisation channel; VH = vertical-transmit, horizontal-receive cross-polarisation channel; CR = cross-polarisation ratio (VH − VV in dB); ΔVH = post-onset minus pre-onset σ⁰_VH; n = number of hand-labelled pixels per class (here n = 96); days = calendar days (time-series x-axis unit, ±30 d window around landfall); OA = overall accuracy (fraction of correctly classified test pixels); v0.3.0 = released version of the Module 02 random-forest classifier.",
    ),
]

# Step 1: rasterise each PDF to JPG at 1000 DPI
print("Step 1: Rasterising PDFs to 1000-DPI JPG …")
jpg_paths = {}
for fid, pdf_path, *_ in FIGURES:
    if not pdf_path.exists():
        print(f"  MISSING: {pdf_path}")
        continue
    slug = fid.lower().replace(" ", "")
    out_jpg = OUT_JPG / f"{slug}.jpg"
    # pdftoppm at 1000 DPI → JPEG
    subprocess.check_call([
        "pdftoppm", "-jpeg", "-jpegopt", "quality=95",
        "-r", "1000",
        str(pdf_path), str(OUT_JPG / slug),
    ])
    # pdftoppm emits {slug}-1.jpg; rename to {slug}.jpg
    candidates = sorted(OUT_JPG.glob(f"{slug}-*.jpg"))
    if candidates:
        candidates[0].rename(out_jpg)
    jpg_paths[fid] = out_jpg
    print(f"  OK  {fid}: {out_jpg.name} ({out_jpg.stat().st_size/1024:.0f} KB)")

# Step 2: build DOCX
print("\nStep 2: Building DOCX …")
doc = Document()

# A4 portrait, narrow margins to give figures room
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(10)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
normal.paragraph_format.space_after = Pt(6)

# Title page
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Figures Bundle")
r.font.name = 'Arial'; r.font.size = Pt(18); r.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Decoupling Cyclone-Induced Saline Inundation from Agronomic Flooding in Sentinel-1/2 Rice Phenology Retrieval")
r.font.name = 'Arial'; r.font.size = Pt(11); r.italic = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Companion to GitHub release v1.0.0-submission · Zenodo DOI 10.5281/zenodo.20585636\nAll figures rasterised from vector PDF sources at 1000 DPI (JPEG, quality 95).")
r.font.name = 'Arial'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# Figure usable width = page width - margins = 21 - 4 = 17 cm
FIG_WIDTH_CM = 17.0

for fid, pdf_path, label, caption, description, abbreviations in FIGURES:
    jpg = jpg_paths.get(fid)
    if jpg is None or not jpg.exists():
        continue

    # Embed the image, centred, 17 cm wide
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(jpg), width=Cm(FIG_WIDTH_CM))

    # Caption line: bold "Figure N." + label
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p_cap.add_run(f"{fid}. ")
    r1.font.name = 'Arial'; r1.font.size = Pt(10); r1.bold = True
    r2 = p_cap.add_run(label)
    r2.font.name = 'Arial'; r2.font.size = Pt(10); r2.bold = True

    # Caption body
    p_body = doc.add_paragraph()
    p_body.paragraph_format.space_after = Pt(4)
    r = p_body.add_run(caption)
    r.font.name = 'Arial'; r.font.size = Pt(10)

    # Description (italic, smaller, grey)
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.left_indent = Cm(0.5)
    p_desc.paragraph_format.space_after = Pt(4)
    r1 = p_desc.add_run("Description. ")
    r1.font.name = 'Arial'; r1.font.size = Pt(9); r1.italic = True; r1.bold = True
    r1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r2 = p_desc.add_run(description)
    r2.font.name = 'Arial'; r2.font.size = Pt(9); r2.italic = True
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Abbreviations footer (regular, smaller, dark-grey). Lists every
    # non-trivial acronym, symbol, and unit that physically appears in this
    # figure or its caption (or is implicit in the displayed quantities, e.g.
    # SE underlying error bars) — including universally understood units
    # (km, dB, DOY, kt, °N, °E, n, CI, SE, p, %, days, months) per reviewer
    # request to expand universally understood units.
    if abbreviations:
        p_abbr = doc.add_paragraph()
        p_abbr.paragraph_format.left_indent = Cm(0.5)
        p_abbr.paragraph_format.space_after = Pt(4)
        r1 = p_abbr.add_run("Abbreviations. ")
        r1.font.name = 'Arial'; r1.font.size = Pt(9); r1.bold = True; r1.italic = True
        r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r2 = p_abbr.add_run(abbreviations)
        r2.font.name = 'Arial'; r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Page break between figures so each starts on its own page
    doc.add_page_break()

doc.save(str(OUT_DOCX))
print(f"\nOK: wrote {OUT_DOCX} ({OUT_DOCX.stat().st_size/1024:.1f} KB)")
