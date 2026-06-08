"""Build Table_S10 (RF feature importance + falsifiability) and Table_S13a
(per-(district, year, metric) v2.1 correction summary) as standalone docx
files for inclusion in the Supplement bundle.

All inputs are real CSVs in analysis/results/.  No synthetic data is used.
"""
from __future__ import annotations
import csv
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path("/home/user/workspace/RiceBaCI-GEE")
RES = ROOT / "analysis/results"
SUPPL = ROOT / "manuscript/supplement"

# -------------------------------------------------------------- helpers
def _style_doc(doc):
    for s in doc.sections:
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    for st in doc.styles:
        try:
            st.font.name = "Arial"; st.font.size = Pt(10)
        except Exception:
            pass


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(13 if level == 1 else 11)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def _caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = True
    r.font.size = Pt(9)


def _table_from_rows(doc, headers, rows, col_widths_cm=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3A5F")
        tcPr.append(shd)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = tbl.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(str(val)); r.font.size = Pt(9)
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)


# ============================================================== TABLE S10
def build_s10():
    doc = Document(); _style_doc(doc)

    _heading(doc, "Table S10 — Random-forest classifier feature importance "
                  "and falsifiability checks", level=1)

    # --- Panel A: feature importance ---
    _heading(doc, "Panel A. Gini feature importance (full-feature model, "
                  "n_features = 7)", level=2)
    _caption(doc,
        "Gini importance of features in the v0.3.0 saline-flood random-forest "
        "classifier (n_train = 384, n_test = 96, OA = 0.990, F1 = 0.990). "
        "Source: analysis/results/rf_feature_importance_real.csv.")

    feat_path = RES / "rf_feature_importance_real.csv"
    with open(feat_path) as f:
        rdr = list(csv.DictReader(f))
    rows = [(r["rank"], r["feature"], r["label"], f'{float(r["gini_importance"]):.4f}')
            for r in rdr]
    _table_from_rows(doc,
        headers=["Rank", "Feature", "Description", "Gini importance"],
        rows=rows,
        col_widths_cm=[1.2, 3.6, 6.5, 2.5])

    doc.add_paragraph()

    # --- Panel B: falsifiability checks ---
    _heading(doc, "Panel B. Pre-registered falsifiability checks "
                  "(OSF c4mp8, §S3.7)", level=2)
    _caption(doc,
        "Four pre-registered mechanistic checks applied to the v0.3.0 "
        "classifier outputs. All four PASS. Source: "
        "analysis/results/rf_falsifiability_checks_real.csv.")

    chk_path = RES / "rf_falsifiability_checks_real.csv"
    with open(chk_path) as f:
        rdr = list(csv.DictReader(f))
    rows = []
    for r in rdr:
        # Reduce verbose check name to readable form
        chk_label = r["check"].replace("_", " ")
        obs = (r.get("observed_value_db") or
               r.get("observed_value_pct") or
               (f'cyc {r.get("observed_value_cyc","")}, agr '
                f'{r.get("observed_value_agr","")}'
                if r.get("observed_value_cyc") else ""))
        rows.append((chk_label, r["pre_registered_threshold"],
                     obs, r["status"], r["rationale"][:120]))
    _table_from_rows(doc,
        headers=["Check", "Pre-registered threshold", "Observed",
                 "Status", "Rationale (truncated)"],
        rows=rows,
        col_widths_cm=[3.5, 2.5, 2.5, 1.3, 5.5])

    doc.add_paragraph()
    _caption(doc,
        "Together, Panel A and Panel B confirm that the v0.3.0 classifier "
        "satisfies both the statistical-performance criterion (Panel A, "
        "Gini-importance distribution dominated by physically-meaningful "
        "Sentinel-2 NDWI/LSWI and SAR ΔVH features) and the mechanistic "
        "falsifiability criterion (Panel B, all four pre-registered checks "
        "PASS). Cross-references: Section 4.1 (Module 02 acceptance), "
        "Section 4.2 (backscatter confound demonstration), Section 3.7.4 "
        "(falsifiability-test framework).")

    out = SUPPL / "Table_S10_rf_feature_importance.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ============================================================== TABLE S13a
def build_s13a():
    doc = Document(); _style_doc(doc)

    _heading(doc, "Table S13a — Per-(district, year, metric) v2.1 correction "
                  "summary", level=1)
    _caption(doc,
        "All 35 cyclone-affected district-year-metric cells in the v2.1 "
        "real-data panel. For each cell: cyclone-flood pixel share f, "
        "saline-flood delta δ_cyc (days, from the cyclone-mask correction), "
        "raw DOY, corrected DOY, and the resulting correction in days "
        "(corrected − raw). Source: "
        "analysis/results/real_v21/v21_correction_summary.csv.")

    src = RES / "real_v21/v21_correction_summary.csv"
    with open(src) as f:
        rdr = list(csv.DictReader(f))

    rows = []
    for r in rdr:
        rows.append((
            r["district"],
            r["year"],
            r["cyclone"],
            r["metric"],
            f'{float(r["f"])*100:.3f}',          # % share
            r["delta_cyc"],
            f'{float(r["raw_doy"]):.2f}',
            f'{float(r["corrected_doy"]):.2f}',
            f'{float(r["correction_days"]):+.2f}',
        ))

    _table_from_rows(doc,
        headers=["District", "Year", "Cyclone", "Metric",
                 "f (%)", "δ_cyc (d)", "Raw DOY", "Corrected DOY",
                 "Correction (d)"],
        rows=rows,
        col_widths_cm=[2.2, 1.0, 1.5, 1.0, 1.4, 1.4, 1.5, 1.7, 1.7])

    doc.add_paragraph()

    # Summary stats
    sos = [float(r["correction_days"]) for r in rdr if r["metric"] == "SOS"]
    pos = [float(r["correction_days"]) for r in rdr if r["metric"] == "POS"]
    eos = [float(r["correction_days"]) for r in rdr if r["metric"] == "EOS"]
    import statistics as st
    def _stat(v):
        if not v: return ("—","—","—","—")
        return (len(v),
                f'{st.mean([abs(x) for x in v]):.3f}',
                f'{st.stdev([abs(x) for x in v]) if len(v)>1 else 0:.3f}',
                f'{min(v):+.2f} / {max(v):+.2f}')

    _heading(doc, "Summary by phenometric", level=2)
    summ_rows = [
        ("SOS", *_stat(sos)),
        ("POS", *_stat(pos)),
        ("EOS", *_stat(eos)),
    ]
    _table_from_rows(doc,
        headers=["Metric", "n cells", "Mean |Δ| (d)", "SD |Δ| (d)",
                 "Range (d)"],
        rows=summ_rows,
        col_widths_cm=[2.0, 1.8, 2.4, 2.2, 3.6])

    doc.add_paragraph()
    _caption(doc,
        "The largest single-cell correction is Bhadrak / Yaas 2021 at "
        "−1.01 d (SOS), −0.50 d (POS), and −1.51 d (EOS), consistent with "
        "that district's cyclone-flood pixel share f = 7.21 %. Across all "
        "cells the direction of every correction is negative (corrected "
        "DOY < raw DOY), consistent with the cyclone-surge backscatter "
        "trough being mis-interpreted as an early transplanting signal in "
        "the uncorrected pipeline. Cross-references: Section 4.3 (raw vs. "
        "corrected pipeline), Table S5 (jackknife leverage), Figure 4 "
        "(district-level SOS panel).")

    out = SUPPL / "Table_S13a_v21_correction_summary.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build_s10()
    build_s13a()
