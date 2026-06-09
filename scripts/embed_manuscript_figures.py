#!/usr/bin/env python3
"""Embed figure PNGs into manuscript/Manuscript.docx inline at their textual callouts.

Pandoc renders the source manuscript_text.md without embedding figures because
the source has no Markdown ![](path) image syntax — figures are referenced by
caption text only ("Figure 1", "Figure 2", ...). This script post-processes
the pandoc output so every figure appears inline in the rendered DOCX.

Strategy:
  - For each figure file in figures/, locate the first paragraph whose text
    contains the matching callout (e.g. "Figure 1.", "Figure 2.").
  - Insert a new paragraph with the image AFTER the callout-bearing paragraph.
  - Image width = 6.0 in (RSE standard one-column width).

Usage:
    python3 scripts/embed_manuscript_figures.py
"""
from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from copy import deepcopy

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "manuscript" / "Manuscript.docx"
FIGDIR = ROOT / "figures"

# Single source of truth for figure captions: import the FIGURES list that
# scripts/build_figures_bundle.py uses for Figures_Bundle.docx. This guarantees
# the inline manuscript figures and the stand-alone Figures Bundle carry
# identical labels and caption text, and that any Pass 21-style caption fix
# automatically propagates to both outputs.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_figures_bundle import FIGURES as BUNDLE_FIGURES  # noqa: E402

_BUNDLE_CAPTIONS = {fname: (label, caption)
                    for label, fname, caption in BUNDLE_FIGURES}

# Figure embed order: ALWAYS serial (1, 1B, 2, 3, 4, 5, 6, S1, S2). Each
# figure is embedded after the first paragraph whose text matches the callout
# regex AND whose paragraph index is strictly greater than the previous
# figure's anchor index. This guarantees the inline figures appear in serial
# order regardless of where each figure is first textually cited in the
# manuscript (Pass 21c fix: prior versions anchored to the FIRST callout
# globally, producing out-of-order embedding 1 -> 2 -> 1B -> 3 -> 6 -> 5 -> 4
# because Figure 6 is cited in section 3.7.4 before Figure 4/5 callouts).
# The callout regexes are deliberately broad so each figure can find SOME
# downstream callout to anchor against.
FIGURES = [
    ("Figure 1",   FIGDIR / "figure1_study_area.png",
     re.compile(r"\bFigure\s+1\b(?!B)")),
    ("Figure 1B",  FIGDIR / "fig1b_identification_dag.png",
     re.compile(r"\bFigure\s+1B\b")),
    ("Figure 2",   FIGDIR / "fig2_did_coefplot.png",
     re.compile(r"\bFigure\s+2\b")),
    ("Figure 3",   FIGDIR / "fig3_event_study.png",
     re.compile(r"\bFigure\s+3\b")),
    ("Figure 4",   FIGDIR / "fig4_district_sos_panel.png",
     re.compile(r"\bFigure\s+4\b")),
    ("Figure 5",   FIGDIR / "fig5_power_curves.png",
     re.compile(r"\bFigure\s+5\b")),
    ("Figure 6",   FIGDIR / "fig6_placebo_distribution.png",
     re.compile(r"\bFigure\s+6\b")),
    ("Figure S1",  FIGDIR / "figS1_cyclone_climatology.png",
     re.compile(r"\bFigure\s+S1\b")),
    ("Figure S2",  FIGDIR / "figS2_backscatter_signatures.png",
     re.compile(r"\bFigure\s+S2\b")),
]


def insert_paragraph_after(paragraph, doc):
    """Create a new empty paragraph immediately after `paragraph` and return it."""
    new_p = deepcopy(paragraph._p)  # clone for namespacing
    # strip its content
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    # wrap in python-docx Paragraph
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def _resolve_anchors(paragraphs, figures):
    """Return a list of (fig_name, fig_path, anchor_idx, err) in serial order.

    Strategy: pick each figure's natural first-callout paragraph, then enforce
    serial-order monotonicity by promoting any out-of-order anchor to
    max(natural_first_callout, prev_anchor + 1). This guarantees the inline
    figures appear in the order 1, 1B, 2, 3, 4, 5, 6, S1, S2 while still
    placing each figure as close to its natural textual context as possible.
    Figures with no callout anywhere fall back to (prev_anchor + 1) so they
    are still embedded immediately after the previous figure in serial order.
    """
    # First pass: each figure's natural first-callout index (independent).
    naturals = []
    for fig_name, fig_path, callout_re in figures:
        if not fig_path.exists():
            naturals.append((fig_name, fig_path, None, "figure file does not exist"))
            continue
        first = None
        for i, p in enumerate(paragraphs):
            if callout_re.search(p.text):
                first = i
                break
        naturals.append((fig_name, fig_path, first, None))

    # Second pass: enforce monotonicity by promoting out-of-order anchors.
    anchors = []
    prev_idx = -1
    last_valid_p = len(paragraphs) - 1
    for fig_name, fig_path, natural, err in naturals:
        if err == "figure file does not exist":
            anchors.append((fig_name, fig_path, None, err))
            continue
        if natural is None:
            # No callout anywhere: embed immediately after the previous figure
            # so serial order is preserved.
            chosen = min(prev_idx + 1, last_valid_p) if prev_idx >= 0 else 0
            note = "no callout found; placed after previous figure"
            anchors.append((fig_name, fig_path, chosen, note))
            prev_idx = chosen
            continue
        # Promote natural anchor to prev_idx + 1 if it would otherwise break
        # serial order.
        chosen = max(natural, prev_idx + 1)
        note = None if chosen == natural else (
            f"natural callout at idx={natural} promoted to {chosen} "
            f"to preserve serial order")
        anchors.append((fig_name, fig_path, chosen, note))
        prev_idx = chosen
    return anchors


def main():
    if not DOCX.exists():
        print(f"FAIL: {DOCX} does not exist")
        raise SystemExit(1)

    doc = Document(str(DOCX))

    embedded = []
    missing = []

    # Resolve monotonic anchors from the snapshot BEFORE any edits, then walk
    # the figures in REVERSE order so each insertion does not perturb the
    # paragraph indices of figures yet to be inserted.
    snapshot = list(doc.paragraphs)
    anchors = _resolve_anchors(snapshot, FIGURES)

    # Print resolution table for diagnostics.
    print("Anchor resolution (paragraph index in pandoc snapshot):")
    for fig_name, fig_path, idx, err in anchors:
        status = err if err else "OK"
        print(f"  {fig_name:11} -> idx={idx} ({status})")

    # Cache each anchor's paragraph element so reverse iteration can find it
    # after later (earlier-in-document) reverse insertions have added new
    # paragraphs.
    anchor_pairs = []
    for fig_name, fig_path, idx, err in anchors:
        if idx is None:
            missing.append((fig_name, fig_path, err or "no anchor"))
            continue
        anchor_pairs.append((fig_name, fig_path, snapshot[idx]))

    # Reverse iteration so we never invalidate the paragraph references for
    # figures inserted later in the document.
    for fig_name, fig_path, live_target in reversed(anchor_pairs):
        # create new paragraph after the callout and add image
        img_p = insert_paragraph_after(live_target, doc)
        run = img_p.add_run()
        run.add_picture(str(fig_path), width=Inches(6.0))
        img_p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        # Insert caption paragraph immediately after the image so the reader
        # sees "Figure N. <caption text>" below every embedded figure. Pull
        # the caption from the bundle builder's FIGURES table so the inline
        # manuscript caption is identical to the Figures Bundle caption.
        label_caption = _BUNDLE_CAPTIONS.get(fig_path.name)
        if label_caption is not None:
            label, caption_text = label_caption
            cap_p = insert_paragraph_after(img_p, doc)
            cap_p.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT
            label_run = cap_p.add_run(f"{label}. ")
            label_run.bold = True
            label_run.font.size = Pt(10)
            text_run = cap_p.add_run(caption_text)
            text_run.font.size = Pt(10)
        else:
            print(f"  WARN: no bundle caption found for {fig_path.name} "
                  f"(figure embedded without caption)")

        embedded.append(fig_name)
        print(f"  OK: embedded {fig_name} ({fig_path.name}) at chosen anchor")

    doc.save(str(DOCX))
    print()
    print(f"Embedded {len(embedded)} figures in {DOCX.relative_to(ROOT)}")
    if missing:
        print(f"WARNING: {len(missing)} figures NOT embedded:")
        for m in missing:
            print(f"   {m}")
        raise SystemExit(2)


if __name__ == "__main__":
    main()
