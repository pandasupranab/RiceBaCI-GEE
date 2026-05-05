"""
05b_bulbul_transferability.py — out-of-sample transferability probe.

Pre-registered claim
--------------------
The corrected pipeline learns a SOS-shift response to pre-Kharif saline-
surge cyclones (Fani / Amphan / Yaas).  Bulbul (Nov 2019) is a different
class of event:

    • Post-monsoon, not pre-Kharif (no transplanting overlap; Bulbul
      arrived during ripening/maturity, not establishment).
    • Rainfall-dominant, not surge-dominant (landfall on Sagar Island
      ~290 km NE of the study area; Odisha received heavy rainfall but
      no saline ingress).
    • Outside the 8 study districts.

If our corrected pipeline is mechanistic — flood-pixel masking removes a
generic excess-water bias rather than memorising the Fani/Amphan/Yaas
cohort — then the *trained* DiD should transfer to a Bulbul probe panel
with residuals centred near zero. If instead we see large positive
residuals, the corrected pipeline is over-fitting the surge mechanism.

Probe panel
-----------
Built from the same Module 03 phenology pipeline applied to:

    • 5 Bulbul-rainfall districts in coastal Odisha (Mayurbhanj, parts
      of Baleshwar/Bhadrak/Kendrapara not in the main treatment cohort,
      Khordha) — these are the OOS coastal-but-non-surge cells.
    • 5 inland Odisha districts that received Bulbul rainfall but were
      far from the track (Anugul, Dhenkanal, Cuttack, Khordha, Ganjam)
      — these absorb the year FE.

For each Bulbul-affected pixel cohort we compute:
    Δ_obs = SOS_2020_corrected − mean(SOS_pre_corrected, years 2017-2018)

and the prediction:
    Δ_pred = tau_hat_corrected_SOS  (from Table S1)

Residual = Δ_obs − Δ_pred.

The synthetic mode here generates a small Bulbul-ish probe panel
(no real GEE data needed) so the wiring and Table S3 plumbing can be
validated end-to-end.

Outputs
-------
analysis/results/bulbul_transferability.csv
manuscript/supplement/Table_S3_bulbul_transferability.docx (regenerated)

Author : Supranab Panda
Date   : 2026-05-05
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

import numpy as np
import pandas as pd

# ---------------------------------------------------------------------------
# Probe-panel district roster (Bulbul-rainfall extension)
#   Must NOT overlap the main DiD treatment set; uses non-treatment
#   coastal Odisha + inland Odisha districts.
# ---------------------------------------------------------------------------
BULBUL_PROBE_DISTRICTS = [
    # Bulbul rainfall belt: coastal-OUTSIDE-treatment + inland
    ("Mayurbhanj", "coastal_rainfall"),
    ("Khordha",    "coastal_rainfall"),
    ("Ganjam",     "coastal_rainfall"),
    ("Nayagarh",   "inland_rainfall"),
    ("Boudh",      "inland_rainfall"),
    ("Kandhamal",  "inland_rainfall"),
]

PRE_YEARS = [2017, 2018]
EVENT_YEAR = 2019                    # Bulbul = Nov 2019; we measure 2020 SOS
POST_YEAR = 2020                     # first Kharif AFTER Bulbul

BASELINE_DOY_SOS = 195.0             # same climatology as main pipeline
DIST_FE_SD       = 2.0
YEAR_FE_SD       = 1.5
PROBE_NOISE      = 1.4


# ---------------------------------------------------------------------------
def make_probe_panel(seed: int = 20260505,
                     true_bulbul_effect: float = 0.5) -> pd.DataFrame:
    """
    Generate a synthetic Bulbul probe panel.

    `true_bulbul_effect` is the *actual* SOS shift Bulbul induces on
    the probe districts — by default a small positive value (~0.5 d)
    that is much smaller than the trained Fani/Amphan/Yaas effect
    (corrected ~+2 d).  Setting this near zero models the
    "post-monsoon Bulbul barely touches Kharif transplanting" scenario.
    """
    rng = np.random.default_rng(seed + 1)        # different stream from main panel
    rows = []
    for dist, exposure in BULBUL_PROBE_DISTRICTS:
        d_fe = rng.normal(0, DIST_FE_SD)
        for year in PRE_YEARS + [POST_YEAR]:
            y_fe = rng.normal(0, YEAR_FE_SD)
            shock = true_bulbul_effect if year == POST_YEAR else 0.0
            sos = (BASELINE_DOY_SOS + d_fe + y_fe + shock
                   + rng.normal(0, PROBE_NOISE))
            rows.append({
                "district":        dist,
                "exposure":        exposure,
                "year":            year,
                "metric":          "SOS",
                "pipeline":        "corrected",
                "median_sos_doy":  round(sos, 3),
            })
    return pd.DataFrame(rows)


# ---------------------------------------------------------------------------
def compute_transferability(probe_df: pd.DataFrame,
                            tau_pred: float, tau_se: float) -> pd.DataFrame:
    """
    Per-district transferability computation.

    For each district:
      - SOS_pre   = mean(SOS over PRE_YEARS)
      - SOS_post  = SOS in POST_YEAR
      - delta_obs = SOS_post − SOS_pre
      - delta_pred = tau_pred (from main DiD)
      - residual  = delta_obs − delta_pred
      - 95 % PI on prediction = tau_pred ± 1.96 * tau_se
        (treating tau_se as the dominant predictive variance; conservative)
    """
    rows = []
    for dist, sub in probe_df.groupby("district"):
        pre  = sub[sub["year"].isin(PRE_YEARS)]["median_sos_doy"].mean()
        post = sub[sub["year"] == POST_YEAR]["median_sos_doy"].iloc[0]
        delta_obs   = post - pre
        delta_pred  = tau_pred
        residual    = delta_obs - delta_pred
        pi_lo       = delta_pred - 1.96 * tau_se
        pi_hi       = delta_pred + 1.96 * tau_se
        inside      = (delta_obs >= pi_lo) and (delta_obs <= pi_hi)
        rows.append({
            "district":           dist,
            "exposure":           sub["exposure"].iloc[0],
            "delta_obs_d":        round(delta_obs,  3),
            "delta_pred_d":       round(delta_pred, 3),
            "residual_d":         round(residual,   3),
            "pi95_lo":            round(pi_lo,      3),
            "pi95_hi":            round(pi_hi,      3),
            "inside_95pct_PI":    "yes" if inside else "no",
        })
    return pd.DataFrame(rows)


# ---------------------------------------------------------------------------
def regenerate_table_s3(transfer_df: pd.DataFrame, out_path: Path,
                         tau_pred: float, tau_se: float) -> None:
    """Replace the Table S3 stub with realised numbers."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        f"Table S3.  Bulbul (Nov 2019) transferability probe. "
        f"Predicted SOS shift uses the corrected-pipeline DiD "
        f"coefficient \u03c4\u0302 = {tau_pred:+.2f} d "
        f"(SE = {tau_se:.2f}); observed shift is the per-district 2020 SOS "
        f"minus the 2017\u20132018 baseline. Residuals centred near zero "
        f"and inside the 95 % prediction interval indicate the corrected "
        f"pipeline generalises beyond the Fani / Amphan / Yaas training "
        f"cohort to a different cyclone class (post-monsoon, "
        f"rainfall-dominant) outside the 8 study districts."
    )
    run.bold = True
    run.font.size = Pt(10)

    cols = ["district", "exposure", "delta_obs_d", "delta_pred_d",
            "residual_d", "pi95_lo", "pi95_hi", "inside_95pct_PI"]
    labels = ["District", "Exposure class", "\u0394 obs (d)",
              "\u0394 pred (d)", "Residual (d)",
              "PI\u2082.\u2085", "PI\u2089\u2087.\u2085",
              "Inside 95 % PI?"]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, lab in enumerate(labels):
        cell = table.rows[0].cells[i]
        cell.text = lab
        for paragraph in cell.paragraphs:
            for r in paragraph.runs:
                r.bold = True

    for _, r in transfer_df.iterrows():
        row = table.add_row().cells
        for i, c in enumerate(cols):
            row[i].text = str(r[c])

    # Footer summary
    n      = len(transfer_df)
    n_in   = (transfer_df["inside_95pct_PI"] == "yes").sum()
    mean_r = transfer_df["residual_d"].mean()
    sd_r   = transfer_df["residual_d"].std()

    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run(
        f"Summary. {n_in}/{n} districts inside 95 % PI.  "
        f"Mean residual = {mean_r:+.2f} d (SD {sd_r:.2f}). "
        "If 'mean residual is near zero AND \u22657 of 8 districts inside PI', "
        "transferability is supported; if 'mean residual is large positive' "
        "the corrected pipeline over-fits the surge mechanism and Bulbul "
        "is genuinely a different DGP."
    )
    note_run.font.size = Pt(9)

    for paragraph in doc.paragraphs:
        for r in paragraph.runs:
            if r.font.name is None:
                r.font.name = "Arial"

    doc.save(out_path)


# ---------------------------------------------------------------------------
def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--probe-panel", default=None,
                    help="real-data probe panel CSV; if omitted, synthetic mode")
    ap.add_argument("--did-results", default="analysis/results/did_static.csv")
    ap.add_argument("--outdir",      default="analysis/results")
    ap.add_argument("--table-out",
                    default="manuscript/supplement/Table_S3_bulbul_transferability.docx")
    ap.add_argument("--bulbul-effect", type=float, default=0.5,
                    help="(synthetic mode) true SOS shift Bulbul induces; "
                         "set near 0 for 'post-monsoon doesn't matter' scenario")
    ap.add_argument("--seed",         type=int, default=20260505)
    args = ap.parse_args()

    # Load main DiD coefficient (corrected, SOS) for prediction
    did_df = pd.read_csv(args.did_results)
    main_row = did_df.query(
        "pipeline == 'corrected' and metric == 'SOS'"
    ).iloc[0]
    tau_pred = float(main_row["tau_days"])
    tau_se   = float(main_row["se_days"])

    # Probe panel: real or synthetic
    if args.probe_panel:
        probe = pd.read_csv(args.probe_panel)
        print(f"loaded real probe panel: {len(probe)} rows")
        mode = "real"
    else:
        probe = make_probe_panel(
            seed=args.seed,
            true_bulbul_effect=args.bulbul_effect,
        )
        print(f"synthetic probe panel: {len(probe)} rows, "
              f"true Bulbul effect = {args.bulbul_effect:+.2f} d")
        mode = "synthetic"

    transfer = compute_transferability(probe, tau_pred, tau_se)

    out_dir = Path(args.outdir); out_dir.mkdir(parents=True, exist_ok=True)
    out_csv = out_dir / "bulbul_transferability.csv"
    transfer.to_csv(out_csv, index=False)

    table_out = Path(args.table_out)
    table_out.parent.mkdir(parents=True, exist_ok=True)
    regenerate_table_s3(transfer, table_out, tau_pred, tau_se)

    print(f"\nUsing trained \u03c4\u0302 = {tau_pred:+.2f} d "
          f"(SE = {tau_se:.2f}, from {args.did_results})")
    print("\nTransferability per district:")
    print(transfer.to_string(index=False))

    n_in   = (transfer["inside_95pct_PI"] == "yes").sum()
    mean_r = transfer["residual_d"].mean()
    print(f"\nSummary [{mode} mode]: "
          f"{n_in}/{len(transfer)} inside 95 % PI; "
          f"mean residual = {mean_r:+.2f} d")
    print(f"\nwrote: {out_csv}")
    print(f"wrote: {table_out}")


if __name__ == "__main__":
    main()
