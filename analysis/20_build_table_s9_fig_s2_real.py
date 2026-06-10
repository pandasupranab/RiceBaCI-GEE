"""20_build_table_s9_fig_s2_real.py - Build Table S9 + Figure S2 from
real Sentinel-1 RTC backscatter signatures over the 4 paddy probe districts.

Inputs:
  analysis/results/real_v21/s1_backscatter_real_signatures.csv  (85 rows)

Outputs:
  manuscript/supplement/Table_S9_backscatter_signatures.docx
  figures/figS2_backscatter_signatures.pdf  (also .png)
  manuscript/supplement/methods_module12_backscatter_REAL.md (rewrite of Note S3)

Phase definitions (anchored to Odisha Kharif 2019 + Bulbul 9-Nov-2019):
  - PRE-TRANSPLANT BASELINE: doy 121-160 (May 1 - Jun 9)
  - AGRONOMIC TRANSPLANTING: doy 171-220 (Jun 20 - Aug 8)
  - PEAK CANOPY:             doy 230-270 (Aug 18 - Sep 27)
  - BULBUL EVENT WINDOW:     doy 305-325 (Nov 1 - Nov 21; landfall Nov 9)

For each phase x district we compute mean VH (dB), mean VV (dB), mean CR
(linear); we also compute Delta_VH (event vs pre-transplant baseline) for
the agronomic and Bulbul phases.
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "analysis" / "results" / "real_v21" / "s1_backscatter_real_signatures.csv"
OUT_DOCX = ROOT / "manuscript" / "supplement" / "Table_S9_backscatter_signatures.docx"
OUT_PDF = ROOT / "figures" / "figS2_backscatter_signatures.pdf"
OUT_PNG = ROOT / "figures" / "figS2_backscatter_signatures.png"

PHASES = [
    ("Pre-transplant baseline (May)", 121, 160),
    ("Agronomic transplanting (Jun-early Aug)", 171, 220),
    ("Peak canopy (mid-Aug-late Sep)", 230, 270),
    ("Bulbul event window (Nov 1-21)", 305, 325),
]

DISTRICTS = ["Boudh", "Ganjam", "Khordha", "Nayagarh"]


def phase_stats(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for d in DISTRICTS:
        sub = df[df.district == d]
        for phase_name, lo, hi in PHASES:
            mask = (sub.doy_centre >= lo) & (sub.doy_centre <= hi)
            phase = sub[mask]
            if phase.empty:
                rows.append({"district": d, "phase": phase_name,
                             "n_dekads": 0, "mean_VH_dB": np.nan,
                             "mean_VV_dB": np.nan, "mean_CR_lin": np.nan})
                continue
            rows.append({
                "district": d,
                "phase": phase_name,
                "n_dekads": int(len(phase)),
                "mean_VH_dB": round(float(phase["vh_db_mean"].mean()), 2),
                "mean_VV_dB": round(float(phase["vv_db_mean"].mean()), 2),
                "mean_CR_lin": round(float(phase["cr_lin_mean"].mean()), 3),
            })
    return pd.DataFrame(rows)


def add_event_deltas(stats: pd.DataFrame) -> pd.DataFrame:
    """For each district, compute delta-VH vs pre-transplant baseline
    for the agronomic and bulbul phases."""
    out_rows = []
    for d in DISTRICTS:
        sub = stats[stats.district == d].set_index("phase")
        base = sub.loc["Pre-transplant baseline (May)"]
        for phase_name, _, _ in PHASES[1:]:
            row = sub.loc[phase_name]
            dvh = round(float(row["mean_VH_dB"]) - float(base["mean_VH_dB"]), 2)
            dvv = round(float(row["mean_VV_dB"]) - float(base["mean_VV_dB"]), 2)
            dcr = round(float(row["mean_CR_lin"]) - float(base["mean_CR_lin"]), 3)
            out_rows.append({
                "district": d, "phase": phase_name,
                "delta_VH_dB": dvh, "delta_VV_dB": dvv, "delta_CR_lin": dcr,
            })
    return pd.DataFrame(out_rows)


def build_table_s9_docx(stats: pd.DataFrame, deltas: pd.DataFrame, path: Path):
    doc = Document()
    # Page setup
    sect = doc.sections[0]
    sect.left_margin = Inches(0.8); sect.right_margin = Inches(0.8)
    sect.top_margin = Inches(0.8); sect.bottom_margin = Inches(0.8)

    # Title
    t = doc.add_paragraph()
    r = t.add_run("Table S9. Real Sentinel-1 RTC dual-polarisation backscatter "
                  "signatures across rice-phenology phases in the four "
                  "Bulbul probe districts, 2019.")
    r.font.name = "Arial"; r.font.size = Pt(11); r.font.bold = True

    p = doc.add_paragraph()
    r = p.add_run(
        "Source: Microsoft Planetary Computer Sentinel-1 RTC collection "
        "(10 m, IW GRD, gamma_0 terrain-corrected), district-mean monthly "
        "aggregates at 100-m resolution within GADM v4.1 L2 polygons, "
        "2019-05-01 to 2019-12-15. n_dekads = number of 10-day epochs "
        "averaged. VH and VV in dB; CR = VH_linear / VV_linear (linear "
        "ratio). Phase definitions: baseline pre-transplant = DOY 121-160; "
        "agronomic transplanting = DOY 171-220; peak canopy = DOY 230-270; "
        "Bulbul event window (landfall 9-Nov-2019) = DOY 305-325. "
        "Delta values are phase mean minus pre-transplant baseline, per "
        "district. n = 85 district-dekads total across the four districts."
    )
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.italic = True

    # Phase-mean table
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run("(a) Phase means")
    r.font.bold = True; r.font.size = Pt(10); r.font.name = "Arial"

    cols = ["District", "Phase", "n_dekads",
            "Mean VH (dB)", "Mean VV (dB)", "Mean CR"]
    tbl = doc.add_table(rows=1 + len(stats), cols=len(cols))
    tbl.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        cell = tbl.rows[0].cells[i]
        for p_ in cell.paragraphs:
            for run_ in p_.runs:
                run_.font.bold = True
        cell.text = c
        for p_ in cell.paragraphs:
            for run_ in p_.runs:
                run_.font.bold = True; run_.font.size = Pt(9)
                run_.font.name = "Arial"
    for i, (_, row) in enumerate(stats.iterrows(), 1):
        vals = [str(row["district"]), str(row["phase"]),
                str(row["n_dekads"]),
                f'{row["mean_VH_dB"]:+.2f}' if pd.notna(row["mean_VH_dB"]) else "NA",
                f'{row["mean_VV_dB"]:+.2f}' if pd.notna(row["mean_VV_dB"]) else "NA",
                f'{row["mean_CR_lin"]:.3f}' if pd.notna(row["mean_CR_lin"]) else "NA"]
        for j, v in enumerate(vals):
            cell = tbl.rows[i].cells[j]
            cell.text = v
            for p_ in cell.paragraphs:
                for run_ in p_.runs:
                    run_.font.size = Pt(9); run_.font.name = "Arial"

    # Delta table
    doc.add_paragraph()
    h2 = doc.add_paragraph()
    r = h2.add_run("(b) Delta vs pre-transplant baseline")
    r.font.bold = True; r.font.size = Pt(10); r.font.name = "Arial"

    cols2 = ["District", "Phase",
             "Delta VH (dB)", "Delta VV (dB)", "Delta CR"]
    tbl2 = doc.add_table(rows=1 + len(deltas), cols=len(cols2))
    tbl2.style = "Light Grid Accent 1"
    for i, c in enumerate(cols2):
        cell = tbl2.rows[0].cells[i]
        cell.text = c
        for p_ in cell.paragraphs:
            for run_ in p_.runs:
                run_.font.bold = True; run_.font.size = Pt(9)
                run_.font.name = "Arial"
    for i, (_, row) in enumerate(deltas.iterrows(), 1):
        vals = [str(row["district"]), str(row["phase"]),
                f'{row["delta_VH_dB"]:+.2f}',
                f'{row["delta_VV_dB"]:+.2f}',
                f'{row["delta_CR_lin"]:+.3f}']
        for j, v in enumerate(vals):
            cell = tbl2.rows[i].cells[j]
            cell.text = v
            for p_ in cell.paragraphs:
                for run_ in p_.runs:
                    run_.font.size = Pt(9); run_.font.name = "Arial"

    doc.save(path)
    print(f"[OK] wrote {path}")


def build_fig_s2(df: pd.DataFrame, pdf: Path, png: Path):
    plt.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 9,
        "axes.spines.top": False, "axes.spines.right": False,
        "axes.linewidth": 0.6,
    })
    fig, axes = plt.subplots(3, 1, figsize=(7.2, 8.0), sharex=True)
    colors = {"Boudh": "#1f77b4", "Ganjam": "#d62728",
              "Khordha": "#2ca02c", "Nayagarh": "#9467bd"}

    bulbul_doy = 313  # 2019-11-09

    for d in DISTRICTS:
        sub = df[df.district == d].sort_values("doy_centre")
        if sub.empty:
            continue
        axes[0].plot(sub["doy_centre"], sub["vh_db_mean"], "-o",
                     color=colors[d], lw=1.2, ms=4, label=d)
        axes[1].plot(sub["doy_centre"], sub["vv_db_mean"], "-o",
                     color=colors[d], lw=1.2, ms=4, label=d)
        axes[2].plot(sub["doy_centre"], sub["cr_lin_mean"], "-o",
                     color=colors[d], lw=1.2, ms=4, label=d)

    titles = ("VH (dB)", "VV (dB)", "CR = VH/VV (linear)")
    for ax, t in zip(axes, titles):
        ax.set_ylabel(t)
        ax.axvline(bulbul_doy, color="orange", lw=1.5, ls="--", alpha=0.7,
                   label="_nolegend_" if ax is not axes[0] else "Bulbul landfall")
        ax.axvspan(171, 220, color="#cccccc", alpha=0.2, zorder=0)
        ax.grid(alpha=0.25)
    # Legend placed above the figure, outside the axes, so it cannot
    # overlap any data points. Caption text lives only in the supplement.
    axes[0].legend(loc="lower center", bbox_to_anchor=(0.5, 1.02),
                   frameon=False, fontsize=8, ncol=5,
                   handlelength=1.6, columnspacing=1.2)
    axes[2].set_xlabel("Day of year (2019)")

    plt.tight_layout(rect=(0, 0, 1, 0.97))
    fig.savefig(pdf, bbox_inches="tight", pad_inches=0.05)
    fig.savefig(png, bbox_inches="tight", pad_inches=0.05, dpi=1000)
    plt.close(fig)
    print(f"[OK] wrote {pdf}")
    print(f"[OK] wrote {png}")


def main():
    df = pd.read_csv(SRC)
    print(f"[OK] read {len(df)} district-dekads")

    stats = phase_stats(df)
    deltas = add_event_deltas(stats)
    stats.to_csv(ROOT / "analysis" / "results" / "real_v21" /
                 "s1_backscatter_phase_means.csv", index=False)
    deltas.to_csv(ROOT / "analysis" / "results" / "real_v21" /
                  "s1_backscatter_phase_deltas.csv", index=False)

    print("\n--- Phase means ---")
    print(stats.to_string(index=False))
    print("\n--- Deltas vs baseline ---")
    print(deltas.to_string(index=False))

    build_table_s9_docx(stats, deltas, OUT_DOCX)
    build_fig_s2(df, OUT_PDF, OUT_PNG)


if __name__ == "__main__":
    main()
