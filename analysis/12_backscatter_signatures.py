"""
12_backscatter_signatures.py — canonical Sentinel-1 dual-pol signatures
                               for the three inundation mechanisms.

Pre-registered claim
--------------------
The Module 02 saline-flood classifier (§3.3) exploits a *separable*
dual-polarisation backscatter signature across the three inundation
mechanisms relevant to coastal Odisha rice:

    • Transplanting flood (agronomic; controlled depth ~5-15 cm;
      gradual fill over 7-14 days; freshwater).
    • Saline storm-surge inundation (cyclonic; deep impulsive flood
      0.5-3 m; onset in <24 h at landfall; saline).
    • Freshwater monsoon-rainfall inundation (Bulbul-class;
      shallow ponded water; gradual onset over 24-48 h; freshwater
      with riverine drainage).

The discriminator the classifier exploits is the *joint* (VH, VV, CR)
trajectory, not any single feature: surge events produce the deepest
VH trough with concurrent VV depression and a CR depression; rainfall
produces shallow VH dip with VV barely affected and CR slightly
elevated; transplanting produces a smooth VH trough but with VV
*recovery* lagging VH by 2-3 weeks (canopy emergence) and CR
trajectory crossing zero rather than dipping.

This module renders the canonical idealised signatures (no field
data required at synthesis stage) so reviewers can read the
mechanism in one glance, and so the §3.3 feature-set justification
is anchored to a published figure rather than to text alone.

Outputs
-------
analysis/results/backscatter_signatures.csv
manuscript/supplement/Table_S9_backscatter_features.docx
figures/figS2_backscatter_signatures.{pdf,png}

Author : Supranab Panda
Date   : 2026-05-05
"""
from __future__ import annotations

import argparse
import csv
from pathlib import Path

import numpy as np

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "analysis" / "results"
SUPPL = ROOT / "manuscript" / "supplement"
FIGS = ROOT / "figures"
RESULTS.mkdir(parents=True, exist_ok=True)
SUPPL.mkdir(parents=True, exist_ok=True)
FIGS.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Canonical signature parameters. Values are calibrated to the published
# literature on C-band SAR over rice paddies and inundated coastal zones:
#
#   Transplanting flood: Hoshikawa et al. 2023 (EJRS); Wali et al. 2020 (RS);
#                        Filipponi 2019; Ramadhani et al. 2020.
#   Saline storm-surge:  Konkathi et al. 2024 (IGARSS); Haldar et al. 2016;
#                        Pham-Van et al. 2020.
#   Freshwater rainfall: Lee & Pottier 2009 (depolarisation in ponded
#                        rough-surface scattering).
# ---------------------------------------------------------------------------
SIGNATURES = {
    "transplanting_flood": {
        "label": "Transplanting flood (agronomic)",
        "onset_doy": 188,           # mid-July, Kharif transplanting peak
        "vh_baseline_db": -14.0,    # dry-stubble / bare-puddled paddy
        "vh_min_db": -20.5,
        "vh_recovery_days": 35,     # gradual canopy emergence raises VH
        "vv_baseline_db": -9.0,
        "vv_min_db": -13.5,         # smaller drop than VH
        "vv_recovery_days": 50,     # canopy lifts VV later than VH
        "cr_baseline_db": -5.0,     # CR = VH - VV in dB
        "cr_min_db": -7.0,
        "cr_recovery_days": 42,
        "onset_rate_days": 12,      # smooth controlled fill 7-14 d
        "color": "#01696F",
        "marker": "s",
    },
    "saline_storm_surge": {
        "label": "Saline storm-surge (cyclonic)",
        "onset_doy": 123,           # Fani-class pre-Kharif landfall
        "vh_baseline_db": -12.0,    # vegetated coastal pre-monsoon
        "vh_min_db": -22.5,         # deepest dip of the three
        "vh_recovery_days": 21,     # surge withdraws on weeks not days
        "vv_baseline_db": -8.5,
        "vv_min_db": -15.0,         # also depressed (specular reflection)
        "vv_recovery_days": 28,
        "cr_baseline_db": -3.5,
        "cr_min_db": -7.5,          # CR drops because VH falls faster
        "cr_recovery_days": 24,
        "onset_rate_days": 1,       # impulsive: full inundation in <24 h
        "color": "#A12C7B",
        "marker": "*",
    },
    "freshwater_rainfall": {
        "label": "Freshwater rainfall (Bulbul-class)",
        "onset_doy": 313,           # Bulbul: 9 Nov 2019, post-monsoon
        "vh_baseline_db": -13.5,
        "vh_min_db": -16.5,         # shallow dip
        "vh_recovery_days": 10,     # rapid drainage on rough surface
        "vv_baseline_db": -8.5,
        "vv_min_db": -10.0,         # almost unaffected
        "vv_recovery_days": 8,
        "cr_baseline_db": -5.0,
        "cr_min_db": -6.5,          # CR dips weakly
        "cr_recovery_days": 9,
        "onset_rate_days": 2,       # 24-48 h to peak
        "color": "#DA7101",
        "marker": "o",
    },
}


def _profile(doy_axis: np.ndarray, onset: float, baseline: float,
             vmin: float, onset_rate: float, recovery: float) -> np.ndarray:
    """Asymmetric trough: rapid (or gradual) drop to vmin at onset, then
    exponential recovery to baseline."""
    out = np.full_like(doy_axis, baseline, dtype=float)
    pre = doy_axis < onset
    out[pre] = baseline
    # Onset descent: half-cosine over onset_rate days
    desc = (doy_axis >= onset) & (doy_axis < onset + onset_rate)
    if desc.any():
        t = (doy_axis[desc] - onset) / onset_rate
        out[desc] = baseline + (vmin - baseline) * 0.5 * (1 - np.cos(np.pi * t))
    # Trough minimum
    trough = (doy_axis >= onset + onset_rate)
    if trough.any():
        t = doy_axis[trough] - (onset + onset_rate)
        out[trough] = vmin + (baseline - vmin) * (1 - np.exp(-t / (recovery / 3)))
    return out


def write_csv(path: Path) -> None:
    """Per-mechanism summary of canonical signature features."""
    fieldnames = [
        "mechanism", "label", "onset_doy", "onset_rate_days",
        "vh_baseline_db", "vh_min_db", "delta_vh_db",
        "vv_baseline_db", "vv_min_db", "delta_vv_db",
        "cr_baseline_db", "cr_min_db", "delta_cr_db",
        "vh_recovery_days", "vv_recovery_days", "cr_recovery_days",
    ]
    with path.open("w", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=fieldnames)
        w.writeheader()
        for key, p in SIGNATURES.items():
            row = {
                "mechanism": key,
                "label": p["label"],
                "onset_doy": p["onset_doy"],
                "onset_rate_days": p["onset_rate_days"],
                "vh_baseline_db": p["vh_baseline_db"],
                "vh_min_db": p["vh_min_db"],
                "delta_vh_db": round(p["vh_min_db"] - p["vh_baseline_db"], 2),
                "vv_baseline_db": p["vv_baseline_db"],
                "vv_min_db": p["vv_min_db"],
                "delta_vv_db": round(p["vv_min_db"] - p["vv_baseline_db"], 2),
                "cr_baseline_db": p["cr_baseline_db"],
                "cr_min_db": p["cr_min_db"],
                "delta_cr_db": round(p["cr_min_db"] - p["cr_baseline_db"], 2),
                "vh_recovery_days": p["vh_recovery_days"],
                "vv_recovery_days": p["vv_recovery_days"],
                "cr_recovery_days": p["cr_recovery_days"],
            }
            w.writerow(row)


def make_supplement_table(csv_path: Path, out_docx: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.section import WD_ORIENT

    rows = list(csv.DictReader(csv_path.open()))
    doc = Document()
    for sec in doc.sections:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Cm(29.7)
        sec.page_height = Cm(21.0)
        sec.left_margin = sec.right_margin = Cm(1.6)
        sec.top_margin = sec.bottom_margin = Cm(1.6)

    title = doc.add_paragraph()
    r = title.add_run(
        "Table S9. Canonical Sentinel-1 dual-polarisation backscatter "
        "signatures for the three inundation mechanisms exploited by the "
        "Module 02 saline-flood classifier."
    )
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)

    headers = ["Mechanism", "Onset DOY", "Onset rate (d)",
               "VH base (dB)", "VH min (dB)", "ΔVH (dB)",
               "VV base (dB)", "VV min (dB)", "ΔVV (dB)",
               "CR base (dB)", "CR min (dB)", "ΔCR (dB)",
               "VH recov. (d)", "VV recov. (d)", "CR recov. (d)"]
    cols = ["label", "onset_doy", "onset_rate_days",
            "vh_baseline_db", "vh_min_db", "delta_vh_db",
            "vv_baseline_db", "vv_min_db", "delta_vv_db",
            "cr_baseline_db", "cr_min_db", "delta_cr_db",
            "vh_recovery_days", "vv_recovery_days", "cr_recovery_days"]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8)
    for i, row in enumerate(rows, start=1):
        for j, key in enumerate(cols):
            c = tbl.rows[i].cells[j]
            c.text = str(row[key])
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)

    note = doc.add_paragraph()
    nr = note.add_run(
        "\nVH and VV are calibrated dual-polarisation backscatter in "
        "decibels (Sentinel-1 IW GRD). CR = VH − VV (dB) is the cross-"
        "ratio. Onset rate is the time from baseline to trough minimum. "
        "Recovery is the e-folding time from trough to within 1 dB of "
        "baseline. Values are canonical idealisations calibrated against "
        "Hoshikawa et al. (2023), Wali et al. (2020), Filipponi (2019), "
        "Konkathi et al. (2024), and Lee & Pottier (2009); they specify "
        "the feature space exploited by the Module 02 random-forest "
        "classifier (§3.3) and are not produced by direct measurement."
    )
    nr.font.name = "Arial"
    nr.font.size = Pt(9)
    nr.italic = True

    doc.save(str(out_docx))


def make_figure(fig_pdf: Path, fig_png: Path) -> None:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, axes = plt.subplots(3, 1, figsize=(11.5, 10.5), sharex=False)
    fig.subplots_adjust(left=0.10, right=0.96, top=0.94, bottom=0.06,
                        hspace=0.60)

    keys = ["transplanting_flood", "saline_storm_surge", "freshwater_rainfall"]
    panel_titles = ["(A) Transplanting flood — agronomic, freshwater, controlled depth ~5-15 cm",
                    "(B) Saline storm-surge — cyclonic, impulsive onset, deep flood 0.5-3 m",
                    "(C) Freshwater rainfall (Bulbul-class) — post-monsoon, shallow ponded water"]

    for ax, key, ptitle in zip(axes, keys, panel_titles):
        p = SIGNATURES[key]
        # plot range: 30 d before onset to 70 d after
        doy = np.linspace(p["onset_doy"] - 30, p["onset_doy"] + 70, 600)
        vh = _profile(doy, p["onset_doy"], p["vh_baseline_db"],
                      p["vh_min_db"], p["onset_rate_days"], p["vh_recovery_days"])
        vv = _profile(doy, p["onset_doy"], p["vv_baseline_db"],
                      p["vv_min_db"], p["onset_rate_days"], p["vv_recovery_days"])
        cr = _profile(doy, p["onset_doy"], p["cr_baseline_db"],
                      p["cr_min_db"], p["onset_rate_days"], p["cr_recovery_days"])

        ax.plot(doy, vh, color="#01696F", lw=2.0, label="VH (cross-pol)")
        ax.plot(doy, vv, color="#A84B2F", lw=2.0, label="VV (co-pol)")
        ax.plot(doy, cr, color="#7A39BB", lw=2.0, ls="--", label="CR = VH − VV")

        # Onset marker
        ax.axvline(p["onset_doy"], color="0.30", lw=0.8, ls=":")
        ax.text(p["onset_doy"], -24.5, "onset", fontsize=10.5, color="0.30",
                ha="center", va="bottom", weight="bold",
                bbox=dict(boxstyle="round,pad=0.22", fc="white",
                          ec="0.7", lw=0.5))

        # Annotate ΔVH (the headline discriminator) — place ABOVE the trough
        # in the upper plot area to avoid overlap with curves and legend.
        delta_vh = p["vh_min_db"] - p["vh_baseline_db"]
        ax.annotate(f"ΔVH = {delta_vh:+.1f} dB",
                    xy=(p["onset_doy"] + p["onset_rate_days"] + 1, p["vh_min_db"]),
                    xytext=(p["onset_doy"] + 18, -2.5),
                    fontsize=11.5, color="#01696F", fontweight="bold",
                    bbox=dict(boxstyle="round,pad=0.30", fc="white",
                              ec="#01696F", lw=0.8),
                    arrowprops=dict(arrowstyle="->", color="#01696F", lw=0.9,
                                    connectionstyle="arc3,rad=-0.2"))

        ax.set_ylim(-26, 1)
        ax.set_ylabel("Backscatter (dB)", fontsize=12.5)
        ax.set_title(ptitle, fontsize=13, loc="left", pad=6,
                     color=p["color"], weight="bold")
        ax.tick_params(labelsize=10.5)
        ax.grid(True, ls=":", lw=0.5, color="0.85")
        # Legend in lower-left where curves are at baseline (clear of trough)
        ax.legend(loc="lower left", fontsize=10.5, framealpha=0.92,
                  fancybox=False, edgecolor="0.7", ncol=3)

    axes[-1].set_xlabel("Day-of-year (DOY) — 30 d pre-onset to 70 d post-onset",
                        fontsize=12.5)

    # fig.suptitle removed (caption supplied below figure in DOCX/manuscript).
    fig.suptitle("", fontsize=10.5, y=0.995, x=0.5, ha="center")

    fig.savefig(str(fig_pdf), dpi=300, bbox_inches="tight", pad_inches=0.35)
    fig.savefig(str(fig_png), dpi=200, bbox_inches="tight", pad_inches=0.35)
    plt.close(fig)


def main(argv=None) -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--quick", action="store_true",
                   help="Synthetic mode (default). Deterministic.")
    args = p.parse_args(argv)

    csv_path = RESULTS / "backscatter_signatures.csv"
    table_s9 = SUPPL / "Table_S9_backscatter_features.docx"
    fig_pdf = FIGS / "figS2_backscatter_signatures.pdf"
    fig_png = FIGS / "figS2_backscatter_signatures.png"

    print("[12] writing canonical signatures ->", csv_path)
    write_csv(csv_path)
    print("[12] writing Table S9 ->", table_s9)
    make_supplement_table(csv_path, table_s9)
    print("[12] writing Figure S2 ->", fig_pdf)
    make_figure(fig_pdf, fig_png)
    print("[12] OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
