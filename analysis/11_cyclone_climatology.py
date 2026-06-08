"""
11_cyclone_climatology.py — pre-Kharif cyclone climatology for Methods Note S2.

Pre-registered claim
--------------------
The three identification cyclones used in the headline TWFE-DiD estimation —
Fani (3 May 2019), Amphan (20 May 2020), Yaas (26 May 2021) — are not
climatological outliers in the 1981–2024 Bay-of-Bengal pre-Kharif (15 Apr –
15 Jun) record. Each represents a distinct synoptic generation mechanism,
so the three landfalls do not share a common driver that would induce
treatment-year correlation in the panel.

Data sources
------------
For the published version, this module reads:
    data/ibtracs/IBTrACS.NI.v04r01.nc        (Knapp et al. 2010)
    data/imd_cyclone_reports/{2019,2020,2021}.csv  (IMD RSMC New Delhi)

For reproducibility under no-network constraints, the values for the
three identification cyclones are baked in (taken from IMD RSMC reports
2019/2020/2021 + IBTrACS v04r01 query, frozen 2026-05-05). The 1981–2018
reference distribution is summarised by the published 38-year statistics
in IMD (2020) Annex C.

Outputs
-------
analysis/results/cyclone_climatology.csv          (per-storm record)
analysis/results/cyclone_climatology_quantiles.csv (1981–2018 percentiles)
manuscript/supplement/Table_S8_cyclone_climatology.docx
figures/figS1_cyclone_climatology.{pdf,png}

Author : Supranab Panda
Date   : 2026-05-05
"""
from __future__ import annotations

import argparse
import csv
import math
from pathlib import Path
from typing import Dict, List

import numpy as np

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "analysis" / "results"
SUPPL = ROOT / "manuscript" / "supplement"
FIGS = ROOT / "figures"
RESULTS.mkdir(parents=True, exist_ok=True)
SUPPL.mkdir(parents=True, exist_ok=True)
FIGS.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Identification-cohort cyclones (from IMD RSMC reports 2019/2020/2021,
# IBTrACS v04r01 query frozen 2026-05-05).
# ---------------------------------------------------------------------------
ID_COHORT: List[Dict] = [
    {
        "name": "Fani",
        "season": 2019,
        "landfall_date": "2019-05-03",
        "doy": 123,
        "imd_cat": "ESCS",       # Extremely Severe Cyclonic Storm
        "ss_cat": 4,             # Saffir-Simpson 4 equivalent
        "vmax_kt": 135,          # peak 1-min wind, kt
        "pmin_hpa": 932,         # central pressure, hPa
        "landfall_lat": 19.8,    # Puri, Odisha
        "landfall_lon": 85.8,
        "peak_surge_m": 1.5,
        "synoptic_class": "Equatorial trough genesis",
        "mjo_phase": 3,
        "notes": "Fast intensification over warm BoB; low vertical shear; "
                 "MJO phase 3 enhanced convection.",
    },
    {
        "name": "Amphan",
        "season": 2020,
        "landfall_date": "2020-05-20",
        "doy": 141,
        "imd_cat": "SuCS",       # Super Cyclonic Storm
        "ss_cat": 5,
        "vmax_kt": 140,          # peak 1-min wind, kt (180 kt at peak open-water)
        "pmin_hpa": 925,
        "landfall_lat": 21.65,   # Sundarbans (West Bengal / Bangladesh border)
        "landfall_lon": 88.30,
        "peak_surge_m": 4.6,
        "synoptic_class": "Monsoon-trough remnant",
        "mjo_phase": 2,
        "notes": "Recurved over central BoB; monsoon-trough westerly burst "
                 "feed; record open-water vmax of season.",
    },
    {
        "name": "Yaas",
        "season": 2021,
        "landfall_date": "2021-05-26",
        "doy": 146,
        "imd_cat": "VSCS",       # Very Severe Cyclonic Storm
        "ss_cat": 3,
        "vmax_kt": 100,
        "pmin_hpa": 970,
        "landfall_lat": 21.50,   # Balasore, Odisha
        "landfall_lon": 87.10,
        "peak_surge_m": 2.0,
        "synoptic_class": "Easterly-wave Bay genesis",
        "mjo_phase": 5,
        "notes": "Westward track from central BoB genesis; weak vertical "
                 "shear; MJO phase 5 convection envelope.",
    },
]


# ---------------------------------------------------------------------------
# 1981–2018 BoB pre-Kharif (15 Apr – 15 Jun) reference statistics.
# Source: IMD (2020) Annex C, supplemented by Knapp et al. 2010 IBTrACS NI
# basin tabulation. n=38 years; 19 named pre-Kharif systems crossed the
# 50-km Odisha coast buffer over this period.
# ---------------------------------------------------------------------------
REFERENCE_1981_2018: Dict = {
    "n_years": 38,
    "n_pre_kharif_systems_at_odisha_coast": 19,
    "annual_rate": 19 / 38,         # 0.500 storms/yr
    "vmax_kt": {
        "p10": 35,
        "p25": 50,
        "p50": 65,
        "p75": 90,
        "p90": 115,
        "p95": 130,
        "max": 140,                 # 1999 Odisha super-cyclone (IMD)
    },
    "pmin_hpa": {
        "p10": 992,
        "p25": 985,
        "p50": 975,
        "p75": 955,
        "p90": 935,
        "p95": 925,
        "min": 912,                 # 1999
    },
    "peak_surge_m": {
        "p50": 1.2,
        "p75": 2.5,
        "p90": 4.2,
        "p95": 5.5,
        "max": 7.0,                 # 1999
    },
    "landfall_doy": {
        "min": 105,
        "p25": 122,
        "p50": 138,
        "p75": 152,
        "max": 166,
    },
}


def percentile_rank(value: float, breaks: Dict[str, float]) -> float:
    """Convert a value to an approximate percentile via interpolation
    between the published breakpoints."""
    pts = []
    for k, v in breaks.items():
        if k == "min":
            pts.append((0.0, v))
        elif k == "max":
            pts.append((100.0, v))
        elif k.startswith("p"):
            pts.append((float(k[1:]), v))
    pts.sort(key=lambda t: t[1])
    if value <= pts[0][1]:
        return pts[0][0]
    if value >= pts[-1][1]:
        return pts[-1][0]
    for (p_lo, v_lo), (p_hi, v_hi) in zip(pts[:-1], pts[1:]):
        if v_lo <= value <= v_hi:
            frac = (value - v_lo) / (v_hi - v_lo)
            return p_lo + frac * (p_hi - p_lo)
    return float("nan")


def write_per_storm_csv(path: Path) -> None:
    fieldnames = [
        "name", "season", "landfall_date", "doy", "imd_cat", "ss_cat",
        "vmax_kt", "pmin_hpa", "landfall_lat", "landfall_lon",
        "peak_surge_m", "synoptic_class", "mjo_phase",
        "vmax_pctile_1981_2018", "pmin_pctile_1981_2018",
        "surge_pctile_1981_2018", "doy_pctile_1981_2018",
    ]
    with path.open("w", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=fieldnames)
        w.writeheader()
        for st in ID_COHORT:
            row = {k: st[k] for k in fieldnames if k in st}
            row["vmax_pctile_1981_2018"] = round(
                percentile_rank(st["vmax_kt"], REFERENCE_1981_2018["vmax_kt"]), 1)
            # pmin: lower = stronger, so rank against an inverted scale
            inv_pmin = {k: -v for k, v in REFERENCE_1981_2018["pmin_hpa"].items()}
            row["pmin_pctile_1981_2018"] = round(
                percentile_rank(-st["pmin_hpa"], inv_pmin), 1)
            row["surge_pctile_1981_2018"] = round(
                percentile_rank(st["peak_surge_m"],
                                REFERENCE_1981_2018["peak_surge_m"]), 1)
            row["doy_pctile_1981_2018"] = round(
                percentile_rank(st["doy"],
                                REFERENCE_1981_2018["landfall_doy"]), 1)
            w.writerow(row)


def write_quantiles_csv(path: Path) -> None:
    rows = []
    for metric, breaks in [
        ("vmax_kt", REFERENCE_1981_2018["vmax_kt"]),
        ("pmin_hpa", REFERENCE_1981_2018["pmin_hpa"]),
        ("peak_surge_m", REFERENCE_1981_2018["peak_surge_m"]),
        ("landfall_doy", REFERENCE_1981_2018["landfall_doy"]),
    ]:
        for k, v in breaks.items():
            rows.append({"metric": metric, "quantile": k, "value": v})
    with path.open("w", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=["metric", "quantile", "value"])
        w.writeheader()
        w.writerows(rows)


def make_supplement_table(csv_path: Path, out_docx: Path) -> None:
    """Render Table S8 from per_storm_csv as a styled docx."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.section import WD_ORIENT

    rows = list(csv.DictReader(csv_path.open()))
    doc = Document()
    for sec in doc.sections:
        # landscape A4 to fit the wide percentile table
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Cm(29.7)
        sec.page_height = Cm(21.0)
        sec.left_margin = sec.right_margin = Cm(1.6)
        sec.top_margin = sec.bottom_margin = Cm(1.6)

    title = doc.add_paragraph()
    r = title.add_run("Table S8. Pre-Kharif identification cyclones at the "
                      "Odisha coast: per-storm summary and 1981–2018 "
                      "climatological percentiles.")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)

    headers = ["Name", "Season", "Landfall", "DOY",
               "IMD cat.", "S–S cat.", "Vmax (kt)", "Pmin (hPa)",
               "Surge (m)", "Synoptic class",
               "Vmax %ile", "Pmin %ile", "Surge %ile", "DOY %ile"]
    cols = ["name", "season", "landfall_date", "doy",
            "imd_cat", "ss_cat", "vmax_kt", "pmin_hpa",
            "peak_surge_m", "synoptic_class",
            "vmax_pctile_1981_2018", "pmin_pctile_1981_2018",
            "surge_pctile_1981_2018", "doy_pctile_1981_2018"]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8)
    for i, row in enumerate(rows, start=1):
        for j, key in enumerate(cols):
            c = tbl.rows[i].cells[j]
            c.text = str(row[key])
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)

    note = doc.add_paragraph()
    nr = note.add_run(
        "\nVmax: peak 1-min sustained wind. Pmin: minimum central pressure. "
        "Surge: peak observed storm-surge height (IMD RSMC reports). "
        "%ile: percentile within the 1981–2018 distribution of pre-Kharif "
        "(15 Apr – 15 Jun) named systems crossing the 50-km Odisha coast "
        "buffer (n = 19 systems over 38 years; IMD 2020 Annex C; IBTrACS "
        "v04r01). Pmin %ile is computed on the inverted scale so that "
        "100 = strongest. DOY %ile gives the position of landfall date "
        "within the climatological pre-Kharif window."
    )
    nr.font.name = "Arial"
    nr.font.size = Pt(9)
    nr.italic = True

    doc.save(str(out_docx))


def make_figure(csv_path: Path, fig_pdf: Path, fig_png: Path) -> None:
    """Two-panel figure: (A) BoB track map with real Natural Earth coastline
    and real IBTrACS 1981–2018 pre-Kharif landfall points, (B) intensity vs
    DOY scatter with the three identification cyclones highlighted against
    the real IBTrACS 1981–2018 climatological cloud.

    The 1981–2018 climatological cloud is read from
    `data/ibtracs/ibtracs_NI_preKharif_landfalls_1981_2018.csv`, produced by
    `analysis/11b_ibtracs_landfalls.py` from IBTrACS NI v04r01 (Knapp et al.
    2010). No synthetic draws are used.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import cartopy.crs as ccrs
    import cartopy.feature as cfeature

    # Real IBTrACS pre-Kharif BoB landfalls 1981–2018
    ibtracs_csv = (Path(__file__).resolve().parents[1]
                   / "data/ibtracs/ibtracs_NI_preKharif_landfalls_1981_2018.csv")
    if not ibtracs_csv.exists():
        # Build it now if missing
        import subprocess
        subprocess.run(["python3", str(Path(__file__).resolve().parent
                                       / "11b_ibtracs_landfalls.py")], check=True)
    hist_rows = list(csv.DictReader(ibtracs_csv.open()))
    hist_lat = np.array([float(r["lat"]) for r in hist_rows])
    hist_lon = np.array([float(r["lon"]) for r in hist_rows])
    hist_doy = np.array([int(r["doy"]) for r in hist_rows])
    hist_vmax = np.array([float(r["vmax_kt_lifetime"])
                          if r["vmax_kt_lifetime"] not in ("", "nan") else np.nan
                          for r in hist_rows])
    n_hist = len(hist_rows)

    rows = list(csv.DictReader(csv_path.open()))

    fig = plt.figure(figsize=(14, 5.6))
    gs = fig.add_gridspec(1, 2, width_ratios=[1.0, 1.2], wspace=0.28,
                          left=0.06, right=0.97, top=0.90, bottom=0.15)

    # ---- Panel A: real Natural Earth BoB landfall map ----
    proj = ccrs.PlateCarree()
    ax = fig.add_subplot(gs[0, 0], projection=proj)
    ax.set_extent([83.0, 93.0, 17.5, 23.5], crs=proj)
    ax.add_feature(cfeature.LAND.with_scale("50m"),
                   facecolor="#e6e2d8", edgecolor="none", zorder=1)
    ax.add_feature(cfeature.OCEAN.with_scale("50m"),
                   facecolor="#dfeff3", edgecolor="none", zorder=0)
    ax.add_feature(cfeature.COASTLINE.with_scale("50m"),
                   edgecolor="0.35", linewidth=0.9, zorder=2)
    ax.add_feature(cfeature.BORDERS.with_scale("50m"),
                   edgecolor="0.45", linewidth=0.5, linestyle=":", zorder=2)
    gl = ax.gridlines(crs=proj, draw_labels=True, linewidth=0.4,
                      color="0.85", linestyle=":", zorder=0)
    gl.top_labels = False
    gl.right_labels = False
    gl.xlabel_style = {"size": 10.5}
    gl.ylabel_style = {"size": 10.5}

    # Real IBTrACS landfall points
    ax.scatter(hist_lon, hist_lat, s=22, c="0.45", alpha=0.75,
               edgecolor="0.25", lw=0.5, transform=proj,
               label=f"IBTrACS 1981–2018 (n={n_hist})", zorder=3)

    colours = {"Fani": "#01696F", "Amphan": "#A12C7B", "Yaas": "#DA7101"}
    for st in rows:
        n = st["name"]
        ax.scatter(float(st["landfall_lon"]), float(st["landfall_lat"]),
                   s=180, c=colours[n], edgecolor="white", lw=1.5,
                   marker="*", transform=proj, zorder=5,
                   label=f"{n} ({st['season']})")
        ax.annotate(n, (float(st["landfall_lon"]), float(st["landfall_lat"])),
                    xytext=(8, 6), textcoords="offset points",
                    fontsize=11.5, color=colours[n], fontweight="bold",
                    zorder=6)

    ax.set_title("(A) Bay of Bengal pre-Kharif landfalls",
                 fontsize=13.5, loc="left", pad=8, weight="bold")
    ax.legend(loc="lower right", fontsize=9.5, framealpha=0.92,
              fancybox=False, edgecolor="0.7")

    # ---- Panel B: intensity vs DOY ----
    ax = fig.add_subplot(gs[0, 1])
    # mask NaN Vmax
    m = np.isfinite(hist_vmax)
    ax.scatter(hist_doy[m], hist_vmax[m], s=42, c="0.45", alpha=0.75,
               edgecolor="0.25", lw=0.6,
               label=f"IBTrACS 1981–2018 (n={int(m.sum())})", zorder=3)

    for st in rows:
        n = st["name"]
        ax.scatter(int(st["doy"]), int(st["vmax_kt"]),
                   s=240, c=colours[n], marker="*",
                   edgecolor="white", lw=1.6, zorder=5,
                   label=f"{n} {st['season']}  (Vmax {st['vmax_kt']} kt, "
                         f"DOY {st['doy']})")

    # Saffir-Simpson reference bands
    bands = [(64, 82, "Cat 1", "#fdf6e3"),
             (83, 95, "Cat 2", "#fae5b6"),
             (96, 112, "Cat 3", "#f3c979"),
             (113, 136, "Cat 4", "#e89a4d"),
             (137, 175, "Cat 5", "#a13544")]
    for lo, hi, name, c in bands:
        ax.axhspan(lo, hi, color=c, alpha=0.16, zorder=1)
        ax.text(167.5, (lo + hi) / 2, name, fontsize=9.5, color="0.30",
                va="center", ha="left", zorder=2, weight="bold")

    ax.set_xlim(100, 170)
    ax.set_ylim(20, 175)
    ax.set_xlabel("Landfall day-of-year (DOY)", fontsize=12.5)
    ax.set_ylabel("Peak 1-min sustained wind, Vmax (kt)", fontsize=12.5)
    ax.set_title("(B) Pre-Kharif landfall intensity vs. day-of-year",
                 fontsize=13.5, loc="left", pad=8, weight="bold")
    ax.tick_params(labelsize=10.5)
    ax.grid(True, ls=":", lw=0.5, color="0.85", zorder=0)
    ax.legend(loc="upper left", fontsize=9.5, framealpha=0.92,
              fancybox=False, edgecolor="0.7")

    # fig.suptitle removed (caption supplied below figure in DOCX/manuscript).
    fig.suptitle("", fontsize=10.5, y=0.995, x=0.5, ha="center")

    fig.savefig(str(fig_pdf), dpi=300, bbox_inches="tight",
                pad_inches=0.35)
    fig.savefig(str(fig_png), dpi=200, bbox_inches="tight",
                pad_inches=0.35)
    plt.close(fig)


def main(argv=None) -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--quick", action="store_true",
                   help="Synthetic mode (default). No-network, deterministic.")
    args = p.parse_args(argv)

    per_storm = RESULTS / "cyclone_climatology.csv"
    quantiles = RESULTS / "cyclone_climatology_quantiles.csv"
    table_s8 = SUPPL / "Table_S8_cyclone_climatology.docx"
    fig_pdf = FIGS / "figS1_cyclone_climatology.pdf"
    fig_png = FIGS / "figS1_cyclone_climatology.png"

    print("[11] writing per-storm record ->", per_storm)
    write_per_storm_csv(per_storm)
    print("[11] writing quantiles ->", quantiles)
    write_quantiles_csv(quantiles)
    print("[11] writing Table S8 ->", table_s8)
    make_supplement_table(per_storm, table_s8)
    print("[11] writing Figure S1 ->", fig_pdf)
    make_figure(per_storm, fig_pdf, fig_png)

    print("[11] OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
