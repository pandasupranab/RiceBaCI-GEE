"""
07_supplement_tables.py — generate journal-ready supplementary tables.

Outputs (manuscript/supplement/):
    Table_S1_did_static.docx        — full coefficient table (DiD ATT)
    Table_S2_pretrends.docx         — pre-period interaction tests
    Table_S3_bulbul_transferability.docx — Bulbul transferability probe stub
    Table_S1_S3_combined.csv        — machine-readable consolidation

The DOCX path is intentional — Elsevier and Frontiers both accept it
directly without conversion.

Author: Supranab Panda
Date  : 2026-05-05
"""
from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
def _set_table_style(table, font_size: int = 10):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Arial"


def _style_header(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def _add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"


# ---------------------------------------------------------------------------
def table_s1_did_static(static_df: pd.DataFrame, out_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    _add_caption(doc,
                 "Table S1.  Static difference-in-differences estimates of "
                 "pre-Kharif cyclone effect on rice phenology (8 districts × "
                 "8 years; SEs clustered at the district level).")

    cols = ["pipeline", "metric", "n_obs", "n_districts",
            "tau_days", "se_days", "t_stat", "p_value",
            "ci_lo_95", "ci_hi_95", "r2_within"]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = {
            "pipeline":    "Pipeline",
            "metric":      "Metric",
            "n_obs":       "N obs",
            "n_districts": "N dist.",
            "tau_days":    "τ (d)",
            "se_days":     "SE",
            "t_stat":      "t",
            "p_value":     "p",
            "ci_lo_95":    "CI₂.₅",
            "ci_hi_95":    "CI₉₇.₅",
            "r2_within":   "R²ᵥᵥ",
        }[c]
    _style_header(table.rows[0])

    for _, r in static_df.iterrows():
        row = table.add_row().cells
        for i, c in enumerate(cols):
            v = r[c]
            if isinstance(v, float):
                row[i].text = f"{v:+.3f}" if c == "tau_days" else f"{v:.3f}"
            else:
                row[i].text = str(v)
    _set_table_style(table)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "Notes. τ is the average treatment effect on the treated, in days "
        "of phenology shift. *** p<0.001; ** p<0.01; * p<0.05. "
        "R²ᵥᵥ is the partial within-R² of the DiD term after absorbing "
        "district and year fixed effects. Treatment cohort: coastal-treatment "
        "districts (Baleshwar, Bhadrak, Kendrapara, Jagatsinghpur, Puri) × "
        "cyclone years 2019, 2020, 2021. Bulbul (2019, post-monsoon, "
        "outside the study area) and Hudhud (2014) are excluded as "
        "transferability hold-outs."
    ).font.size = Pt(9)

    doc.save(out_path)


# ---------------------------------------------------------------------------
def table_s2_pretrends(pre_df: pd.DataFrame, out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Arial"; style.font.size = Pt(10)

    _add_caption(doc,
                 "Table S2.  Parallel-trends test: regression of phenology "
                 "metric on year × treat interaction in pre-treatment period "
                 "(years < 2019). A non-significant coefficient supports the "
                 "DiD identifying assumption.")

    cols = ["pipeline", "metric", "interaction_coef", "se", "p_value",
            "n_pre", "note"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    labels = ["Pipeline", "Metric", "β (year×treat)", "SE", "p",
              "N pre", "Verdict"]
    for i, lab in enumerate(labels):
        hdr[i].text = lab
    _style_header(table.rows[0])

    for _, r in pre_df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["pipeline"])
        row[1].text = str(r["metric"])
        row[2].text = f"{r['interaction_coef']:+.3f}"
        row[3].text = f"{r['se']:.3f}"
        row[4].text = f"{r['p_value']:.4f}"
        row[5].text = str(int(r["n_pre"]))
        row[6].text = str(r["note"])
    _set_table_style(table)

    doc.save(out_path)


# ---------------------------------------------------------------------------
def table_s3_bulbul_stub(out_path: Path) -> None:
    """Stub table — populated after Module 05b transferability run."""
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Arial"; style.font.size = Pt(10)

    _add_caption(doc,
                 "Table S3.  Bulbul (Nov 2019) transferability probe: "
                 "out-of-sample prediction of inland Bulbul-affected districts' "
                 "SOS using coefficients trained on Fani / Amphan / Yaas. "
                 "[Populated by Module 05b after main DiD is locked.]")

    cols = ["district", "predicted_SOS_shift_d", "observed_SOS_shift_d",
            "residual_d", "within_95pct_PI"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    hdr_labels = ["District", "Predicted ΔSOS (d)", "Observed ΔSOS (d)",
                  "Residual (d)", "Inside 95% PI?"]
    for i, lab in enumerate(hdr_labels):
        table.rows[0].cells[i].text = lab
    _style_header(table.rows[0])

    # Placeholder rows for the 8 districts so reviewers see the shape
    placeholders = [
        ("Baleshwar",     "—", "—", "—", "—"),
        ("Bhadrak",       "—", "—", "—", "—"),
        ("Kendrapara",    "—", "—", "—", "—"),
        ("Jagatsinghpur", "—", "—", "—", "—"),
        ("Puri",          "—", "—", "—", "—"),
        ("Dhenkanal",     "—", "—", "—", "—"),
        ("Anugul",        "—", "—", "—", "—"),
        ("Cuttack",       "—", "—", "—", "—"),
    ]
    for vals in placeholders:
        row = table.add_row().cells
        for i, v in enumerate(vals):
            row[i].text = v
    _set_table_style(table)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "Notes. Awaiting Module 05b (transferability) execution. "
        "Prediction uses the corrected-pipeline DiD coefficients from "
        "Table S1 to forecast Bulbul-induced SOS shift; observed shift "
        "is computed from the corrected phenology pipeline applied to "
        "Bulbul (Nov 2019). Residuals centred near zero (and within the "
        "95% prediction interval) indicate the trained model generalises "
        "to a different cyclone class (post-monsoon, rainfall-dominant) "
        "outside the study area."
    ).font.size = Pt(9)

    doc.save(out_path)


# ---------------------------------------------------------------------------
def table_s4_wild_bootstrap(wcb_df: pd.DataFrame, out_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    _add_caption(doc,
                 "Table S4.  Wild-cluster bootstrap inference (Cameron, Gelbach "
                 "& Miller 2008) for the DiD coefficient. Rademacher weights, "
                 "residuals imposed under the null. CIs by inversion on a "
                 "41-point grid.")

    cols   = ["pipeline", "metric", "tau_hat", "t_obs", "B",
              "p_wcb_2sided", "ci_lo_95_wcb", "ci_hi_95_wcb"]
    labels = ["Pipeline", "Metric", "\u03c4 (d)", "t (CR1)", "B",
              "p (WCR)", "CI\u2082.\u2085 (WCR)", "CI\u2089\u2087.\u2085 (WCR)"]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, lab in enumerate(labels):
        table.rows[0].cells[i].text = lab
    _style_header(table.rows[0])

    for _, r in wcb_df.iterrows():
        row = table.add_row().cells
        for i, c in enumerate(cols):
            v = r[c]
            if c == "tau_hat":
                row[i].text = f"{v:+.3f}"
            elif isinstance(v, float):
                row[i].text = f"{v:.4f}" if c == "p_wcb_2sided" else f"{v:.3f}"
            else:
                row[i].text = str(v)
    _set_table_style(table)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "Notes. WCR = wild cluster restricted (residuals imposed under H0: \u03c4=0). "
        "Floor on bootstrap p-value at B=9999 is 1/(B+1) = 0.0001. "
        "District clusters: 8."
    ).font.size = Pt(9)
    doc.save(out_path)


# ---------------------------------------------------------------------------
def table_s5_jackknife(jk_df: pd.DataFrame, verdict_df: pd.DataFrame,
                       out_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    _add_caption(doc,
                 "Table S5a.  Leave-one-district-out (LOO) sensitivity "
                 "verdicts. Each cell shows the maximum percentage change in "
                 "\u03c4\u0302 across the 8 LOO refits and the most-influential "
                 "dropped district.")

    cols  = ["pipeline", "metric", "tau_full",
             "max_abs_delta_pct", "most_leveraging_district", "verdict"]
    labs  = ["Pipeline", "Metric", "\u03c4\u0302 full (d)",
             "max |\u0394\u03c4| (%)", "Driver district", "Verdict"]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, lab in enumerate(labs):
        table.rows[0].cells[i].text = lab
    _style_header(table.rows[0])

    for _, r in verdict_df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["pipeline"])
        row[1].text = str(r["metric"])
        row[2].text = f"{r['tau_full']:+.3f}"
        row[3].text = f"{r['max_abs_delta_pct']:.1f}"
        row[4].text = str(r["most_leveraging_district"])
        row[5].text = str(r["verdict"])
    _set_table_style(table)

    # Full LOO district detail
    doc.add_paragraph()
    _add_caption(doc,
                 "Table S5b.  Full leave-one-district-out detail: \u03c4\u0302 "
                 "and 95 % CI when each district is dropped.")

    cols2 = ["pipeline", "metric", "dropped_district", "exposure",
             "tau_loo", "se_loo", "p_loo",
             "ci_lo", "ci_hi", "delta_pct"]
    labs2 = ["Pipeline", "Metric", "Dropped", "Exposure",
             "\u03c4\u0302 LOO", "SE", "p",
             "CI\u2082.\u2085", "CI\u2089\u2087.\u2085", "\u0394 (%)"]
    table2 = doc.add_table(rows=1, cols=len(cols2))
    table2.style = "Light Grid Accent 1"
    for i, lab in enumerate(labs2):
        table2.rows[0].cells[i].text = lab
    _style_header(table2.rows[0])

    for _, r in jk_df.iterrows():
        if r["note"] != "ok":
            continue
        row = table2.add_row().cells
        for i, c in enumerate(cols2):
            v = r[c]
            if isinstance(v, float):
                row[i].text = f"{v:+.3f}" if c == "tau_loo" else f"{v:.3f}"
            else:
                row[i].text = str(v)
    _set_table_style(table2, font_size=9)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "Notes. Verdict legend: stable (max |\u0394\u03c4| < 25 % AND no sign "
        "flip), leverage (one district drives > 25 % of \u03c4\u0302), "
        "fragile (some LOO flips the sign of \u03c4\u0302). The Goodman-Bacon "
        "decomposition is not applicable (single-cohort design); LOO "
        "sensitivity is the binding leverage check."
    ).font.size = Pt(9)
    doc.save(out_path)


# ---------------------------------------------------------------------------
def table_s6_mde(mde_df: pd.DataFrame, out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    _add_caption(
        doc,
        "Table S6. Minimum detectable effect (MDE) at \u03b1 = 0.05, "
        "power = 0.80, df = G\u22121 = 7. Cluster-robust SE recovered "
        "from Module 05a (WCR). \u201cdetectable\u201d = yes when "
        "|\u03c4\u0302| \u2265 MDE_2sided.",
    )

    cols = ["pipeline", "metric", "tau_hat_d", "se_d",
            "MDE_2sided_d", "MDE_1sided_d", "tau_over_MDE", "detectable"]
    headers = ["Pipeline", "Metric", "\u03c4\u0302 (d)", "SE (d)",
               "MDE 2-sided", "MDE 1-sided", "|\u03c4\u0302|/MDE",
               "Detectable"]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h

    for _, r in mde_df.iterrows():
        row = table.add_row().cells
        for i, c in enumerate(cols):
            v = r[c]
            if isinstance(v, float):
                row[i].text = f"{v:.3f}" if c not in ("tau_over_MDE",) else f"{v:.2f}"
            else:
                row[i].text = str(v)

    _set_table_style(table, font_size=9)
    _style_header(table.rows[0])

    doc.add_paragraph().add_run(
        "Notes. MDE = (t_{\u03b1/2,G\u22121} + t_{1\u2212\u03b2,G\u22121}) \u00d7 SE. "
        "At G = 8, t_{0.025,7} = 2.365 and t_{0.80,7} = 0.896, so the "
        "multiplier is 3.261. The single non-detectable cell "
        "(corrected/EOS) is the same cell whose null is confirmed by "
        "the wild-cluster bootstrap (Table S4)."
    ).font.size = Pt(9)
    doc.save(out_path)


# ---------------------------------------------------------------------------
def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--results", default="analysis/results")
    ap.add_argument("--outdir",  default="manuscript/supplement")
    args = ap.parse_args()

    res_dir = Path(args.results)
    out_dir = Path(args.outdir)
    out_dir.mkdir(parents=True, exist_ok=True)

    static_df = pd.read_csv(res_dir / "did_static.csv")
    pre_df    = pd.read_csv(res_dir / "parallel_trends.csv")

    table_s1_did_static(static_df, out_dir / "Table_S1_did_static.docx")
    table_s2_pretrends (pre_df,    out_dir / "Table_S2_pretrends.docx")

    # Bulbul S3: regenerated by Module 05b (skip stub if real run exists)
    bulbul_csv = res_dir / "bulbul_transferability.csv"
    if not (out_dir / "Table_S3_bulbul_transferability.docx").exists():
        table_s3_bulbul_stub(out_dir / "Table_S3_bulbul_transferability.docx")

    # Wild-cluster bootstrap (Table S4)
    wcb_path = res_dir / "wild_bootstrap.csv"
    if wcb_path.exists():
        table_s4_wild_bootstrap(
            pd.read_csv(wcb_path),
            out_dir / "Table_S4_wild_bootstrap.docx",
        )
        print(f"wrote: {out_dir}/Table_S4_wild_bootstrap.docx")

    # LOO sensitivity (Table S5)
    jk_path = res_dir / "jackknife_district.csv"
    vd_path = res_dir / "jackknife_verdicts.csv"
    if jk_path.exists() and vd_path.exists():
        table_s5_jackknife(
            pd.read_csv(jk_path), pd.read_csv(vd_path),
            out_dir / "Table_S5_jackknife.docx",
        )
        print(f"wrote: {out_dir}/Table_S5_jackknife.docx")

    # MDE / power (Table S6)
    mde_path = res_dir / "power_mde.csv"
    if mde_path.exists():
        table_s6_mde(
            pd.read_csv(mde_path),
            out_dir / "Table_S6_mde.docx",
        )
        print(f"wrote: {out_dir}/Table_S6_mde.docx")

    # Also write consolidated CSVs for the GitHub release attachment
    static_df.assign(table="S1").to_csv(
        out_dir / "Table_S1_did_static.csv", index=False
    )
    pre_df.assign(table="S2").to_csv(
        out_dir / "Table_S2_pretrends.csv", index=False
    )

    print(f"wrote: {out_dir}/Table_S1_did_static.docx")
    print(f"wrote: {out_dir}/Table_S2_pretrends.docx")
    print(f"wrote: {out_dir}/Table_S1_did_static.csv")
    print(f"wrote: {out_dir}/Table_S2_pretrends.csv")


if __name__ == "__main__":
    main()
