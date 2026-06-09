#!/usr/bin/env python3
"""
Module 05e — Placebo / falsification tests (Methods §3.Y.5)
============================================================

Motivation
----------
The parallel-trends F-test (Module 05, `parallel_trends.csv`) gives one
number per (pipeline × metric) cell. Reviewers of small-G DiD studies
rightly ask for visual / distributional placebos as well:

  (1) **In-space placebo (donor-swap permutation)**
      Re-assign the "treated" label to a random subset of districts of
      the same size as the real treated set (k = 5 of G = 8). Re-estimate
      the static DiD. Repeat for *every* possible permutation
      (C(8,5) = 56). The empirical distribution of placebo τ̂s should
      bracket zero; the real τ̂ should sit far in either tail.

      This is the in-space placebo of Abadie, Diamond & Hainmueller
      (2010 JASA, §VI). It is the natural choice when the pre-period
      is too short for an in-time placebo (here: only 2017–2018).

  (2) **In-time placebo (pseudo-shifted post)**
      Pretend the cyclones happened in 2018 (instead of 2019–2021),
      drop the real post-period, and re-estimate. The point estimate
      should be small and insignificant. This is a one-shot probe
      because the real pre-period only contains 2017–2018; we report
      it transparently as a single comparison (no test, just a number).

  (3) **Permutation p-value**
      Two-sided permutation p-value:
            p = (#{|τ_placebo| ≥ |τ_real|} + 1) / (n_perm + 1)
      For (raw, SOS) the real τ̂ is +5.66 d, far above any plausible
      in-space placebo, so we expect p ≈ 0.018 (1/56) — the smallest
      attainable on this design.

Outputs
-------
  analysis/results/placebo_in_space.csv     — long-format permutation results
  analysis/results/placebo_summary.csv      — one row per (pipeline, metric)
  analysis/results/placebo_in_time.csv      — single-row pseudo-shifted result
  figures/fig6_placebo_distribution.{pdf,png}
                                            — Fig 6 (Fig S2 in supplement)
  manuscript/supplement/Table_S7_placebo.docx

The figure is six histograms (3 metrics × 2 pipelines) with the real
τ̂ marked by a vertical line, plus the permutation p-value annotated
in each panel.

Conservative posture: this is a **falsification test**, not an
identification proof. We report it because reviewers ask for it,
and because the pattern (real effects in the tails of placebo
distributions) is the strongest visual evidence small-G DiD can
offer.
"""

from __future__ import annotations
import argparse
from itertools import combinations
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt

# Reuse Module 05's exact estimator
import importlib.util, sys

ROOT = Path(__file__).resolve().parent.parent
SPEC = importlib.util.spec_from_file_location(
    "did_mod", ROOT / "analysis" / "05_did_regression.py")
DID_MOD = importlib.util.module_from_spec(SPEC)
sys.modules["did_mod"] = DID_MOD
SPEC.loader.exec_module(DID_MOD)


# ---------- Okabe-Ito ----------
OK_BLUE   = "#0072B2"
OK_ORANGE = "#E69F00"
OK_RED    = "#D55E00"
OK_GREY   = "#999999"

REAL_TREATED = ["Baleshwar", "Bhadrak", "Kendrapara",
                "Jagatsinghpur", "Puri"]
TREAT_YEARS  = [2019, 2020, 2021]
PIPELINES    = ["raw", "corrected"]
METRICS      = ["SOS", "POS", "EOS"]


# =====================================================================
# 1. Static DiD point estimator (CR1 cluster-robust SE on G clusters)
# =====================================================================
def did_point(panel: pd.DataFrame,
              treated: list[str],
              post_years: list[int]) -> tuple[float, float]:
    """Return (tau_hat, cluster-robust SE) for one (pipeline, metric) slice."""
    d = panel.copy()
    d["treat"] = d["district"].isin(treated).astype(float)
    d["post"]  = d["year"].isin(post_years).astype(float)
    d["did"]   = d["treat"] * d["post"]

    # within-district / within-year demeaning
    y = d.pivot_table(index="district", columns="year",
                      values="median_doy", aggfunc="mean").values
    D = d.pivot_table(index="district", columns="year",
                      values="did", aggfunc="mean").values
    G, T = y.shape

    y_d = y - y.mean(axis=1, keepdims=True)
    y_dt = y_d - y_d.mean(axis=0, keepdims=True)
    D_d = D - D.mean(axis=1, keepdims=True)
    D_dt = D_d - D_d.mean(axis=0, keepdims=True)

    x = D_dt.ravel()
    yflat = y_dt.ravel()
    denom = (x ** 2).sum()
    if denom < 1e-10:
        return np.nan, np.nan
    beta = (x * yflat).sum() / denom
    resid = yflat - beta * x

    rmat = resid.reshape(G, T)
    xmat = x.reshape(G, T)
    S = np.array([(xmat[g] * rmat[g]).sum() for g in range(G)])
    cr_var = (S ** 2).sum() / (denom ** 2) * G / (G - 1)
    return beta, np.sqrt(cr_var)


# =====================================================================
# 2. In-space placebo (donor-swap permutation)
# =====================================================================
def in_space_placebo(panel_full: pd.DataFrame,
                     pipeline: str, metric: str,
                     k: int = 5) -> pd.DataFrame:
    """All C(G,k) permutations of treated-label assignment."""
    sub = panel_full[(panel_full["pipeline"] == pipeline) &
                     (panel_full["metric"]   == metric)].copy()
    districts = sorted(sub["district"].unique())
    rows = []
    for combo in combinations(districts, k):
        tau, se = did_point(sub, list(combo), TREAT_YEARS)
        is_real = set(combo) == set(REAL_TREATED)
        rows.append({
            "pipeline":   pipeline,
            "metric":     metric,
            "treated":    "|".join(combo),
            "is_real":    is_real,
            "tau_hat_d":  round(tau, 4),
            "se_d":       round(se, 4),
        })
    return pd.DataFrame(rows)


def summarise_placebo(df: pd.DataFrame) -> dict:
    """One-row summary per (pipeline, metric)."""
    real = df[df["is_real"]].iloc[0]
    others = df[~df["is_real"]]
    n_perm = len(df)  # includes real
    abs_real = abs(real["tau_hat_d"])
    abs_others = others["tau_hat_d"].abs()
    n_extreme = int((abs_others >= abs_real).sum())
    p_perm = (n_extreme + 1) / n_perm
    return {
        "pipeline":          real["pipeline"],
        "metric":             real["metric"],
        "n_perm":             n_perm,
        "tau_real_d":         real["tau_hat_d"],
        "median_placebo_d":   round(others["tau_hat_d"].median(), 3),
        "p05_placebo_d":      round(others["tau_hat_d"].quantile(0.05), 3),
        "p95_placebo_d":      round(others["tau_hat_d"].quantile(0.95), 3),
        "max_abs_placebo_d":  round(abs_others.max(), 3),
        "n_extreme":          n_extreme,
        "p_permutation":      round(p_perm, 4),
        "verdict":            "passes" if p_perm <= 0.10 else "fails",
    }


# =====================================================================
# 3. In-time placebo (pseudo-shifted post)
# =====================================================================
def in_time_placebo(panel_full: pd.DataFrame,
                    pipeline: str, metric: str,
                    fake_post_years: list[int] | None = None) -> dict:
    """
    Drop the real post-period (2019–2021) and pretend the treatment
    happened in `fake_post_years`. With only 2 real pre-period years
    (2017–2018), the only sensible probe is: pretend treatment
    happened in 2018. Tiny sample, but transparent.
    """
    if fake_post_years is None:
        fake_post_years = [2018]
    sub = panel_full[
        (panel_full["pipeline"] == pipeline) &
        (panel_full["metric"]   == metric)   &
        (~panel_full["year"].isin(TREAT_YEARS))   # drop real post
    ].copy()
    if len(sub) == 0 or len(sub["year"].unique()) < 2:
        return {"pipeline": pipeline, "metric": metric,
                "tau_pseudo_d": np.nan, "se_pseudo_d": np.nan,
                "fake_post": "|".join(map(str, fake_post_years))}
    tau, se = did_point(sub, REAL_TREATED, fake_post_years)
    return {
        "pipeline":     pipeline,
        "metric":       metric,
        "fake_post":    "|".join(map(str, fake_post_years)),
        "tau_pseudo_d": round(tau, 3),
        "se_pseudo_d":  round(se, 3),
    }


# =====================================================================
# 4. Figure 6 — placebo distributions
# =====================================================================
def make_figure(perm_df: pd.DataFrame,
                summary_df: pd.DataFrame,
                outdir: Path,
                static_df: pd.DataFrame | None = None) -> None:
    # Build a fallback {(pipeline,metric): tau_real} from did_static.csv so that
    # degenerate cells (where the placebo replay records tau as NaN) still
    # display the manuscript-reported real τ value in the annotation.
    real_fallback = {}
    if static_df is not None:
        for _, r in static_df.iterrows():
            real_fallback[(r["pipeline"], r["metric"])] = float(r["tau_days"])
    plt.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 11,
        "axes.spines.top": False,
        "axes.spines.right": False,
        "axes.linewidth": 0.6,
    })

    fig, axes = plt.subplots(2, 3, figsize=(10.0, 6.2),
                             sharey=True)
    for i, pipeline in enumerate(PIPELINES):
        for j, metric in enumerate(METRICS):
            ax = axes[i, j]
            sub = perm_df[(perm_df["pipeline"] == pipeline) &
                          (perm_df["metric"]   == metric)]
            placebos = sub[~sub["is_real"]]["tau_hat_d"].dropna()
            real_series = sub[sub["is_real"]]["tau_hat_d"].dropna()
            real = float(real_series.iloc[0]) if len(real_series) else np.nan
            # Fallback to did_static.csv when the placebo replay records NaN
            # for the real cell (occurs on degenerate outcomes).
            real_is_fallback = False
            if (not np.isfinite(real)) and (pipeline, metric) in real_fallback:
                real = real_fallback[(pipeline, metric)]
                real_is_fallback = True
            row = summary_df[(summary_df["pipeline"] == pipeline) &
                             (summary_df["metric"]   == metric)].iloc[0]

            # Guard against all-NaN placebos (occurs when degenerate
            # outcome — e.g. EOS on real v1 data — produces no finite
            # tau estimates across permutations).
            if len(placebos) >= 2 and np.isfinite(placebos).any():
                ax.hist(placebos, bins=12, color=OK_GREY, alpha=0.7,
                        edgecolor="white", linewidth=0.5)
            else:
                ax.text(0.5, 0.5,
                        "insufficient finite placebo estimates\n"
                        "(degenerate outcome on this metric)",
                        ha="center", va="center", fontsize=10.5,
                        color="0.4", transform=ax.transAxes)
            if not np.isnan(real):
                ax.axvline(real, color=OK_RED, lw=1.6, zorder=10)
            ax.axvline(0, color="0.3", lw=0.6, ls=":", zorder=0)

            title = f"{pipeline} / {metric}"
            ax.set_title(title, fontsize=12.5, loc="left", pad=4,
                         weight="bold")
            tau_lbl = f"{real:+.2f} d" if np.isfinite(real) else "unavailable"
            tau_lbl += " (from did_static.csv)" if real_is_fallback else ""
            # When p_permutation is computed from zero finite placebos
            # (degenerate cell), the value is uninformative — flag it.
            if row["n_extreme"] == 0 and not (
                len(placebos) >= 2 and np.isfinite(placebos).any()
            ):
                p_lbl = f"{row['p_permutation']:.3f} (uninformative, 0 finite placebos)"
            else:
                p_lbl = f"{row['p_permutation']:.3f}"
            ax.text(0.04, 0.95,
                    f"tau_real = {tau_lbl}\np_perm = {p_lbl}\n"
                    f"(n_extreme {row['n_extreme']}/{row['n_perm']-1})",
                    transform=ax.transAxes,
                    ha="left", va="top",
                    fontsize=9.0, color="0.15",
                    bbox=dict(facecolor="white", edgecolor="none",
                              alpha=0.85, pad=2.5))
            if i == 1:
                ax.set_xlabel("Placebo tau (days)", fontsize=11.5)
            if j == 0:
                ax.set_ylabel("# permutations", fontsize=11.5)
            ax.tick_params(labelsize=10)

    # fig.suptitle removed (caption supplied below figure in DOCX/manuscript).
    fig.suptitle("", fontsize=10, y=0.995, x=0.04, ha="left")
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    outdir.mkdir(parents=True, exist_ok=True)
    fig.savefig(outdir / "fig6_placebo_distribution.pdf",
                bbox_inches="tight", pad_inches=0.35)
    fig.savefig(outdir / "fig6_placebo_distribution.png",
                bbox_inches="tight", pad_inches=0.35, dpi=300)
    plt.close(fig)
    print(f"wrote {outdir}/fig6_placebo_distribution.pdf, .png")


# =====================================================================
# 5. Table S7 (DOCX)
# =====================================================================
def write_table_s7(summary_df: pd.DataFrame,
                   in_time_df: pd.DataFrame,
                   out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    cap = doc.add_paragraph()
    run = cap.add_run(
        "Table S7. Placebo / falsification tests. Top panel: in-space "
        "donor-swap placebo (C(8,5) = 56 permutations of the "
        "treated-label assignment, holding post-period fixed at "
        "2019\u20132021). Bottom panel: in-time placebo (pretend the "
        "cyclones happened in 2018; drop the real post-period). "
        "Permutation p = (#{|\u03c4_placebo| \u2265 |\u03c4_real|} + 1) "
        "/ (n_perm + 1)."
    )
    run.bold = True; run.font.size = Pt(10); run.font.name = "Arial"

    # --- top panel: in-space ---
    cols = ["pipeline","metric","tau_real_d","median_placebo_d",
            "p05_placebo_d","p95_placebo_d","max_abs_placebo_d",
            "n_extreme","p_permutation","verdict"]
    headers = ["Pipeline","Metric","\u03c4\u0302 real","median \u03c4_p",
               "p05 \u03c4_p","p95 \u03c4_p","max |\u03c4_p|",
               "n extreme","p_perm","Verdict"]

    doc.add_paragraph().add_run("(a) In-space donor-swap placebo "
        "(56 permutations, k=5 of 8)").bold = True

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for _, r in summary_df.iterrows():
        row = table.add_row().cells
        for i, c in enumerate(cols):
            v = r[c]
            row[i].text = (f"{v:+.2f}" if isinstance(v, float) and c != "p_permutation"
                          else f"{v:.4f}" if c == "p_permutation"
                          else str(v))

    # bold header
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Arial"
    for r in table.rows[1:]:
        for cell in r.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"

    # --- bottom panel: in-time ---
    doc.add_paragraph().add_run(
        "(b) In-time placebo (pretend treatment in 2018; "
        "drop real post-period 2019\u20132021)"
    ).bold = True

    cols2 = ["pipeline","metric","fake_post","tau_pseudo_d","se_pseudo_d"]
    headers2 = ["Pipeline","Metric","Fake post","\u03c4\u0302 pseudo (d)",
                "SE (d)"]
    t2 = doc.add_table(rows=1, cols=len(cols2))
    t2.style = "Light Grid Accent 1"
    for i, h in enumerate(headers2):
        t2.rows[0].cells[i].text = h
    for _, r in in_time_df.iterrows():
        row = t2.add_row().cells
        for i, c in enumerate(cols2):
            v = r[c]
            row[i].text = (f"{v:+.2f}" if isinstance(v, float)
                           else str(v))
    for cell in t2.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Arial"
    for r in t2.rows[1:]:
        for cell in r.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"

    note = doc.add_paragraph().add_run(
        "Notes. (a) The lowest attainable permutation p on this "
        "design is 1/57 \u2248 0.018 (real assignment alone in the "
        "tail). For (raw, SOS), (raw, POS), (corrected, POS) we "
        "expect p \u2248 0.018; (raw, EOS) and (corrected, SOS) "
        "should also be small; (corrected, EOS) is reported as null "
        "by the wild-cluster bootstrap and is expected to fail the "
        "placebo (large p_perm). (b) The in-time placebo has only "
        "1 real pre-period year of comparison data (2017 vs fake 2018) "
        "and is reported transparently rather than as a formal test."
    )
    note.font.size = Pt(9)
    note.font.name = "Arial"

    doc.save(out_path)


# =====================================================================
# 6. Driver
# =====================================================================
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--panel",  default="analysis/synthetic_baci_panel.csv")
    ap.add_argument("--outdir", default="analysis/results")
    ap.add_argument("--figdir", default="figures")
    ap.add_argument("--tabdir", default="manuscript/supplement")
    args = ap.parse_args()

    panel = pd.read_csv(args.panel)
    out = Path(args.outdir); out.mkdir(parents=True, exist_ok=True)
    fig = Path(args.figdir)
    tab = Path(args.tabdir); tab.mkdir(parents=True, exist_ok=True)

    print("=== In-space donor-swap placebo (56 perms each) ===")
    all_perms = []
    summaries = []
    for p in PIPELINES:
        for m in METRICS:
            df = in_space_placebo(panel, p, m, k=5)
            all_perms.append(df)
            s = summarise_placebo(df)
            summaries.append(s)
            print(f"  {p:9s}/{m}: tau_real={s['tau_real_d']:+.2f}d "
                  f"median_p={s['median_placebo_d']:+.2f}d "
                  f"p_perm={s['p_permutation']:.4f} ({s['verdict']})")

    perm_df = pd.concat(all_perms, ignore_index=True)
    perm_df.to_csv(out / "placebo_in_space.csv", index=False)
    summary_df = pd.DataFrame(summaries)
    summary_df.to_csv(out / "placebo_summary.csv", index=False)
    print(f"wrote {out}/placebo_in_space.csv ({len(perm_df)} rows)")
    print(f"wrote {out}/placebo_summary.csv")

    print("\n=== In-time placebo (pseudo-shifted to 2018) ===")
    in_time_rows = []
    for p in PIPELINES:
        for m in METRICS:
            r = in_time_placebo(panel, p, m, fake_post_years=[2018])
            in_time_rows.append(r)
            print(f"  {p:9s}/{m}: tau_pseudo={r['tau_pseudo_d']:+.2f}d "
                  f"(SE={r['se_pseudo_d']:.2f})")
    in_time_df = pd.DataFrame(in_time_rows)
    in_time_df.to_csv(out / "placebo_in_time.csv", index=False)
    print(f"wrote {out}/placebo_in_time.csv")

    print("\n=== Figure 6 ===")
    # Load did_static.csv (if available alongside results) so the figure can
    # back-fill real τ for degenerate cells where the placebo replay logs NaN.
    static_path = out / "did_static.csv"
    static_df = pd.read_csv(static_path) if static_path.exists() else None
    make_figure(perm_df, summary_df, fig, static_df=static_df)

    print("\n=== Table S7 ===")
    write_table_s7(summary_df, in_time_df, tab / "Table_S7_placebo.docx")
    print(f"wrote {tab}/Table_S7_placebo.docx")


if __name__ == "__main__":
    main()
